# -*- coding: utf-8 -*-
"""生成 CLI运维管理系统 软著说明书 + 源码 DOCX (AGENT版)"""
import os, sys, io, textwrap
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

try:
    from PIL import Image, ImageDraw, ImageFont
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

BASE = os.path.dirname(os.path.abspath(__file__))
SCR = os.path.join(BASE, "screenshots")
OUT_MANUAL = os.path.join(os.path.dirname(BASE), "软著_CLI_软件说明书.docx")
OUT_SOURCE = os.path.join(os.path.dirname(BASE), "软著_CLI_源代码.docx")

# ─── screenshot map ───
IMG_MAP = {
    "1_help": "01_help.png", "2_status": "02_status.png",
    "3_config": "03_config_info.png", "4_health": "04_health.png",
    "5_scan": "05_scan.png", "6_exports": "06_list_exports.png",
    "7_rag": "07_check_rag.png", "8_monitor": "08_monitor_snap.png",
    "9_test": "09_test.png", "10_report": "10_report.png",
}

# ─── module files ───
MODULES = [
    "__init__.py", "config.py", "errors.py", "logger.py", "utils.py",
    "validators.py", "models.py", "cache_manager.py", "monitor.py",
    "export_engine.py", "preprocessor.py", "data_scanner.py",
    "health_checker.py", "report_generator.py", "batch_processor.py",
    "cli_main.py", "web/__init__.py", "web/app.py",
]


# ─── PIL diagram generators ───

def _draw_rounded_rect(draw, xy, fill, radius=8, outline=None):
    x1, y1, x2, y2 = xy
    draw.pieslice([x1, y1, x1 + radius * 2, y1 + radius * 2], 180, 270, fill=fill)
    draw.pieslice([x2 - radius * 2, y1, x2, y1 + radius * 2], 270, 360, fill=fill)
    draw.pieslice([x1, y2 - radius * 2, x1 + radius * 2, y2], 90, 180, fill=fill)
    draw.pieslice([x2 - radius * 2, y2 - radius * 2, x2, y2], 0, 90, fill=fill)
    draw.rectangle([x1 + radius, y1, x2 - radius, y1 + radius * 2], fill=fill)
    draw.rectangle([x1 + radius, y2 - radius * 2, x2 - radius, y2], fill=fill)
    draw.rectangle([x1, y1 + radius, x2, y2 - radius], fill=fill)
    if outline:
        draw.arc([x1, y1, x1 + radius * 2, y1 + radius * 2], 180, 270, fill=outline, width=2)
        draw.arc([x2 - radius * 2, y1, x2, y1 + radius * 2], 270, 360, fill=outline, width=2)
        draw.arc([x1, y2 - radius * 2, x1 + radius * 2, y2], 90, 180, fill=outline, width=2)
        draw.arc([x2 - radius * 2, y2 - radius * 2, x2, y2], 0, 90, fill=outline, width=2)
        draw.line([x1 + radius, y1, x2 - radius, y1], fill=outline, width=2)
        draw.line([x1 + radius, y2, x2 - radius, y2], fill=outline, width=2)
        draw.line([x1, y1 + radius, x1, y2 - radius], fill=outline, width=2)
        draw.line([x2, y1 + radius, x2, y2 - radius], fill=outline, width=2)


def _draw_arrow(draw, x1, y1, x2, y2, color=(88, 166, 255), width=2):
    draw.line([x1, y1, x2, y2], fill=color, width=width)
    mx, my = (x1 + x2) / 2, (y1 + y2) / 2
    angle = ((x2 - x1) ** 2 + (y2 - y1) ** 2) ** 0.5
    if angle > 0:
        dx, dy = (x2 - x1) / angle, (y2 - y1) / angle
        px, py = -dy * 8, dx * 8
        draw.polygon([(x2, y2), (x2 - dx * 12 + px, y2 - dy * 12 + py),
                       (x2 - dx * 12 - px, y2 - dy * 12 - py)], fill=color)


def generate_architecture_diagram():
    """生成AGENT架构图: AI AGENT → HTTP API → CLI Commands → System"""
    w, h = 720, 400
    img = Image.new("RGB", (w, h), (13, 17, 23))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 16)
        font_small = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 13)
        font_title = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 18)
    except Exception:
        font = ImageFont.load_default()
        font_small = font
        font_title = font

    # Title
    draw.text((w // 2, 16), "AI AGENT CLI 运维架构图", fill=(230, 237, 243), font=font_title, anchor="mt")

    # Layer 1: AI AGENT
    x1, y1 = w // 2 - 80, 60
    _draw_rounded_rect(draw, [x1, y1, x1 + 160, y1 + 50], fill=(45, 70, 120), radius=10)
    draw.text((x1 + 80, y1 + 25), "AI AGENT", fill=(88, 166, 255), font=font, anchor="mm")

    # Arrow down
    lx, ly = w // 2 - 100, y1 + 50
    _draw_rounded_rect(draw, [lx, ly, lx + 200, ly + 36], fill=(30, 40, 55), radius=6)
    draw.text((lx + 100, ly + 18), "远程 Hook · HTTP 调用", fill=(139, 148, 158), font=font_small, anchor="mm")
    _draw_arrow(draw, w // 2, y1 + 50, w // 2, ly)

    # Layer 2: Web API (RESTful)
    y2 = ly + 50
    _draw_rounded_rect(draw, [w // 2 - 120, y2, w // 2 + 120, y2 + 50], fill=(13, 65, 146), radius=10)
    draw.text((w // 2, y2 + 25), "RESTful API 服务", fill=(88, 166, 255), font=font, anchor="mm")
    _draw_arrow(draw, w // 2, ly + 36, w // 2, y2)

    # Layer 3: CLI Commands (3 boxes)
    y3 = y2 + 70
    cmd_boxes = [
        (w // 4 - 40, "scan/export\n数据操作"),
        (w // 2 - 80, "health/monitor\n系统监控"),
        (3 * w // 4 - 120, "report/batch\n报告与批处理"),
    ]
    for cx, clabel in cmd_boxes:
        _draw_rounded_rect(draw, [cx, y3, cx + 160, y3 + 60], fill=(18, 30, 45), radius=8, outline=(48, 54, 61))
        lines = clabel.split("\n")
        for li, line in enumerate(lines):
            draw.text((cx + 80, y3 + 18 + li * 22), line, fill=(201, 209, 217), font=font_small, anchor="mm")

    # Fan-out arrows from API to CLI boxes
    for cx, _ in cmd_boxes:
        _draw_arrow(draw, w // 2, y2 + 50, cx + 80, y3)

    # Layer 4: System Resources
    y4 = y3 + 80
    sys_boxes = [
        (w // 6, "数据目录"),
        (w // 2 - 40, "日志/缓存"),
        (5 * w // 6 - 80, "导出/报告"),
    ]
    for sx, slabel in sys_boxes:
        _draw_rounded_rect(draw, [sx, y4, sx + 100, y4 + 40], fill=(13, 17, 23), radius=6, outline=(48, 54, 61))
        draw.text((sx + 50, y4 + 20), slabel, fill=(139, 148, 158), font=font_small, anchor="mm")

    for i, (cx, _) in enumerate(cmd_boxes):
        sx = [w // 6, w // 2 - 40, 5 * w // 6 - 80][i]
        _draw_arrow(draw, cx + 80, y3 + 60, sx + 50, y4)

    # Layer 5: End users
    y5 = y4 + 60
    _draw_rounded_rect(draw, [w // 4 - 40, y5, w // 4 + 100, y5 + 36], fill=(25, 45, 30), radius=6)
    draw.text((w // 4 + 10, y5 + 18), "人类运维人员", fill=(63, 185, 80), font=font_small, anchor="mm")
    _draw_rounded_rect(draw, [w // 4 + 160, y5, w // 4 + 300, y5 + 36], fill=(25, 30, 50), radius=6)
    draw.text((w // 4 + 230, y5 + 18), "AI AGENT (自动调用)", fill=(88, 166, 255), font=font_small, anchor="mm")
    _draw_arrow(draw, w // 2, y2 + 50, w // 4 + 10, y5)
    _draw_arrow(draw, w // 2, y2 + 50, w // 4 + 230, y5)

    buf = io.BytesIO()
    img.save(buf, "PNG")
    buf.seek(0)
    return buf


def generate_api_diagram():
    """生成API结构图: 三层级API"""
    w, h = 600, 320
    img = Image.new("RGB", (w, h), (13, 17, 23))
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 16)
        font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 13)
    except Exception:
        font_title = ImageFont.load_default()
        font = font_title

    draw.text((w // 2, 14), "RESTful API 结构", fill=(230, 237, 243), font=font_title, anchor="mt")

    tiers = [
        ("/api/v1/cli/*", 50, (45, 70, 120), (88, 166, 255), [
            "POST /cli/run   远程执行命令",
            "GET  /cli/health 健康检查",
            "GET  /cli/status 系统状态",
            "GET  /cli/history 操作历史",
        ]),
        ("/api/v1/data/*", 150, (13, 65, 146), (88, 166, 255), [
            "POST /data/scan   数据扫描",
            "POST /data/export 数据导出",
            "GET  /data/info   目录信息",
        ]),
        ("/api/v1/system/*", 240, (25, 45, 30), (63, 185, 80), [
            "GET /system/info   系统信息",
            "GET /system/health 公开健康检查",
        ]),
    ]

    for title, y_start, bg_color, tx_color, items in tiers:
        _draw_rounded_rect(draw, [20, y_start, w - 20, y_start + 38], fill=bg_color, radius=6)
        draw.text((w // 2, y_start + 19), title, fill=tx_color, font=font, anchor="mm")
        for li, item in enumerate(items):
            draw.text((40, y_start + 42 + li * 18), item, fill=(201, 209, 217), font=font)

    buf = io.BytesIO()
    img.save(buf, "PNG")
    buf.seek(0)
    return buf


def generate_dashboard_mockup():
    """生成Web看板模拟截图"""
    w, h = 640, 420
    img = Image.new("RGB", (w, h), (13, 17, 23))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 14)
        font_s = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 12)
        font_xs = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 10)
    except Exception:
        font = ImageFont.load_default()
        font_s = font
        font_xs = font

    # Navbar
    _draw_rounded_rect(draw, [0, 0, w, 40], fill=(22, 27, 34), radius=0)
    draw.text((20, 20), "CLI 运维管理  |  主页  控制台  退出", fill=(139, 148, 158), font=font_s, anchor="lm")

    # Hero area
    _draw_rounded_rect(draw, [15, 50, w - 15, 130], fill=(13, 40, 60), radius=8)
    draw.text((w // 2, 70), "AI AGENT 驱动的运维管理系统", fill=(230, 237, 243), font=font, anchor="mt")
    draw.text((w // 2, 95), "CLI命令行可直接被AGENT操控 · 远程hook实现AI提效", fill=(150, 180, 210), font=font_xs, anchor="mt")

    # 3 stat cards
    stats = [("系统版本", "V1.0"), ("健康状态", "正常"), ("CLI命令", "16")]
    for i, (label, val) in enumerate(stats):
        cx = 30 + i * 200
        _draw_rounded_rect(draw, [cx, 140, cx + 180, 200], fill=(22, 27, 34), radius=6, outline=(48, 54, 61))
        draw.text((cx + 90, 160), val, fill=(88, 166, 255), font=font, anchor="mm")
        draw.text((cx + 90, 185), label, fill=(139, 148, 158), font=font_xs, anchor="mm")

    # Table: CLI commands
    yt = 215
    draw.text((20, yt), "CLI 命令列表 (AGENT可操控)", fill=(230, 237, 243), font=font_s, anchor="la")
    cols = ["命令", "功能", "状态"]
    col_x = [20, 180, 400]
    for j, (c, cx) in enumerate(zip(cols, col_x)):
        draw.text((cx, yt + 25), c, fill=(139, 148, 158), font=font_xs, anchor="la")
    draw.line([(15, yt + 40), (w - 15, yt + 40)], fill=(48, 54, 61), width=1)
    commands_data = [
        ("scan", "数据目录扫描", "OK"), ("export", "多格式导出", "OK"),
        ("monitor-snap", "系统资源快照", "OK"), ("health", "健康检查", "OK"),
        ("report", "运维报告生成", "OK"), ("clear-cache", "缓存清理", "OK"),
    ]
    for li, (cmd, desc, status) in enumerate(commands_data):
        yy = yt + 48 + li * 22
        draw.text((col_x[0], yy), cmd, fill=(201, 209, 217), font=font_xs, anchor="la")
        draw.text((col_x[1], yy), desc, fill=(139, 148, 158), font=font_xs, anchor="la")
        draw.text((col_x[2], yy), "● 就绪", fill=(63, 185, 80), font=font_xs, anchor="la")

    buf = io.BytesIO()
    img.save(buf, "PNG")
    buf.seek(0)
    return buf


# ─── DOCX helpers ───

def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5
    for i in range(1, 4):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "黑体"
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        h.font.color.rgb = RGBColor(0, 0, 0)


def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    doc.add_paragraph()


def add_img(doc, num, title, description, custom_buf=None):
    p = doc.add_paragraph()
    r = p.add_run(f"【图{num}】 {title}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    if custom_buf:
        doc.add_picture(custom_buf, width=Inches(5.8))
        return

    fname = IMG_MAP.get(f"{num}_{title.split('_')[0]}")
    if not fname:
        for k, v in IMG_MAP.items():
            if k.startswith(str(num)):
                fname = v
                break
    if fname:
        fpath = os.path.join(SCR, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.8))
            return
    # fallback
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    r2 = p2.add_run(description)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ─── Chapter: Design Philosophy ───

def add_design_philosophy(doc):
    """1.4 设计初衷 — AGENT驱动的CLI架构"""
    doc.add_heading("1.4 设计初衷 — AGENT驱动的CLI架构", level=2)
    doc.add_paragraph(
        "本系统的核心设计理念围绕'模型发展趋势'展开。随着大语言模型（LLM）和 AI AGENT 技术的快速发展，"
        "传统的图形用户界面（GUI）已不再是人与系统交互的唯一方式。AI AGENT 能够理解自然语言指令、"
        "自主规划任务步骤、调用工具完成复杂操作。在这一趋势下，CLI（命令行界面）因其结构化、可编程、"
        "易解析的特性，重新成为人机交互和机机交互的核心范式。"
    )
    doc.add_paragraph(
        "基于模型发展趋势，使用CLI命令行可以让AGENT直接操控系统全部功能，实现远程hook的AI提效。"
        "具体而言，本系统在设计上贯彻了以下三个核心原则："
    )
    doc.add_paragraph(
        "一、CLI优先设计。系统的16个子命令均采用标准化的输入输出格式，每个命令接收明确的参数、"
        "执行确定的功能、输出结构化的结果（支持Text/JSON/HTML三种格式）。这种设计使AGENT无需"
        "理解复杂的GUI语义，只需构造正确的命令字符串即可完成数据扫描、导出、监控、报告等全部操作。"
    )
    doc.add_paragraph(
        "二、远程Hook机制。系统内置RESTful API层（/api/v1/cli/* 端点），将CLI命令封装为HTTP端点。"
        "AGENT通过API代理模式远程触发CLI命令执行，返回结构化结果。这一机制使得AGENT可以在任何网络可达的位置"
        "操控系统，无需直接登录服务器。Hook端点支持任务队列、异步执行、结果回调等高级模式，满足AGENT工作流编排需求。"
    )
    doc.add_paragraph(
        "三、人类+AI双模式。Web界面提供人类可读的运维看板和操作界面（登录注册、控制台、系统状态可视化），"
        "RESTful API提供AGENT可调用的程序化接口。同一套CLI底层能力同时服务于人类操作员和AI AGENT，"
        "既可人工逐条执行命令，也可由AGENT 7×24小时自动化运维，最大限度释放运维效率。"
    )
    doc.add_paragraph(
        '这一设计使本系统区别于传统运维工具：它不是\'给人看的CLI加个Web壳\'，而是原生面向AGENT时代设计的'
        '"API-first + CLI-native"运维基础设施。无论未来GUI形态如何演变，AGENT始终可以通过CLI和API直接操控系统。'
    )


# ─── Chapter: Web Module ───

def add_web_module(doc):
    """四、Web前端模块"""
    doc.add_heading("四、Web前端模块", level=1)

    doc.add_heading("4.1 登录与注册", level=2)
    doc.add_paragraph(
        "Web前端提供完整的用户认证体系，包括登录和注册功能。用户通过注册页面创建账号（用户名、密码、确认密码），"
        "系统对密码进行SHA-256哈希存储。登录后使用Flask Session管理用户会话。未登录用户访问受保护页面时自动跳转"
        "到登录页。登录页面在底部标注设计理念：'设计理念：CLI → AGENT 可直接操控 · 远程 hook · AI 提效'。"
    )
    if HAS_PIL:
        buf = generate_dashboard_mockup()
        add_img(doc, "11", "Web登录界面与运维看板", "Web前端界面截图（模拟）", custom_buf=buf)
    else:
        add_img(doc, "11", "Web登录界面", "登录界面包含用户名/密码输入、注册链接、设计理念标注", custom_buf=None)

    doc.add_heading("4.2 运维看板 (Dashboard)", level=2)
    doc.add_paragraph(
        "控制台页面以暗色主题展示系统运维信息：顶部导航栏包含主页、控制台、退出链接；三列统计卡片显示系统版本"
        "（V1.0）、健康状态（通过/降级/失败）和CLI子命令总数（16个）；健康检查模块展示8项自动化检查的详细结果"
        "（检查项名称、通过/失败状态、详情消息）；系统目录模块列出7个核心目录路径（data_dir/export_dir/log_dir/"
        "cache_dir/report_dir等）；CLI命令一览表列出全部16个命令及其功能描述，每行标注'AGENT可操控'的状态标识。"
        "页面底部再次展示设计初衷文字：'基于模型发展趋势，使用CLI命令行可以让AGENT直接操控，实现远程hook的AI提效'。"
    )

    doc.add_heading("4.3 主页 (Home)", level=2)
    doc.add_paragraph(
        "主页同样采用暗色主题，页面顶部为渐变色Hero区域，突出显示系统的设计理念：'基于模型发展趋势，本系统以CLI"
        "命令行作为核心交互方式，使AI AGENT可直接操控系统全部功能。通过远程hook机制实现AI提效，无需人工逐条输入"
        "命令。Web界面提供人类可读的运维看板和RESTful API，AI AGENT可通过API直接调用CLI底层能力，形成人类+AI"
        "双模式的运维工作流。' Hero区域下方为三张功能卡片（数据扫描与监控、多格式数据导出、报告与健康检查），"
        "快捷操作按钮区（数据扫描、监控快照、健康检查、生成报告、清除缓存），以及API接口列表区。"
    )

    doc.add_heading("4.4 三层级RESTful API", level=2)
    doc.add_paragraph(
        "系统提供三层级RESTful API，分别对应CLI代理、数据管理和系统信息。所有API（除公开健康检查外）均需认证。"
    )
    add_table(doc, ["层级", "端点", "方法", "功能", "认证"],
        [
            ["CLI代理", "/api/v1/cli/run", "POST", "远程执行CLI命令, AGENT调用入口", "是"],
            ["CLI代理", "/api/v1/cli/health", "GET", "健康检查, AGENT可调用", "是"],
            ["CLI代理", "/api/v1/cli/status", "GET", "系统状态概览", "是"],
            ["CLI代理", "/api/v1/cli/history", "GET", "操作历史查询", "是"],
            ["数据管理", "/api/v1/data/scan", "POST", "数据目录扫描", "是"],
            ["数据管理", "/api/v1/data/export", "POST", "多格式数据导出", "是"],
            ["数据管理", "/api/v1/data/info", "GET", "数据目录信息", "是"],
            ["系统信息", "/api/v1/system/info", "GET", "系统信息（含设计哲学描述）", "是"],
            ["系统信息", "/api/v1/system/health", "GET", "公开健康检查（免认证）", "否"],
        ])
    doc.add_paragraph(
        "其中 /api/v1/cli/run 是AGENT远程操控的核心接口，接收command和args参数，将CLI命令封装为任务（含task_id），"
        "加入执行队列后返回状态信息。/api/v1/system/info 返回的系统信息中包含设计哲学字段（design_philosophy），"
        "明确标注'CLI设计使AGENT可直接操控，通过远程hook实现AI提效'。"
    )
    if HAS_PIL:
        buf = generate_api_diagram()
        add_img(doc, "12", "RESTful API三层级结构", "API结构示意图", custom_buf=buf)
    else:
        add_img(doc, "12", "API结构", "三层级API结构示意", custom_buf=None)

    doc.add_heading("4.5 AGENT远程调用流程", level=2)
    doc.add_paragraph(
        "AI AGENT通过以下流程远程操控CLI系统：(1) AGENT调用POST /api/v1/cli/run 接口发送命令和参数；"
        "(2) 系统接收请求，生成唯一task_id，将命令加入执行队列，立即返回202 Accepted状态；"
        "(3) AGENT轮询GET /api/v1/cli/status 或查询操作历史 GET /api/v1/cli/history 获取命令执行结果；"
        "(4) 系统执行完成后，AGENT获取结构化输出结果（JSON格式），可直接用于后续分析或决策。"
        "这一流程使得AGENT可以将CLI运维能力嵌入到自主工作流中，例如：每日自动执行健康检查→分析结果→"
        "发现磁盘不足→自动执行日志清理→生成运维报告→发送摘要通知。全程无需人工干预。"
    )
    doc.add_page_break()


# ─── Main Builder ───

def build_manual():
    """软著申请书标准说明书 — 目标 22-26 页"""
    doc = Document()
    setup_styles(doc)

    # ── Cover ──
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("唐诗意象数据运维管理系统"); r.font.size = Pt(30); r.bold = True
    r.font.name = "黑体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("软件说明书"); r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("设计初衷：基于模型发展趋势，使用CLI命令行可以让AGENT直接操控，实现远程hook的AI提效")
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)
    r.font.name = "黑体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    doc.add_paragraph()
    for text in ["版本：V1.0", "开发完成日期：2026年5月", "类型：命令行工具（CLI）+ Web前端",
                 "开发语言：Python 3.10+", "总代码行数：约 6,000 行（18 个模块）"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.size = Pt(14)
    doc.add_page_break()

    # ── TOC ──
    doc.add_heading("目录", level=1)
    for item in ["一、引言", "    1.1 项目背景", "    1.2 软件功能概述", "    1.3 运行环境",
                 "    1.4 设计初衷 — AGENT驱动的CLI架构",
                 "二、系统架构", "    2.1 总体架构", "    2.2 模块清单",
                 "三、核心功能模块与操作流程",
                 "四、Web前端模块",
                 "    4.1 登录与注册", "    4.2 运维看板", "    4.3 主页",
                 "    4.4 三层级RESTful API", "    4.5 AGENT远程调用流程",
                 "五、数据处理流程",
                 "六、部署与安全",
                 "七、测试覆盖",
                 "八、文件清单"]:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ====================================================================
    # 一、引言
    # ====================================================================
    doc.add_heading("一、引言", level=1)

    doc.add_heading("1.1 项目背景", level=2)
    doc.add_paragraph(
        "古典文学数字化研究涉及大量结构化标注数据的管理工作。唐诗意象数据通常以 JSON 格式存储多层嵌套结构"
        "（诗歌编号、作者、诗行、分析单元、情感轨迹等），在实际研究和运维过程中，需要频繁进行数据目录扫描、"
        "多格式导出、系统资源监控、健康状态检查及运维报告生成。传统人工操作方式效率低且易出错，迫切需要"
        "一套自动化的命令行运维工具。本系统独立于 Web 前端，以纯命令行形式运行，可集成到自动化脚本、"
        "定时任务和 CI/CD 流水线中，为古典文学数据管理提供高效可靠的运维支撑。"
    )

    doc.add_heading("1.2 软件功能概述", level=2)
    doc.add_paragraph(
        "唐诗意象数据运维管理系统是一个基于 Python 3.10+ 开发的命令行运维管理工具集。系统围绕古典文学"
        "标注数据的全生命周期管理需求，提供了 16 个功能完备的子命令，涵盖数据扫描、多格式导出、系统资源"
        "监控、双层缓存管理、自动化健康检查、运维报告生成、批量数据处理、数据预处理与备份等核心功能模块。"
        "系统遵循'最小依赖'原则，核心功能完全基于 Python 标准库实现，无需安装第三方包即可运行。"
    )
    doc.add_paragraph("本系统核心功能如下：")
    features = [
        "数据目录扫描（scan）：递归扫描指定目录，生成完整文件清单，支持按扩展名、文件大小、文件年龄三维分类统计，提供变更检测功能。",
        "多格式数据导出（export）：支持 CSV (UTF-8-BOM)、JSON (双层结构)、XML、TXT (TSV)、HTML 五种标准格式，支持字段过滤与重命名、数据子集切片、行数上限控制。",
        "系统资源监控（monitor-snap）：实时采集 CPU/内存/磁盘/网络/进程五类数据，后台定时采集，阈值告警回调，0-100综合健康度评分。",
        "双层缓存管理（clear-cache）：内存缓存（LRU/LFU/TTL淘汰策略）+ 文件缓存（磁盘JSON持久化），@cached 装饰器，命中率统计。",
        "系统健康检查（health）：8 项自动化检查（Python环境、磁盘空间、配置有效性、各目录状态等），依赖拓扑排序，失败重试。",
        "运维报告生成（report）：Text/JSON/HTML 三格式，六大章节（系统信息/目录/资源/健康/缓存/导出文件）。",
        "批量数据处理（batch）：多线程并发执行，分块控制，超时控制，指数退避重试，实时进度回调。",
        "数据预处理（preprocessor + backup）：JSON清洗、结构校验、批量目录验证、数据备份。",
        "Web前端管理：登录注册认证、运维控制台看板、系统主页、三层级RESTful API（CLI代理/数据管理/系统信息）。",
        "远程AGENT调用：/api/v1/cli/run 端点允许AI AGENT远程执行CLI命令，任务队列管理，异步执行。",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("1.3 运行环境", level=2)
    add_table(doc, ["项目", "要求", "备注"], [
        ["操作系统", "Windows 10+ / Linux / macOS", "跨平台"],
        ["Python 版本", "Python 3.10+", "利用 dataclass/match 语法"],
        ["必需依赖", "无（纯标准库）", "核心功能零第三方依赖"],
        ["可选依赖", "psutil, chromadb", "增强监控和向量库功能"],
    ])

    # 1.4 Design Philosophy (NEW)
    add_design_philosophy(doc)
    doc.add_page_break()

    # ====================================================================
    # 二、系统架构
    # ====================================================================
    doc.add_heading("二、系统架构", level=1)

    doc.add_heading("2.1 总体架构", level=2)
    doc.add_paragraph("系统采用四层模块化架构，各层单向依赖：")
    doc.add_paragraph(
        "命令入口层（cli_main.py）：基于 argparse 实现 16 个子命令的注册、路由和参数解析。"
        "业务逻辑层（6 个模块）：monitor.py（系统监控）、export_engine.py（数据导出）、"
        "cache_manager.py（缓存管理）、health_checker.py（健康检查）、report_generator.py（报告生成）、"
        "batch_processor.py（批量处理）。"
        "数据模型层（models.py）：8 个 dataclass 模型，统一序列化和数据传递。"
        "基础设施层（6 个模块）：config.py（配置管理）、errors.py（异常体系）、logger.py（日志系统）、"
        "utils.py（30+通用函数）、validators.py（14个校验函数）、preprocessor.py（数据预处理）。"
    )
    doc.add_paragraph(
        "Web前端层（web/模块）：基于 Flask 框架提供登录注册、控制台看板、主页展示和三层级 RESTful API。"
        "Web层通过API代理模式与CLI底层通信，AI AGENT可通过HTTP端点远程操控CLI命令。"
    )

    if HAS_PIL:
        buf = generate_architecture_diagram()
        add_img(doc, "0", "system_architecture", "AGENT CLI 运维架构图", custom_buf=buf)
    else:
        p = doc.add_paragraph()
        p.add_run("AGENT → HTTP API → CLI Commands → System Resources").font.size = Pt(10)

    doc.add_heading("2.2 模块清单", level=2)
    add_table(doc, ["序号", "模块名", "文件名", "行数", "职责"], [
        ["1", "CLI主入口", "cli_main.py", "725", "16子命令路由、argparse参数解析"],
        ["2", "配置管理", "config.py", "286", "40+配置项、多环境、ConfigManager类"],
        ["3", "异常体系", "errors.py", "407", "7层40+异常类、1000-7099错误码"],
        ["4", "日志模块", "logger.py", "209", "分级日志、RotatingFileHandler"],
        ["5", "工具函数", "utils.py", "500", "30+函数：字符串/文件/JSON/统计"],
        ["6", "校验模块", "validators.py", "224", "14个函数：路径/格式/范围"],
        ["7", "数据模型", "models.py", "450", "8个dataclass模型、序列化"],
        ["8", "缓存管理", "cache_manager.py", "391", "内存+文件双层缓存"],
        ["9", "系统监控", "monitor.py", "365", "CPU/内存/磁盘/网络/进程"],
        ["10", "导出引擎", "export_engine.py", "519", "5格式导出+历史管理"],
        ["11", "数据预处理", "preprocessor.py", "375", "JSON清洗/校验/备份"],
        ["12", "目录扫描", "data_scanner.py", "325", "递归扫描/分类/变更检测"],
        ["13", "健康检查", "health_checker.py", "341", "8项检查、注册表"],
        ["14", "报告生成", "report_generator.py", "306", "Text/JSON/HTML三格式"],
        ["15", "批量处理", "batch_processor.py", "316", "并发/超时/重试"],
        ["16", "Web前端入口", "web/app.py", "286", "Flask应用、登录注册、API"],
        ["17", "Web包初始化", "web/__init__.py", "2", "Web模块声明"],
        ["18", "包初始化", "__init__.py", "47", "版本声明、模块导出"],
    ])
    doc.add_page_break()

    # ====================================================================
    # 三、核心功能模块与操作流程
    # ====================================================================
    doc.add_heading("三、核心功能模块与操作流程", level=1)
    doc.add_paragraph("本章节展示系统核心子命令的操作方式和终端运行效果。所有截图均为实际执行结果的命令行界面截图。")

    sections = [
        ("3.1 帮助信息 (help)", "1_help",
         "执行 python -m tang_cli_ops.cli_main help，显示全部 16 个命令的名称、功能描述和使用示例。"
         "是了解系统功能的第一入口。", "help"),
        ("3.2 系统状态总览 (status)", "2_status",
         "执行 python -m tang_cli_ops.cli_main status，显示版本信息、7 个目录状态、缓存统计（命中率/条目数）、"
         "系统资源使用（CPU/内存/磁盘）、最近导出文件列表。", "status"),
        ("3.3 配置信息查询 (config-info)", "3_config",
         "执行 python -m tang_cli_ops.cli_main config-info，列出 30+ 个可配置项及其当前运行时值。"
         "用户可通过环境变量 TCO_<KEY> 覆盖任何配置项。", "config_info"),
        ("3.4 系统健康检查 (health)", "4_health",
         "执行 python -m tang_cli_ops.cli_main health，按依赖顺序执行 8 项检查并输出通过/失败状态、"
         "修复建议和整体健康判定。", "health"),
        ("3.5 数据目录扫描 (scan)", "5_scan",
         "执行 python -m tang_cli_ops.cli_main scan，递归扫描并生成文件清单，按扩展名/大小/年龄三维分类。", "scan"),
        ("3.6 导出文件管理 (list-exports)", "6_exports",
         "列出所有历史导出文件，便于管理导出历史和清理过期文件。", "list_exports"),
        ("3.7 向量数据库检查 (check-rag)", "7_rag",
         "检查 ChromaDB 向量数据库状态，包括目录、文件数量、总大小。", "check_rag"),
        ("3.8 系统监控快照 (monitor-snap)", "8_monitor",
         "采集系统资源即时快照：磁盘/内存/CPU/进程信息，显示历史趋势对比。", "monitor_snap"),
        ("3.9 单元测试 (test)", "9_test",
         "运行模块导入测试和核心函数功能测试，输出 PASS/FAIL 状态。", "test"),
        ("3.10 运维报告生成 (report)", "10_report",
         "自动采集系统全维度数据并生成 Text/JSON/HTML 三格式报告。", "report"),
    ]

    for title, num, desc, img_key in sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        add_img(doc, num.split("_")[0], img_key, f"【截图】{title}", custom_buf=None)

    # remaining commands
    doc.add_heading("3.11 其余功能命令", level=2)
    doc.add_paragraph("除上述核心命令外，系统还提供以下 6 个功能命令：")
    remaining = [
        "数据导出 (export)：支持 CSV/JSON/XML/TXT/HTML 五种格式。--fields 参数控制导出字段，--rows 控制最大行数。",
        "缓存清理 (clear-cache)：一键清除内存缓存和文件缓存。--force/-y 跳过确认提示。",
        "向量库构建 (build-rag)：扫描数据目录中的 JSON 文件，递归提取诗歌节点，初始化 ChromaDB。",
        "日志清理 (clean-logs)：清理超过指定天数（默认30天）的旧日志文件。",
        "数据备份 (backup)：使用 shutil.copy2 将数据目录完整备份到 backups/ 目录下带时间戳的子目录。",
        "批处理 (batch)：基于 ThreadPoolExecutor 并发引擎，支持分块执行、任务超时控制、失败重试。",
    ]
    for r in remaining:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_page_break()

    # ====================================================================
    # 四、Web前端模块 (NEW)
    # ====================================================================
    add_web_module(doc)

    # ====================================================================
    # 五、数据处理流程
    # ====================================================================
    doc.add_heading("五、数据处理流程", level=1)

    doc.add_heading("5.1 目录扫描流程", level=2)
    doc.add_paragraph(
        "用户执行 scan 命令 → argparse 解析参数 → scan_directory() 根据 recursive 标志选择 os.walk（递归）"
        "或 os.listdir（非递归）遍历目录树。每个文件经过以下处理：跳过隐藏文件 → 检查文件大小是否超过阈值"
        "（500MB，超过则计入 skipped_count）→ get_file_info() 获取元信息 → 追加到 ScanResult.files 列表 → "
        "累加 total_size。扫描完成后调用 classify_files_by_type/size/age 进行三维分类，get_scan_summary() 生成汇总。"
    )

    doc.add_heading("5.2 数据导出流程", level=2)
    doc.add_paragraph(
        "用户执行 export 命令 → 解析 --format 参数确定导出格式 → _load_sample_data() 从 DATA_DIR 加载源数据"
        " → clean_json_content 清洗 → json.loads 解析 → _extract_poems 递归提取诗歌节点 → 遍历分析单元提取"
        "意象级字段 → 构建扁平化行数据字典列表。格式特定的导出函数执行编码、格式化、文件写入。"
    )

    doc.add_heading("5.3 监控采集流程", level=2)
    doc.add_paragraph(
        "monitor-snap 命令或后台定时线程触发 → SystemMonitor.snapshot() 依次采集五类数据："
        "磁盘（os.statvfs）、内存（psutil.virtual_memory）、CPU（psutil.cpu_percent）、"
        "网络（psutil.net_io_counters）、进程（os.getpid + psutil.Process）。"
        "health_score() 按加权公式（磁盘30+内存30+CPU20+稳定性20）计算综合评分。"
    )

    doc.add_heading("5.4 报告生成流程", level=2)
    doc.add_paragraph(
        "report 命令触发 → gather_report_data() 协调六个数据采集函数 → 构建统一报告数据字典 → "
        "根据 --format 参数选择渲染器（Text/JSON/HTML）→ save_report() 写入 reports/ 目录。"
    )

    doc.add_heading("5.5 AGENT远程操作流程", level=2)
    doc.add_paragraph(
        "AI AGENT 调用 POST /api/v1/cli/run 接口 → 系统接收命令参数 → 生成 task_id → 加入执行队列 → "
        "返回 task_id 和状态。AGENT通过轮询 GET /api/v1/cli/history 获取执行结果。"
        "典型自动化场景：每日健康检查 → 分析结果 → 发现磁盘不足 → 自动清理日志 → 生成报告 → 发送通知。"
    )
    doc.add_page_break()

    # ====================================================================
    # 六、部署与安全
    # ====================================================================
    doc.add_heading("六、部署与安全", level=1)

    doc.add_heading("6.1 环境要求与安装", level=2)
    doc.add_paragraph(
        "本系统支持 Windows 10+、Linux（推荐 Ubuntu 22.04+）、macOS 等主流操作系统。Python 版本要求 3.10+。"
        "核心功能完全基于标准库实现，不需要安装任何第三方包。"
        "可选依赖：psutil（增强系统监控），chromadb（向量数据库操作）。"
        "Web前端需额外安装 Flask：pip install flask。"
    )
    doc.add_paragraph("安装步骤：")
    doc.add_paragraph("1. 将 tang_cli_ops 目录放置到目标服务器的项目根目录下。")
    doc.add_paragraph("2. （可选）pip install psutil chromadb flask")
    doc.add_paragraph("3. 验证安装：python -m tang_cli_ops.cli_main help")
    doc.add_paragraph("4. 启动Web服务：python -m tang_cli_ops.web.app")
    doc.add_paragraph("5. 访问 http://127.0.0.1:5001 登录使用。")
    doc.add_paragraph("基本命令格式为 python -m tang_cli_ops.cli_main <command> [options]。"
                       "可集成到 cron 定时任务或 CI/CD 流水线中实现自动化运维。")

    doc.add_heading("6.2 安全机制", level=2)
    doc.add_paragraph(
        "路径安全：validate_file_path() 实现严格路径校验，禁止目录遍历（..序列）和命令注入（|;&$`等特殊字符）。"
    )
    doc.add_paragraph(
        "输入格式校验：validate_file_extension() 基于白名单机制仅允许合法扩展名，validate_export_format() 确保"
        "导出格式在五种合法类型内，validate_command_name() 限制命令在已注册的16个有效命令集合中。"
    )
    doc.add_paragraph(
        "Web认证安全：登录密码经 SHA-256 哈希后存储，Flask Session 管理用户会话，密钥随机生成（os.urandom(24).hex()）。"
        "API接口（除/system/health外）均需登录认证，未认证请求返回401错误。"
    )

    doc.add_heading("6.3 关键配置参数", level=2)
    add_table(doc, ["参数", "默认值", "环境变量", "说明"], [
        ["CACHE_DEFAULT_TTL", "600s", "TCO_CACHE_TTL", "缓存默认过期时间"],
        ["CACHE_CLEANUP_STRATEGY", "lru", "TCO_CACHE_STRATEGY", "淘汰策略"],
        ["MONITOR_COLLECTION_INTERVAL", "60s", "TCO_MONITOR_INTERVAL", "后台监控采集间隔"],
        ["MONITOR_DISK_THRESHOLD", "90%", "TCO_MONITOR_DISK_THRESHOLD", "磁盘使用率告警阈值"],
        ["EXPORT_MAX_ROWS", "50000", "TCO_EXPORT_MAX_ROWS", "单次导出最大行数"],
        ["BATCH_MAX_WORKERS", "4", "TCO_BATCH_WORKERS", "批处理最大并发线程数"],
        ["HEALTH_CHECK_RETRIES", "3", "TCO_HEALTH_RETRIES", "健康检查失败重试次数"],
        ["LOG_MAX_BYTES", "10MB", "TCO_LOG_MAX_BYTES", "单个日志文件最大大小"],
    ])
    doc.add_page_break()

    # ====================================================================
    # 七、测试覆盖
    # ====================================================================
    doc.add_heading("七、测试覆盖", level=1)
    doc.add_paragraph(
        "系统的 test 命令提供两层自动化验证：(1) 模块导入测试，逐一验证每个模块可正确导入且无循环依赖和 ImportError；"
        "(2) 核心函数功能测试（truncate 字符串截断、is_chinese_char 中文字符判断、frequency_count 频次统计、"
        "format_file_size 文件大小格式化），使用 assert 断言验证预期输入输出。"
    )
    doc.add_paragraph(
        "Web前端可通过浏览器手动验证：访问登录页 → 注册新用户 → 登录 → 查看控制台健康状态 → "
        "浏览主页功能卡片 → 测试API接口返回JSON数据。所有API接口均返回标准JSON格式，便于客户端和AGENT解析。"
    )
    doc.add_page_break()

    # ====================================================================
    # 八、文件清单
    # ====================================================================
    doc.add_heading("八、文件清单", level=1)
    add_table(doc, ["文件名", "行数", "说明"], [
        ["cli_main.py", "725", "16子命令入口、argparse路由"],
        ["config.py", "286", "配置管理、ConfigManager类"],
        ["errors.py", "407", "7层40+异常类、错误码体系"],
        ["logger.py", "209", "分级日志、文件轮转"],
        ["utils.py", "500", "30+通用工具函数"],
        ["validators.py", "224", "14个输入校验函数"],
        ["models.py", "450", "8个dataclass数据模型"],
        ["cache_manager.py", "391", "双层缓存+装饰器"],
        ["monitor.py", "365", "系统资源监控+健康评分"],
        ["export_engine.py", "519", "5格式导出引擎"],
        ["preprocessor.py", "375", "JSON清洗/校验/备份"],
        ["data_scanner.py", "325", "目录扫描/分类/变更检测"],
        ["health_checker.py", "341", "8项健康检查"],
        ["report_generator.py", "306", "Text/JSON/HTML报告"],
        ["batch_processor.py", "316", "并发批处理引擎"],
        ["__init__.py", "47", "包初始化、版本声明"],
        ["web/app.py", "286", "Flask Web应用、认证、API"],
        ["web/__init__.py", "2", "Web模块声明"],
        ["TOTAL", "~6,074", "18 个 Python 模块（含Web前端）"],
    ])
    doc.add_paragraph("")

    doc.save(OUT_MANUAL)
    print(f"Manual saved: {OUT_MANUAL} ({os.path.getsize(OUT_MANUAL) / 1024:.0f} KB)")


def build_source_docx():
    """生成源码 DOCX: 前2000行 + 后2000行"""
    all_lines = []
    for fname in MODULES:
        fpath = os.path.join(BASE, fname)
        with open(fpath, "r", encoding="utf-8") as f:
            lines = f.readlines()
        all_lines.append(f"# {'='*65}\n")
        all_lines.append(f"# 文件: tang_cli_ops/{fname}   行数: {len(lines)}\n")
        all_lines.append(f"# {'='*65}\n")
        all_lines.append("\n")
        all_lines.extend(lines)
        all_lines.append("\n")

    total = len(all_lines)
    front = all_lines[:2000]
    back = all_lines[-2000:]
    start_line = total - 2000 + 1

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(8)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.1

    # Cover
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("唐诗意象数据运维管理系统"); r.font.size = Pt(26); r.bold = True
    r.font.name = "黑体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("计算机软件著作权登记 — 源程序"); r.font.size = Pt(16)
    r.font.name = "黑体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    doc.add_paragraph(); doc.add_paragraph()
    for text in ["提交源码行数：前2000行 + 后2000行",
                 f"总源码行数：约 {total} 行（18 个 Python 模块）",
                 "开发语言：Python 3.10+",
                 "开发完成日期：2026年5月",
                 "核心设计：AGENT可直接操控CLI — 远程hook实现AI提效"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.size = Pt(12)
    doc.add_page_break()

    # Front 2000
    doc.add_heading("第一部分：源代码前2000行", level=1)
    for i, line in enumerate(front, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r_num = p.add_run(f"{i:5d} ")
        r_num.font.name = "Courier New"; r_num.font.size = Pt(7)
        r_num.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        display = line.rstrip("\n").replace("\t", "    ")
        r_code = p.add_run(display if display else " ")
        r_code.font.name = "Courier New"; r_code.font.size = Pt(8)

    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（中间部分略）"); r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.name = "宋体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.add_page_break()

    # Back 2000
    doc.add_heading("第二部分：源代码后2000行", level=1)
    for i, line in enumerate(back, start_line):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        r_num = p.add_run(f"{i:5d} ")
        r_num.font.name = "Courier New"; r_num.font.size = Pt(7)
        r_num.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        display = line.rstrip("\n").replace("\t", "    ")
        r_code = p.add_run(display if display else " ")
        r_code.font.name = "Courier New"; r_code.font.size = Pt(8)

    doc.save(OUT_SOURCE)
    print(f"Source saved: {OUT_SOURCE} ({os.path.getsize(OUT_SOURCE) / 1024:.0f} KB)")
    print(f"Total source lines: {total}")
    print(f"Front: {len(front)}, Back: {len(back)} (starting at line {start_line})")


if __name__ == "__main__":
    build_manual()
    build_source_docx()
