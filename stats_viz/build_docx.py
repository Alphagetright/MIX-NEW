# -*- coding: utf-8 -*-
"""生成 诗歌意象多维统计与可视化系统 软著说明书 + 源码 DOCX"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
SCR = os.path.join(BASE, "screenshots")
OUT_MANUAL = os.path.join(os.path.dirname(BASE), "软著_统计可视化系统_软件说明书.docx")
OUT_SOURCE = os.path.join(os.path.dirname(BASE), "软著_统计可视化系统_源代码.docx")

IMG_MAP = {
    "1_dashboard": "01_dashboard.png",
    "2_charts": "02_graph_charts.png",
    "3_table": "03_graph_table.png",
    "4_detail": "04_detail_modal.png",
    "5_filter": "05_filtered_table.png",
    "6_admin": "06_admin_stats.png",
    "7_analyze": "07_analyze_drawer.png",
    "8_ai": "08_ai_query.png",
    "9_recycle": "09_recycle.png",
    "10_overview": "10_stats_overview.png",
}

MODULES = [
    "__init__.py", "config.py", "errors.py", "logger.py", "utils.py", "validators.py",
    "models.py", "preprocessor.py", "data_loader.py", "stats_engine.py",
    "chart_data_builder.py", "data_explorer.py", "export_service.py",
    "report_builder.py", "correlation_analyzer.py",
    "visualization_service.py", "main.py",
]


def setup_styles(doc):
    s = doc.styles["Normal"]
    s.font.name = "宋体"; s.font.size = Pt(14)
    s.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    s.paragraph_format.line_spacing = 1.5
    for i in range(1, 4):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "黑体"; h.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        h.font.color.rgb = RGBColor(0, 0, 0)


def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(11)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = tbl.rows[i + 1].cells[j]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(11)
    doc.add_paragraph()


def add_img(doc, num, title, description=""):
    p = doc.add_paragraph()
    r = p.add_run(f"【图{num}】 {title}"); r.bold = True
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    key = f"{num}_{title.split('_')[0]}"
    fname = IMG_MAP.get(key)
    if not fname:
        for k, v in IMG_MAP.items():
            if k.startswith(str(num)): fname = v; break
    if fname:
        fp = os.path.join(SCR, fname)
        if os.path.exists(fp):
            doc.add_picture(fp, width=Inches(5.8))


def build_manual():
    doc = Document(); setup_styles(doc)

    # ── Cover ──
    for _ in range(8): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("诗歌意象多维统计与可视化系统"); r.font.size = Pt(30); r.bold = True
    r.font.name = "黑体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("软件说明书"); r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
    doc.add_paragraph()
    for t in ["版本：V1.0", "开发完成日期：2026年5月", "类型：Web 应用程序（B/S架构）",
              "开发语言：Python 3.10+ / HTML / JavaScript", "总代码行数：约 4,000 行（17 个模块）"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t); r.font.size = Pt(14)
    doc.add_page_break()

    # ── TOC ──
    doc.add_heading("目录", level=1)
    for item in ["一、引言", "二、系统架构", "三、核心功能与操作截图",
                 "四、数据处理流程", "五、部署与安全", "六、测试覆盖", "七、文件清单"]:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ── 一、引言 ──
    doc.add_heading("一、引言", level=1)
    doc.add_heading("1.1 项目背景", level=2)
    doc.add_paragraph(
        "唐诗是中国古典文学的瑰宝，其中蕴含丰富的意象资源。传统文学研究依赖人工标注和分析，效率较低"
        "且难以量化。本系统旨在通过计算机技术对唐诗意象进行多维度统计分析和可视化展示，为文学研究者"
        "和数字人文学者提供直观的数据分析工具。系统基于 B/S 架构，用户通过浏览器即可访问全部功能。"
    )
    doc.add_heading("1.2 软件功能概述", level=2)
    doc.add_paragraph("本系统提供以下核心统计分析功能：")
    for f in [
        "意象频次统计：自动统计所有意象的出现频次，支持 Top-N 排名和交互式柱状图展示。",
        "多维分类统计：按分类域（11个细分类别）、大类（自然/社会/人文）、情感类别、感知通道、体裁等 15+ 维度进行聚合统计。",
        "情感分析：统计情感类别分布、情感极性（正/负/中性）比例、情感-分类域交叉分析。",
        "诗人画像：统计每位诗人的意象使用量、高频意象、独有意象以及诗人之间的意象相似度。",
        "数据溯源查询：提供多维筛选（文本搜索+分类+体裁+情感）、分页浏览和详细信息查看功能。",
        "交互式图表：基于 ECharts 5.5.0 的柱状图、饼图、散点图、热力图、雷达图等多种图表类型。",
        "多格式数据导出：支持 CSV (UTF-8-BOM)、JSON、HTML 报告三种格式的统计结果导出。",
        "统计报告生成：自动生成包含全维度统计数据的文本/JSON/HTML报告。",
    ]:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("1.3 运行环境", level=2)
    add_table(doc, ["项目", "要求", "备注"], [
        ["操作系统", "Windows / Linux / macOS", "跨平台"],
        ["Python 版本", "Python 3.10+", "核心统计引擎"],
        ["Web 框架", "Flask 3.0+", "B/S 架构服务端"],
        ["浏览器", "Chrome/Firefox/Edge", "现代浏览器"],
        ["JavaScript 库", "ECharts 5.5.0, marked.js", "图表渲染与 Markdown"],
    ])
    doc.add_page_break()

    # ── 二、系统架构 ──
    doc.add_heading("二、系统架构", level=1)
    doc.add_heading("2.1 总体架构", level=2)
    doc.add_paragraph("系统采用 B/S 三层架构：前端展示层（HTML/CSS/JS + ECharts 图表 + SSE 流式渲染）、"
                       "API 路由层（Flask 路由，提供 REST API 和 SSE 端点）、"
                       "业务逻辑层（统计引擎 StatsEngine + 图表构建器 ChartDataBuilder + 数据探索器 DataExplorer + 关联分析器 CorrelationAnalyzer + 导出服务 + 报告生成）。"
                       "基础设施层包含配置管理、异常处理、日志系统、输入校验、工具函数和数据预处理等支撑模块。")

    doc.add_heading("2.2 模块清单", level=2)
    add_table(doc, ["序号", "模块", "文件名", "行数", "职责"], [
        ["1", "统计引擎", "stats_engine.py", "363", "20+维度聚合统计"],
        ["2", "主入口", "main.py", "401", "14个子命令CLI入口"],
        ["3", "图表构建", "chart_data_builder.py", "326", "9种ECharts图表配置"],
        ["4", "关联分析", "correlation_analyzer.py", "318", "共现/诗人/情感/网络分析"],
        ["5", "工具函数", "utils.py", "288", "30+通用函数+TF-IDF"],
        ["6", "报告生成", "report_builder.py", "284", "Text/JSON/HTML/MD/CSV报告"],
        ["7", "数据加载", "data_loader.py", "266", "JSON解析+数据集构建"],
        ["8", "数据模型", "models.py", "251", "15个dataclass模型"],
        ["9", "数据探索", "data_explorer.py", "250", "多维筛选/搜索/分页"],
        ["10", "导出服务", "export_service.py", "246", "CSV/JSON/HTML导出"],
        ["11", "可视化服务", "visualization_service.py", "244", "仪表盘/主题/交互"],
        ["12", "数据预处理", "preprocessor.py", "241", "JSON清洗/校验/备份"],
        ["13", "配置管理", "config.py", "146", "全局配置+分类体系"],
        ["14", "校验模块", "validators.py", "96", "10个输入校验函数"],
        ["15", "异常体系", "errors.py", "88", "分层异常类"],
        ["16", "日志模块", "logger.py", "75", "分级日志+文件轮转"],
        ["17", "包初始化", "__init__.py", "31", "版本+导出列表"],
    ])
    doc.add_page_break()

    # ── 三、核心功能与操作截图 ──
    doc.add_heading("三、核心功能与操作截图", level=1)
    doc.add_paragraph("本章节图文结合展示系统的核心统计分析和可视化功能。")

    sections = [
        ("3.1 系统首页仪表盘", "1_dashboard",
         "登录后进入首页仪表盘。顶部显示面包屑导航和当前用户信息，左侧为功能导航侧边栏。"
         "主区域展示 4 个统计指标卡片（净解析诗歌数、提取意象条目数、分类维度数、已知诗人数）"
         "以及系统概览和快速入口。", "dashboard"),
        ("3.2 意象统计图表", "2_charts",
         "数据图谱-意象统计图表页面。顶部显示整体统计摘要（净解析诗歌313首、提取真意象4,896条）。"
         "左侧为'核心意象 Top50'柱状图（蓝色，支持 dataZoom 滑块缩放），右侧为'意象分类域分布'"
         "柱状图（紫色）。图表基于 ECharts 5.5.0 实现，支持悬浮提示、点击联动筛选、全屏放大。", "charts"),
        ("3.3 溯源数据查询", "3_table",
         "数据图谱-溯源数据查询页面。提供四个筛选条件：文本搜索框（按意象文本关键词搜索）、"
         "分类域下拉筛选、体裁/标签下拉筛选、情感倾向下拉筛选。筛选栏下方为溯源数据表（每页25条），"
         "7列：诗歌编号、意象文本（蓝色加粗）、所属诗歌、体裁/标签、大类归属、情感倾向、操作按钮。"
         "操作按钮包括：详情（弹出模态框）、解析（右侧解析抽屉）、删除（移入回收站）。", "table"),
        ("3.4 意象详情模态框", "4_detail",
         "在溯源数据表中点击'详情'按钮弹出模态框。展示选中意象的完整结构化标注信息，包含25+字段："
         "意象文本（红色加粗）、大类归属、词性、成分类型、感知通道、素材类型、内部结构、指涉来源、"
         "表现功能、结构功能组、文化流通性、跨文化性、认知强度、核心意象、情感极性/类别/置信度、"
         "编码信息（大类编码、子类编码）以及四层摘要和完整诗句原文。", "detail"),
        ("3.5 多维筛选查询", "5_filter",
         "溯源数据表支持多维筛选。用户可通过文本搜索框输入关键词（如'月'），从分类域下拉菜单中选择"
         "特定类别（如'天文意象'），通过体裁下拉菜单筛选特定诗歌体裁，或通过情感倾向下拉菜单筛选。"
         "多个筛选条件可组合使用。筛选结果实时更新表格内容和分页信息。", "filter"),
        ("3.6 管理后台统计面板", "6_admin",
         "管理后台提供系统状态总览和统计分析功能。包含 8 个统计指标卡片（净解析诗歌、意象条目、"
         "分类维度、已知诗人、向量库状态、缓存条目、导出文件数、数据文件数）。"
         "展示两张 ECharts 图表：意象分类域分布图和情感类别分布图。"
         "提供数据管理操作区（刷新缓存、导出CSV/JSON/报告、清理缓存）和系统维护区。", "admin"),
        ("3.7 认知诗学解析抽屉", "7_analyze",
         "在溯源数据表的任意行点击'解析'按钮，从页面右侧滑入解析抽屉（默认宽度420px，可拖拽调整）。"
         "抽屉包含意象元信息区（文本/诗歌/诗句）、标签 chips 区（感知通道/分类域/表现功能等标注字段）"
         "和对话区（AI 从感知层、文化层、情感层、结构层、跨诗比较五个维度进行流式解析回答）。"
         "底部输入框支持多轮追问。", "analyze"),
        ("3.8 AI 智能问答", "8_ai",
         "AI 智能问答页面基于 RAG（检索增强生成）架构。用户输入问题后，系统通过向量检索从诗歌库中"
         "召回 Top-5 相关诗歌（左栏展示诗歌卡片：标题、作者、相似度评分、原文预览），同时由大模型"
         "基于检索结果生成流式回答（右栏，SSE 实时渲染）。底部提供追问输入框，支持多轮对话。", "ai"),
        ("3.9 回收站管理", "9_recycle",
         "回收站页面管理被删除的意象条目。表格列：编号、被删意象文本（红色加粗+删除线样式）、"
         "所属诗歌、大类归属、恢复按钮。支持单条恢复操作，恢复后条目回到溯源数据表。"
         "空状态时显示'回收站为空'提示。", "recycle"),
        ("3.10 统计概览仪表盘", "10_overview",
         "首页仪表盘提供系统的全维度统计概览。包括 4 个彩色统计卡片（蓝色-诗歌数、绿色-意象数、"
         "橙色-分类维度、紫色-诗人数量），系统概览区（两栏：核心能力+技术架构），"
         "以及 3 个快速入口卡片（数据图谱、AI问答、管理后台），方便用户快速导航。", "overview"),
    ]
    for title, num, desc, img_key in sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        add_img(doc, num.split("_")[0], img_key, desc)
    doc.add_page_break()

    # ── 四、数据处理流程 ──
    doc.add_heading("四、数据处理流程", level=1)
    doc.add_paragraph(
        "数据加载流程：系统启动时自动扫描 poem_json/ 目录下的 JSON 数据文件 → clean_json_content "
        "清洗（去BOM、去Markdown包裹、修复尾部逗号）→ json.loads 解析 → extract_poems 递归提取诗歌节点"
        " → 遍历分析单元筛选 is_imagery='1' 的条目 → 构建扁平化溯源数据集 → 全局缓存（避免重复解析）。"
        "数据集按诗歌标题+首句指纹去重，自动跳过 all_data.json 等非数据文件。"
    )
    doc.add_paragraph(
        "统计计算流程：StatsEngine 加载溯源数据集 → 按维度分组聚合（Counter/groupby）→ "
        "频次排序（frequency_count）→ 百分比计算（percentage_distribution）→ "
        "交叉分析（cross_analysis_emotion_category）→ 生成全维度摘要报告（summary_report/"
        "full_dimension_report）。所有统计函数返回 Python 原生数据结构（dict/list），"
        "前端通过 /api/data 和 /api/stats 端点获取 JSON 数据。"
    )
    doc.add_paragraph(
        "图表渲染流程：ChartDataBuilder 接收 StatsEngine 的统计输出 → 按照 ECharts option 规范"
        "构建图表配置对象（title/tooltip/xAxis/yAxis/series）→ 前端 JavaScript 通过 "
        "echarts.init().setOption() 渲染为交互式图表。支持的图表类型：柱状图（bar）、饼图（pie）、"
        "散点图（scatter）、热力图（heatmap）、树图（treemap）、桑基图（sankey）、雷达图（radar）。"
    )
    doc.add_page_break()

    # ── 五、部署与安全 ──
    doc.add_heading("五、部署与安全", level=1)
    doc.add_heading("5.1 安装与运行", level=2)
    doc.add_paragraph("1. 确保 Python 3.10+ 已安装。")
    doc.add_paragraph("2. pip install flask>=3.0.0（Web 服务依赖）。")
    doc.add_paragraph("3. 将 tang_stats_viz/ 和 poem_json/ 目录放置到目标服务器。")
    doc.add_paragraph("4. 启动 Web 服务：python app.py，默认监听 http://0.0.0.0:5000")
    doc.add_paragraph("5. 命令行统计：python -m tang_stats_viz.main stats")
    doc.add_heading("5.2 安全机制", level=2)
    doc.add_paragraph(
        "用户认证：基于 Session 的登录认证，密码 SHA256 哈希存储，未登录自动重定向到登录页。"
        "输入校验：所有用户输入经过 validate_* 函数校验（非空/长度/格式/路径安全）。"
        "HTML 转义：escape_html() 防止 XSS 攻击。速率限制：API 端点 60次/分钟/IP。"
    )
    doc.add_page_break()

    # ── 六、测试覆盖 ──
    doc.add_heading("六、测试覆盖", level=1)
    doc.add_paragraph(
        "系统提供 test 命令运行自动化测试：17个模块导入测试（验证所有模块可正确导入且无循环依赖），"
        "核心函数功能测试（truncate 字符串截断、is_chinese_char 中文字符判断、frequency_count "
        "频次统计、format_file_size 文件大小格式化），数据加载测试（验证数据集构建成功、诗歌数和"
        "意象数均大于零）。命令行参数校验通过 argparse choices 机制自动拦截无效命令。"
    )
    doc.add_page_break()

    # ── 七、文件清单 ──
    doc.add_heading("七、文件清单", level=1)
    row_data = []
    total = 0
    for fname in MODULES:
        fp = os.path.join(BASE, fname)
        with open(fp, "r", encoding="utf-8") as f:
            n = len(f.readlines())
        row_data.append([fname, str(n), "Python"])
        total += n
    row_data.append(["TOTAL", str(total), "17个Python模块"])
    add_table(doc, ["文件名", "行数", "类型"], row_data)

    doc.save(OUT_MANUAL)
    print(f"Manual: {OUT_MANUAL} ({os.path.getsize(OUT_MANUAL)/1024:.0f} KB)")
    return total


def build_source_docx(total_lines):
    all_lines = []
    for fname in MODULES:
        fp = os.path.join(BASE, fname)
        with open(fp, "r", encoding="utf-8") as f:
            lines = f.readlines()
        all_lines.append(f"# {'='*65}\n")
        all_lines.append(f"# 文件: tang_stats_viz/{fname}   行数: {len(lines)}\n")
        all_lines.append(f"# {'='*65}\n\n")
        all_lines.extend(lines)
        all_lines.append("\n")

    total = len(all_lines)
    front = all_lines[:2000]
    back = all_lines[-2000:]
    start_line = total - 2000 + 1

    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21); s.page_height = Cm(29.7)
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Courier New"; style.font.size = Pt(8)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.1

    # Cover
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("诗歌意象多维统计与可视化系统"); r.font.size = Pt(26); r.bold = True
    r.font.name = "黑体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("计算机软件著作权登记 — 源程序"); r.font.size = Pt(16)
    r.font.name = "黑体"; r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    doc.add_paragraph(); doc.add_paragraph()
    for t in ["提交源码行数：前2000行 + 后2000行（共4000行）",
              f"总源码行数：约 {total} 行（17 个 Python 模块）",
              "开发语言：Python 3.10+", "开发完成日期：2026年5月"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t); r.font.size = Pt(12)
    doc.add_page_break()

    doc.add_heading("第一部分：源代码前2000行", level=1)
    for i, line in enumerate(front, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        rn = p.add_run(f"{i:5d} "); rn.font.name = "Courier New"
        rn.font.size = Pt(7); rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        rc = p.add_run(line.rstrip("\n").replace("\t", "    ") or " ")
        rc.font.name = "Courier New"; rc.font.size = Pt(8)

    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（中间部分略）"); r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_page_break()

    doc.add_heading("第二部分：源代码后2000行", level=1)
    for i, line in enumerate(back, start_line):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        rn = p.add_run(f"{i:5d} "); rn.font.name = "Courier New"
        rn.font.size = Pt(7); rn.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        rc = p.add_run(line.rstrip("\n").replace("\t", "    ") or " ")
        rc.font.name = "Courier New"; rc.font.size = Pt(8)

    doc.save(OUT_SOURCE)
    print(f"Source: {OUT_SOURCE} ({os.path.getsize(OUT_SOURCE)/1024:.0f} KB)")


if __name__ == "__main__":
    total = build_manual()
    build_source_docx(total)
    print(f"\nDone! Total lines: {total}")
