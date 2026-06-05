# -*- coding: utf-8 -*-
"""生成软著说明书 .docx 文件，含截图占位"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"C:\Users\Administrator\Desktop\All Mix"
SCR = os.path.join(BASE, "screenshots")
OUT = os.path.join(BASE, "软著_软件说明书.docx")

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '黑体'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h.font.color.rgb = RGBColor(0, 0, 0)

def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    # data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    doc.add_paragraph()
    return table

def add_img_placeholder(doc, num, title, description):
    """添加截图占位"""
    p = doc.add_paragraph()
    run = p.add_run(f'【图{num}】 {title}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # Try to insert actual image
    img_path = os.path.join(SCR, f"{num}_{title}.png")
    # Map to actual filenames
    img_map = {
        "1_login": "1_login_page.png",
        "2_register": "2_register_page.png",
        "3_dashboard": "3_dashboard.png",
        "4_charts": "4_graph_charts.png",
        "5_submenu": "5_submenu_expanded.png",
        "6_table": "6_graph_table.png",
        "7_recycle": "7_recycle.png",
        "8_ai": "8_ai.png",
        "9_admin": "9_admin.png",
        "10_detail_modal": "10_detail_modal.png",
        "11_analyze_drawer": "11_analyze_drawer.png",
    }
    actual = img_map.get(f"{num}_{title.split('_')[0] if '_' in title else ''}", None)
    if not actual:
        # generic lookup
        for k, v in img_map.items():
            if k.startswith(num):
                actual = v
                break

    if actual:
        full = os.path.join(SCR, actual)
        if os.path.exists(full):
            try:
                doc.add_picture(full, width=Inches(5.5))
            except:
                _add_text_placeholder(doc, description)
        else:
            _add_text_placeholder(doc, description)
    else:
        _add_text_placeholder(doc, description)

def _add_text_placeholder(doc, description):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'[ 请插入全屏截图 ]')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    r2 = p2.add_run(description)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p3 = doc.add_paragraph()
    r3 = p3.add_run('注意：全屏截图、无水印无LOGO、截图清晰完整、展示完整界面')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    r3.bold = True

def add_code(doc, text):
    """添加代码块"""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line or ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(8)

# ================================================================
# COVER
# ================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('唐诗意象智能分析系统')
r.font.size = Pt(32)
r.bold = True
r.font.name = '黑体'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('软件说明书')
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
r.font.name = '黑体'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

for text in ['版本：V1.0', '开发完成日期：2026年5月', '类型：Web应用程序（B/S架构）', '总代码行数：约 6,300 行（Python 4,426 行 + HTML 1,887 行）']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(16)

doc.add_page_break()

# ================================================================
# CONTENT
# ================================================================

doc.add_heading('目录', level=1)
doc.add_paragraph('一、引言')
doc.add_paragraph('    1.1 项目背景')
doc.add_paragraph('    1.2 软件概述')
doc.add_paragraph('    1.3 运行环境')
doc.add_paragraph('二、系统架构')
doc.add_paragraph('三、数据结构与数据库设计')
doc.add_paragraph('四、功能模块详细说明（含界面截图）')
doc.add_paragraph('    4.0 用户认证模块（登录与注册）')
doc.add_paragraph('    4.1 首页仪表盘')
doc.add_paragraph('    4.2 数据图谱 — 意象统计图表')
doc.add_paragraph('    4.3 数据图谱 — 溯源数据查询')
doc.add_paragraph('    4.4 数据图谱 — 意象详情与回收站')
doc.add_paragraph('    4.5 认知诗学意象解析')
doc.add_paragraph('    4.6 AI 智能问答')
doc.add_paragraph('    4.7 管理后台')
doc.add_paragraph('    4.8 数据导出')
doc.add_paragraph('    4.9 CLI命令行工具')
doc.add_paragraph('    4.10 缓存/日志/安全等基础设施')
doc.add_paragraph('五、API 接口清单')
doc.add_paragraph('六、部署与运行')
doc.add_paragraph('七、安全与错误处理')
doc.add_paragraph('八、测试覆盖')
doc.add_paragraph('九、关键配置参数')
doc.add_paragraph('十、文件清单')
doc.add_page_break()

# ================================================================
# 一、引言
# ================================================================
doc.add_heading('一、引言', level=1)

doc.add_heading('1.1 项目背景', level=2)
doc.add_paragraph('唐诗是中国古典文学的瑰宝，其中蕴含丰富的意象资源。传统文学研究依赖人工标注和分析，效率较低且难以量化。"唐诗意象智能分析系统"旨在通过计算机技术，对唐诗中的意象进行结构化标注、多维统计分析和智能检索，为文学研究者提供数字化工具支持。')

doc.add_heading('1.2 软件概述', level=2)
doc.add_paragraph('本系统是一个基于Flask框架的B/S架构Web应用程序，集成了以下核心能力：')
doc.add_paragraph('- 意象数据管理：对313首唐诗中的12,738个分析单元进行结构化存储与管理，其中标注为"意象"的条目共4,896条', style='List Bullet')
doc.add_paragraph('- 多维度统计分析：支持意象频次、分类域、情感倾向、感知通道等15个维度的统计与可视化', style='List Bullet')
doc.add_paragraph('- RAG智能问答：基于向量检索（ChromaDB）+大语言模型，实现唐诗相关知识的检索增强生成', style='List Bullet')
doc.add_paragraph('- 认知诗学解析：利用AI对意象进行感知层、文化层、情感层、结构层、跨诗比较的多维解析', style='List Bullet')
doc.add_paragraph('- 用户认证：支持用户注册与登录，基于SHA256加密存储的账号系统', style='List Bullet')
doc.add_paragraph('- 数据导出：支持CSV、JSON、分析报告等多种格式的数据导出', style='List Bullet')

doc.add_heading('1.3 运行环境', level=2)
add_table(doc,
    ['项目', '规格'],
    [
        ['操作系统', 'Windows / Linux / macOS'],
        ['Python 版本', 'Python 3.10+'],
        ['Web 框架', 'Flask 3.0+'],
        ['向量数据库', 'ChromaDB 0.4+'],
        ['大模型 API', '兼容 OpenAI API 格式的本地/远程服务'],
        ['嵌入模型', 'text-embedding-qwen3-embedding-4b'],
        ['对话模型', 'qwen3.6-35b-a3b'],
        ['浏览器', 'Chrome / Firefox / Edge 等现代浏览器'],
    ]
)

doc.add_page_break()

# ================================================================
# 二、系统架构
# ================================================================
doc.add_heading('二、系统架构', level=1)

doc.add_heading('2.1 总体架构', level=2)
doc.add_paragraph('系统采用分层架构设计：')
doc.add_paragraph('- 前端展示层：HTML/CSS/JS + ECharts图表 + SSE流式渲染 + 右侧抽屉UI', style='List Bullet')
doc.add_paragraph('- API路由层：Flask路由，提供REST API和SSE端点', style='List Bullet')
doc.add_paragraph('- 业务逻辑层：统计服务、导出服务、RAG引擎、认知诗学解析', style='List Bullet')
doc.add_paragraph('- 数据访问层：数据模型、JSON数据文件、ChromaDB向量数据库、缓存系统', style='List Bullet')
doc.add_paragraph('- 基础设施层：日志系统、异常处理、中间件、系统监控、校验模块', style='List Bullet')

doc.add_heading('2.2 模块划分', level=2)
add_table(doc,
    ['模块', '文件名', '行数', '职责'],
    [
        ['Web主程序', 'app.py', '497', 'Flask路由、API端点、数据加载'],
        ['用户管理', 'users.py', '91', '用户注册、登录认证、SHA256存储'],
        ['配置管理', 'config.py', '118', '全局配置常量'],
        ['数据模型', 'models.py', '192', '数据类定义'],
        ['统计分析', 'analytics.py', '312', '多维度聚合统计'],
        ['RAG引擎', 'query_rag.py', '237', '向量检索与流式问答'],
        ['RAG构建', 'build_rag.py', '121', '向量数据库构建'],
        ['管理后台', 'admin.py', '341', '系统管理Blueprint'],
        ['CLI工具', 'cli.py', '239', '命令行运维工具'],
        ['工具函数', 'utils.py', '145', '通用工具函数'],
        ['缓存模块', 'cache.py', '143', '内存+文件缓存'],
        ['异常处理', 'errors.py', '111', '统一异常类'],
        ['校验模块', 'validators.py', '83', '请求参数校验'],
        ['日志模块', 'logger.py', '92', '分级日志记录'],
        ['导出服务', 'export_service.py', '117', 'CSV/JSON导出'],
        ['中间件', 'middleware.py', '187', '请求追踪/CORS/限流'],
        ['系统监控', 'monitor.py', '185', '系统资源监控'],
        ['数据预处理', 'preprocessor.py', '148', 'JSON清洗与校验'],
        ['图表页', 'templates/graph_charts.html', '121', '意象统计图表'],
        ['查询页', 'templates/graph_table.html', '568', '溯源数据查询'],
        ['问答页', 'templates/ai.html', '175', 'AI智能问答界面'],
        ['仪表盘', 'templates/dashboard.html', '88', '首页仪表盘'],
        ['登录页', 'templates/login.html', '143', '用户登录界面'],
        ['注册页', 'templates/register.html', '111', '用户注册界面'],
        ['布局', 'templates/base.html', '349', '侧边栏+顶栏基础布局'],
        ['管理页面', 'templates/admin.html', '262', '管理后台UI'],
        ['回收站', 'templates/recycle.html', '70', '回收站页面'],
        ['测试套件', 'tests/*.py', '1067', '单元测试'],
    ]
)

doc.add_page_break()

# ================================================================
# 三、数据结构
# ================================================================
doc.add_heading('三、数据结构与数据库设计', level=1)

doc.add_heading('3.1 数据源', level=2)
doc.add_paragraph('系统数据来源于 poem_json/tang_poetry_313.json 文件（10.9 MB），包含313首唐诗的结构化标注数据。每条诗歌记录包含：诗歌编号、标题、作者、朝代、分类标签、体裁、原文、诗行数组、分析单元数组（25个标注字段）、情感轨迹、意象关系。')

doc.add_heading('3.2 数据统计', level=2)
add_table(doc,
    ['指标', '数值'],
    [
        ['诗歌总数', '313 首'],
        ['分析单元总数', '12,738 条'],
        ['标注为"意象"的条目', '4,896 条'],
        ['去重意象文本数', '2,840 个'],
        ['诗人数量', '77 位'],
        ['分类标签/体裁数', '28 种'],
        ['意象大类编码', '3 类（自然意象、社会意象、人文意象）'],
        ['意象子类编码', '11 类'],
    ]
)

doc.add_heading('3.3 意象分类体系', level=2)
add_table(doc,
    ['大类编码', '名称', '子类'],
    [
        ['1', '自然意象', '1-1天文、1-2地理、1-3植物、1-4动物'],
        ['2', '社会意象', '2-1生产生活、2-2军事战争、2-3制度观念'],
        ['3', '人文意象', '3-1人造器物、3-2人类自身、3-3人物角色、3-4文化'],
    ]
)

doc.add_heading('3.4 向量数据库设计', level=2)
doc.add_paragraph('基于ChromaDB的持久化向量数据库，存储在rag_db/目录。集合名称为poems，距离函数为cosine（余弦相似度），向量模型为text-embedding-qwen3-embedding-4b。文档内容为诗歌原文全文，元数据包含标题、作者、诗歌编号。检索支持按作者过滤+语义向量检索，Top-K默认值为5。')

doc.add_page_break()

# ================================================================
# 四、功能模块详细说明（含界面截图）
# ================================================================
doc.add_heading('四、功能模块详细说明（含界面截图）', level=1)

doc.add_paragraph('本章节图文结合演示系统的完整操作流程。所有截图应为全屏截图、无水印无LOGO、清晰完整。系统采用侧边栏导航布局，一级菜单包含：首页、数据图谱（含三个二级菜单）、AI智能问答、管理后台。')

# ── 4.0 用户认证 ──
doc.add_heading('4.0 用户认证模块（登录与注册）', level=2)

doc.add_paragraph('系统提供完整的用户注册与登录功能。用户账号基于SHA256哈希加密存储于JSON文件中，支持多用户独立账号。访问系统根路径 "/" 时自动重定向到登录页，未登录用户无法访问任何功能页面。')

doc.add_heading('4.0.1 登录界面', level=3)
doc.add_paragraph('登录页面提供用户名和密码输入框。用户输入正确的账号密码后，系统验证通过并创建会话（Session），自动跳转至首页仪表盘。验证失败时显示错误提示。登录页底部提供"立即注册"链接，引导新用户创建账号。默认管理员账号：admin / admin123。')
add_img_placeholder(doc, '1', 'login', '【登录页面全屏截图】深蓝渐变背景，中央白色登录卡片，包含"唐"字LOGO、系统名称、用户名输入框、密码输入框、"登录"按钮、底部"还没有账号？立即注册"链接。')

doc.add_heading('4.0.2 注册界面', level=3)
doc.add_paragraph('注册页面提供用户名和密码输入框。用户名要求2-20个字母数字字符，密码要求至少4个字符。注册成功后自动登录并跳转到首页仪表盘。若用户名已存在则提示错误。密码经过SHA256哈希加密存储。')
add_img_placeholder(doc, '2', 'register', '【注册页面全屏截图】绿色主题，中央白色注册卡片，包含"创建新账号"标题、用户名输入框（带格式提示）、密码输入框（带加密存储提示）、"注册"按钮、底部"已有账号？立即登录"链接。')

# ── 4.1 首页仪表盘 ──
doc.add_heading('4.1 首页仪表盘', level=2)
doc.add_paragraph('登录成功后进入首页仪表盘。页面顶部显示面包屑导航，左侧为深色侧边栏导航（220px宽），右侧主内容区包含：4个统计卡片（净解析诗歌313首、提取真意象4,896条、分类维度11、已知诗人25位）、系统概览区（核心能力+技术架构）、3个快速入口卡片（数据图谱、AI问答、管理后台）。侧边栏支持一级菜单展开/收起二级菜单。')
add_img_placeholder(doc, '3', 'dashboard', '【首页仪表盘全屏截图】左侧深色侧边栏（"首页"菜单高亮），右侧白色顶栏（面包屑+当前用户信息），4个有色统计卡片，两栏系统概览，3个快速入口卡片。')

# ── 4.2 数据图谱-图表 ──
doc.add_heading('4.2 数据图谱 — 意象统计图表', level=2)
doc.add_paragraph('点击侧边栏"数据图谱"一级菜单，展开三个二级菜单项。点击"意象统计图表"进入图表页面。页面展示核心意象Top50柱状图（支持dataZoom滑块缩放、点击柱条联动筛选、全屏放大）和意象分类域分布柱状图。底部提供"进入溯源数据查询"快捷跳转按钮。')
add_img_placeholder(doc, '4', 'charts', '【意象统计图表全屏截图】顶部统计条（诗歌数+意象数），左右两张ECharts柱状图（Top50蓝柱+分类域紫柱），底部"进入溯源数据查询→"按钮。')
add_img_placeholder(doc, '5', 'submenu', '【侧边栏二级菜单展开全屏截图】侧边栏中"数据图谱"菜单展开，显示三个二级菜单项（意象统计图表、溯源数据查询、回收站管理），展开箭头旋转90度。')

# ── 4.3 数据图谱-表格 ──
doc.add_heading('4.3 数据图谱 — 溯源数据查询', level=2)
doc.add_paragraph('点击侧边栏"溯源数据查询"进入数据查询页面。页面提供四个筛选条件：文本搜索框（按意象文本关键词实时搜索）、分类域下拉筛选、体裁/标签下拉筛选、情感倾向下拉筛选。筛选栏下方为溯源数据表，每页25条，最多渲染5,000条。表格包含7列：诗歌编号、意象文本（蓝色加粗）、所属诗歌、体裁/标签、大类归属、情感倾向、操作按钮。操作按钮包括：详情（弹出模态框展示25+标注字段）、解析（打开右侧认知诗学解析抽屉）、删除（移入回收站）。表格支持固定表头滚动和分页控件。页面上方提供"返回统计图表"链接和回收站入口按钮。')
add_img_placeholder(doc, '6', 'table', '【溯源数据查询全屏截图】筛选栏（搜索框+3个下拉筛选器+清除筛选按钮），数据表（7列带分页控件），"返回统计图表"链接和"回收站"按钮。')

# ── 4.4 回收站 ──
doc.add_heading('4.4 数据图谱 — 意象详情与回收站', level=2)
doc.add_paragraph('意象详情模态框：在数据表中点击"详情"按钮弹出。模态框展示选中意象的完整结构化标注信息，包括意象文本（红色加粗）、大类归属、词性、成分类型、感知通道、素材类型、内部结构、指涉来源、表现功能、结构功能组、文化流通性、跨文化性、认知强度、核心意象、情感极性、情感类别、情感置信度、编码信息等25+字段，以及四层摘要和完整诗句原文。')
add_img_placeholder(doc, '10', 'detail_modal', '【意象详情模态框全屏截图】模态弹窗展示意象完整标注信息，含25+结构化字段、四层摘要、诗句原文，红色关闭按钮。')
doc.add_paragraph('回收站：在数据表中点击"删除"按钮将条目移入回收站。点击"回收站"按钮弹出模态框，显示被删意象列表（红色加粗+删除线样式），支持单条恢复操作。恢复后条目回到溯源表。')
add_img_placeholder(doc, '7', 'recycle', '【回收站全屏截图】回收站页面，表格列：编号、被删意象（红色删除线）、所属诗歌、大类归属、恢复按钮。空状态显示"回收站为空"。')

# ── 4.5 认知诗学解析 ──
doc.add_heading('4.5 认知诗学意象解析', level=2)
doc.add_paragraph('在溯源数据表的任意行点击"解析"按钮，从页面右侧滑入解析抽屉（默认宽度420px，可拖拽左侧边框调整）。抽屉包含以下区域：')
doc.add_paragraph('- 标题区域：显示意象文本（蓝色加粗）、所属诗歌名称及作者、楷体诗句原文', style='List Bullet')
doc.add_paragraph('- 标签区：灰蓝色圆角标签chips展示感知通道、分类域、表现功能、认知强度、情感类别/极性等标注字段', style='List Bullet')
doc.add_paragraph('- 对话区：首轮自动发送"请从认知诗学角度深度解析这个意象在诗歌中的多维功能"，AI从5个维度（感知层、文化层、情感层、结构层、跨诗比较）进行流式回答。用户消息蓝色气泡右对齐，AI消息灰色气泡左对齐，支持Markdown渲染', style='List Bullet')
doc.add_paragraph('- 输入区：底部追问输入框+蓝色发送按钮，支持多轮追问', style='List Bullet')
doc.add_paragraph('解析结果基于意象的结构化标注数据，确保分析有据可依。同一意象的解析结果自动缓存，重复打开不重复请求。支持AbortController中断请求。')
add_img_placeholder(doc, '11', 'analyze_drawer', '【认知诗学解析抽屉全屏截图】右侧抽屉滑入，展示意象元信息、标签chips、AI流式解析内容（5维度分析），底部追问输入框。')

# ── 4.6 AI问答 ──
doc.add_heading('4.6 AI 智能问答模块', level=2)
doc.add_paragraph('点击侧边栏"AI智能问答"进入RAG检索增强生成问答页面。初始状态显示hero区域：大号输入框（含placeholder提示）+蓝色"检索并生成回答"按钮。输入问题后点击按钮：')
doc.add_paragraph('1. 系统自动识别问题中涉及的诗人（基于25位已知诗人列表），提取语义查询关键词', style='List Number')
doc.add_paragraph('2. 调用嵌入模型API将问题转为向量，在ChromaDB中检索Top-5相似诗歌', style='List Number')
doc.add_paragraph('3. 检索结果在左栏（40%宽度）以诗歌卡片展示：标题、作者、相似度评分（绿色）、诗歌原文预览（楷体最多120字）', style='List Number')
doc.add_paragraph('4. 右栏（60%宽度）以SSE流式方式逐字输出大模型生成回答', style='List Number')
doc.add_paragraph('5. 底部提供追问输入框+发送按钮，支持多轮对话，保存完整对话历史', style='List Number')
add_img_placeholder(doc, '8', 'ai', '【AI智能问答全屏截图】左栏检索诗歌卡片列表（含相似度），右栏流式回答内容，底部追问输入框。')

# ── 4.7 管理后台 ──
doc.add_heading('4.7 管理后台', level=2)
doc.add_paragraph('点击侧边栏"管理后台"进入系统运维页面（/admin/）。页面包含以下功能区：')
doc.add_paragraph('- 系统状态面板：8个统计卡片（净解析诗歌、提取意象条目、分类维度、已知诗人、向量数据库状态、内存缓存条目、导出文件数、诗歌数据文件数）', style='List Bullet')
doc.add_paragraph('- 统计分析图表：意象分类域分布图（ECharts水平柱状图）+情感类别分布图（ECharts柱状图）', style='List Bullet')
doc.add_paragraph('- 系统配置表：展示当前运行配置（嵌入模型名、对话模型名、服务端口、检索数量、分页大小）', style='List Bullet')
doc.add_paragraph('- 数据管理操作区：刷新数据缓存、导出CSV、导出JSON、导出分析报告、清理全部缓存', style='List Bullet')
doc.add_paragraph('- 系统维护操作区：扫描数据目录、检查向量库、构建向量库、运行单元测试（后台执行，轮询任务状态）', style='List Bullet')
doc.add_paragraph('- 导出文件列表：展示历史导出文件（文件名、大小、修改时间），支持一键清空', style='List Bullet')
add_img_placeholder(doc, '9', 'admin', '【管理后台全屏截图】8个统计卡片网格，两张ECharts图表（分类域分布+情感分布），系统配置表，数据管理按钮区，系统维护按钮区，导出文件列表。')

# ── 4.8~4.17 其他模块（文本描述，省略截图） ──
modules_text = [
    ('4.8 数据导出模块', '文件：export_service.py（117行）\n\n支持三种格式的数据导出：CSV导出（UTF-8-BOM编码，24个字段）、JSON导出（indent格式化）、统计报告导出。辅助功能：list_exports()列出导出文件、clear_exports()清空导出目录、export_filtered_slice()导出数据子集。'),
    ('4.9 CLI命令行工具', '文件：cli.py（239行）\n\n提供8个子命令的系统运维工具：status（系统状态）、scan（扫描数据目录）、export（导出数据）、clear-cache（清理缓存）、list-exports（列出导出文件）、check-rag（检查向量库）、build-rag（构建向量库）、test（运行单元测试）。'),
    ('4.10 RAG向量库构建', '文件：build_rag.py（121行）\n\n独立的向量数据库构建脚本。清除旧库后扫描JSON数据文件，逐首读取诗歌原文，调用嵌入模型API生成文本向量，存入ChromaDB持久化集合（含元数据：标题、作者、诗歌编号）。每首间隔0.3秒避免API限流。'),
    ('4.11 缓存模块', '文件：cache.py（143行）\n\n实现两级缓存体系：MemoryCache（线程安全、字典存储、TTL自动过期）和FileCache（持久化到cache/目录、每个缓存项独立JSON文件）。提供cached装饰器支持自动缓存函数返回值。全局开关CACHE_ENABLED。'),
    ('4.12 日志/异常/校验/中间件/监控/工具', '日志模块（logger.py）：基于Python logging + RotatingFileHandler，10MB滚动备份。\n异常处理（errors.py）：分层异常体系（AppError→DataNotFoundError/EmbeddingError/LLMError等7个子类），Flask统一异常处理器。\n校验模块（validators.py）：validate_question/validate_item/validate_history/validate_poem_id/validate_filename/sanitize_html。\n中间件（middleware.py）：请求上下文（RequestContext）、CORS、rate_limit速率限制（60次/分钟/IP）、timer计时、PerformanceTracker性能追踪。\n系统监控（monitor.py）：SystemMonitor，支持磁盘/内存/CPU/Python运行时/项目文件信息快照，后台定时收集，健康检查。\n工具函数（utils.py）：20个通用工具函数（truncate、extract_numbers、is_chinese_char、chinese_ratio、split_sentences、format_file_size、format_timestamp、safe_get、dict_pick、dict_omit、md5_hash、safe_json_loads、list_chunk、deduplicate_by_key、merge_dicts、normalize_whitespace等）。'),
]

for title, content in modules_text:
    doc.add_heading(title, level=2)
    for para_text in content.split('\n\n'):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())

doc.add_page_break()

# ================================================================
# 五、API接口清单
# ================================================================
doc.add_heading('五、API 接口清单', level=1)

doc.add_heading('5.1 前端与认证 API', level=2)
add_table(doc,
    ['方法', '路径', '功能', '认证'],
    [
        ['GET', '/login', '登录页面', '否'],
        ['POST', '/login', '提交登录', '否'],
        ['GET', '/register', '注册页面', '否'],
        ['POST', '/register', '提交注册', '否'],
        ['GET', '/logout', '退出登录', '否'],
        ['GET', '/', '重定向→登录/首页', '是'],
        ['GET', '/dashboard', '首页仪表盘', '是'],
        ['GET', '/graph', '意象统计图表', '是'],
        ['GET', '/graph/table', '溯源数据查询', '是'],
        ['GET', '/ai', 'AI智能问答', '是'],
        ['GET', '/recycle', '回收站', '是'],
        ['GET', '/api/data', '获取全部数据', '否'],
        ['GET', '/api/stats', '统计分析概览', '否'],
        ['POST', '/api/export', '导出数据', '否'],
        ['POST', '/api/ask', 'RAG问答（SSE流）', '否'],
        ['POST', '/api/analyze', '认知诗学解析（SSE流）', '否'],
        ['GET', '/api/cache/status', '缓存状态', '否'],
    ]
)

doc.add_heading('5.2 管理后台 API', level=2)
add_table(doc,
    ['方法', '路径', '功能'],
    [
        ['GET', '/admin/', '管理后台首页'],
        ['GET', '/admin/api/status', '系统状态总览'],
        ['GET', '/admin/api/stats', '详细统计数据'],
        ['GET/POST', '/admin/api/exports', '列出/创建导出文件'],
        ['POST', '/admin/api/exports/clear', '清空导出目录'],
        ['POST', '/admin/api/cache/clear', '清理缓存'],
        ['POST', '/admin/api/refresh', '刷新数据缓存'],
        ['GET', '/admin/api/scan', '扫描数据目录'],
        ['GET', '/admin/api/rag/check', '检查向量库状态'],
        ['POST', '/admin/api/rag/build', '构建向量库（后台）'],
        ['POST', '/admin/api/test', '运行单元测试（后台）'],
        ['GET', '/admin/api/tasks/<id>', '查询后台任务状态'],
    ]
)

doc.add_page_break()

# ================================================================
# 六、部署与运行
# ================================================================
doc.add_heading('六、部署与运行', level=1)

doc.add_heading('6.1 环境依赖', level=2)
add_code(doc, 'flask>=3.0.0\nrequests>=2.28.0\nchromadb>=0.4.0\nopenpyxl>=3.1.0\npytest>=8.0.0')

doc.add_heading('6.2 启动方式', level=2)
doc.add_paragraph('Web服务启动：在项目根目录下执行 python app.py，默认监听 http://0.0.0.0:5000')
doc.add_paragraph('向量库构建（首次使用前）：python build_rag.py')
doc.add_paragraph('CLI管理工具：python cli.py [status|scan|export|clear-cache|list-exports|check-rag|build-rag|test]')

doc.add_heading('6.3 访问方式', level=2)
doc.add_paragraph('系统首页：http://localhost:5000/（自动跳转到登录页）')
doc.add_paragraph('管理后台：http://localhost:5000/admin/')
doc.add_paragraph('默认管理员账号：admin / admin123')

doc.add_page_break()

# ================================================================
# 七、安全与错误处理
# ================================================================
doc.add_heading('七、安全与错误处理', level=1)

doc.add_heading('7.1 认证安全', level=2)
doc.add_paragraph('- 用户密码使用SHA256哈希加密存储，不保存明文密码', style='List Bullet')
doc.add_paragraph('- 登录后创建Session会话，超时1小时自动失效', style='List Bullet')
doc.add_paragraph('- 所有功能页面（/dashboard /graph /ai /admin等）均需登录认证', style='List Bullet')
doc.add_paragraph('- 未登录访问自动重定向到登录页面', style='List Bullet')

doc.add_heading('7.2 输入校验', level=2)
doc.add_paragraph('- 注册：用户名2-20位字母数字，密码至少4位，用户名唯一性检查', style='List Bullet')
doc.add_paragraph('- 用户提问：非空、不超过2000字符', style='List Bullet')
doc.add_paragraph('- 对话历史：校验role/content结构，限制内容长度5000字符', style='List Bullet')
doc.add_paragraph('- HTML转义：所有用户可见数据escapeHtml处理，防XSS攻击', style='List Bullet')

doc.add_heading('7.3 速率限制与缓存安全', level=2)
doc.add_paragraph('- 进程内速率限制：60次/分钟/IP，超限返回HTTP 429', style='List Bullet')
doc.add_paragraph('- 双层缓存均采用线程锁（RLock）保护，TTL自动过期', style='List Bullet')

doc.add_page_break()

# ================================================================
# 八、测试覆盖
# ================================================================
doc.add_heading('八、测试覆盖', level=1)
add_table(doc,
    ['测试文件', '行数', '测试内容'],
    [
        ['test_utils.py', '172', '工具函数（截断、提取数字、中文字符判断等）'],
        ['test_analytics.py', '160', '统计分析服务各类聚合方法'],
        ['test_preprocessor.py', '123', 'JSON清洗、校验、批量验证'],
        ['test_validators.py', '118', '请求校验函数'],
        ['test_cache.py', '114', '内存缓存、文件缓存的读写与TTL'],
        ['test_models.py', '108', '数据模型创建、序列化、属性'],
        ['test_middleware.py', '90', '中间件、限流装饰器、性能追踪'],
        ['test_export.py', '88', '导出服务的CSV/JSON导出'],
        ['test_config.py', '62', '配置模块的路径和常量'],
        ['test_cli.py', '31', 'CLI命令解析'],
    ]
)

doc.add_page_break()

# ================================================================
# 九、关键配置参数
# ================================================================
doc.add_heading('九、关键配置参数', level=1)
add_table(doc,
    ['参数', '值', '说明'],
    [
        ['FLASK_HOST', '0.0.0.0', '监听地址'],
        ['FLASK_PORT', '5000', '监听端口'],
        ['PAGE_SIZE', '25', '数据表每页条数'],
        ['RENDER_LIMIT', '5000', '前端最大渲染条数'],
        ['TOP_K', '5', 'RAG检索Top-K'],
        ['EMBED_MODEL', 'text-embedding-qwen3-embedding-4b', '嵌入模型'],
        ['CHAT_MODEL', 'qwen3.6-35b-a3b', '对话模型'],
        ['CACHE_DEFAULT_TTL', '300', '缓存默认TTL（秒）'],
        ['SESSION_TIMEOUT', '3600', '登录会话超时（秒）'],
        ['LOG_MAX_BYTES', '10MB', '日志文件最大大小'],
        ['EMBEDDING_RETRIES', '3', '嵌入API重试次数'],
        ['RAG_COLLECTION_NAME', 'poems', '向量库集合名'],
        ['RAG_SPACE_FUNC', 'cosine', '向量距离函数'],
    ]
)

doc.add_page_break()

# ================================================================
# 十、文件清单
# ================================================================
doc.add_heading('十、文件清单', level=1)
file_list = """app.py                  (497行)  Flask Web 主程序
users.py                ( 91行)  用户认证管理
config.py               (118行)  全局配置
models.py               (192行)  数据模型定义
analytics.py            (312行)  数据统计分析服务
query_rag.py            (237行)  RAG 检索与流式问答引擎
build_rag.py            (121行)  RAG 向量数据库构建
admin.py                (341行)  管理后台 Blueprint
cli.py                  (239行)  CLI 命令行工具
utils.py                (145行)  通用工具函数
cache.py                (143行)  缓存模块
errors.py               (111行)  统一异常处理
validators.py           ( 83行)  请求校验
logger.py               ( 92行)  日志模块
export_service.py       (117行)  数据导出服务
middleware.py           (187行)  中间件
monitor.py              (185行)  系统监控
preprocessor.py         (148行)  数据预处理
requirements.txt        (  5行)  依赖清单
templates/base.html     (349行)  侧边栏基础布局
templates/login.html    (143行)  登录界面
templates/register.html (111行)  注册界面
templates/dashboard.html( 88行)  首页仪表盘
templates/graph_charts.html(121行) 意象统计图表页
templates/graph_table.html (568行) 溯源数据查询页
templates/ai.html       (175行)  AI智能问答页
templates/recycle.html  ( 70行)  回收站页
templates/admin.html    (262行)  管理后台页
tests/                  (1067行) 单元测试（10个文件）
poem_json/                      诗歌标注数据目录
rag_db/                         向量数据库
exports/                        导出文件目录"""
add_code(doc, file_list)

# ================================================================
# Save
# ================================================================
doc.save(OUT)
print(f"Document saved: {OUT}")
print(f"Size: {os.path.getsize(OUT) / 1024:.1f} KB")
