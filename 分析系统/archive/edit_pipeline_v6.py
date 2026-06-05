#!/usr/bin/env python3
"""Rebuild pipeline doc: complete flow with 11 screenshots, no CLI."""
import os, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn

SRC = r'C:\Users\Administrator\Desktop\新建文件夹 (14)\软著_数据生产管线_软件说明书.docx'
DST = r'C:\Users\Administrator\Desktop\pipeline_doc.docx'
IMG = r'C:\Users\Administrator\Desktop\All Mix'

shutil.copy2(SRC, DST)
doc = Document(DST)

# ─── Find sections ───
paras = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs)]
body_s18 = next(i for i, t in paras if t == '十八、CLI接口层' and i > 100)
body_s19 = next(i for i, t in paras if t == '十九、Web前端模块' and i > 100)
body_s20 = next(i for i, t in paras if t == '二十、文件清单' and i > 100)
print(f'Section 18 at {body_s18}, 19 at {body_s19}, 20 at {body_s20}')

# ─── Remove old section 18 and section 19 image paragraphs ───
# Remove section 18 paragraphs (18 heading + content)
for i in range(body_s19 - 1, body_s18 - 1, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
print('Removed old section 18')

# Remove images from section 19 body
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '二十、文件清单': break
    if i < body_s18: continue  # will shift but we'll find them
    blips = p._element.findall('.//' + qn('a:blip'))
    if blips:
        p._element.getparent().remove(p._element)
        print(f'Removed image from section 19 at paragraph {i}')

# ─── Find section 19 heading element for insertion point ───
s19_elem = None
for p in doc.paragraphs:
    if p.text.strip() == '十九、Web前端模块':
        s19_elem = p._element
        break

if s19_elem is None:
    print('ERROR: Section 19 heading not found!')
    exit(1)

print(f'Inserting before section 19 heading')

# ─── Helper functions ───
def ins_heading1(text, pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.style = doc.styles['Heading 1']
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    return p

def ins_heading2(text, pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.style = doc.styles['Heading 2']
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    return p

def ins_text(text, pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.first_line_indent = Cm(0.75)
    return p

def ins_fig(text, pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p

def ins_img(path, pos, w=5.2):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(w))
    return p

def ins_page_break(pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    r = p.add_run()
    br = etree.SubElement(r._element, qn('w:br'))
    br.set(qn('w:type'), 'page')

# ─── Screenshot data: ordered by user flow ───
screenshots = [
    ('18.1 用户注册',
     '注册页面（/register）支持新用户自助注册，需填写用户名、密码和确认密码。系统进行两项校验：（1）密码一致性检查——两次输入的密码必须一致；（2）用户名唯一性检查——用户名不可与已有用户重复。注册成功后自动登录并跳转至系统主页。',
     '图1： 用户注册页面', 'p_register.png'),
    ('18.2 用户登录',
     '登录页面（/login）提供用户认证入口，支持用户名密码验证。输入正确的用户名和密码后点击登录按钮，系统验证通过后跳转至系统主页。登录状态基于Flask session管理，未登录用户自动重定向至登录页面。密码使用SHA-256哈希存储，不保存明文密码。',
     '图2： 用户登录页面', 'p_login.png'),
    ('18.3 系统主页',
     '登录后进入系统主页（/home），页面包含五个区域：（1）导航栏——系统Logo、导航链接和当前登录用户；（2）Hero区域——系统名称"古典诗歌文本结构化标注生产系统"和功能描述；（3）功能卡片——数据导入、AI标注引擎、校验质控三大核心功能；（4）快捷操作区——新建标注任务、查看任务列表、导出数据、系统诊断；（5）API接口列表——展示所有可用API端点。',
     '图3： 系统主页', 'p_home.png'),
    ('18.4 控制台界面',
     '控制台页面（/dashboard）提供系统运行状态总览，包含三个区域：（1）统计卡片——标注数据总量、管线运行次数、今日新增标注；（2）最近管线运行——展示最近运行的任务ID、状态和时间；（3）系统功能模块——列出输入适配层、文本预处理层、核心生成引擎、结果解析层、数据校验层、质量控制层、数据输出层共七大模块及其运行状态。',
     '图4： 控制台界面', 'p_dashboard.png'),
    ('18.5 API健康检查',
     '系统提供免认证的健康检查接口GET /api/v1/system/health，返回JSON格式的运行状态信息，包括status（ok/fail）和timestamp（服务器时间戳）。该接口用于监控系统运行状态，是系统运维的基础设施之一。',
     '图5： API健康检查', 'p_api_health.png'),
    ('18.6 系统信息接口',
     '系统信息接口GET /api/v1/system/info返回系统版本、名称、运行时长、注册用户数、标注数据量和管线运行次数。所有业务接口均需用户登录认证，API返回格式统一为JSON，包含必要的状态码和错误信息。',
     '图6： 系统信息接口', 'p_api_info.png'),
    ('18.7 管线任务列表',
     '管线任务列表接口GET /api/v1/pipeline/list返回最近20条管线运行记录，每条记录包含任务ID、运行状态和创建时间。用户可通过该接口查看所有标注任务的执行历史，追踪任务进度。',
     '图7： 管线任务列表', 'p_api_pipelines.png'),
    ('18.8 标注数据列表',
     '标注数据列表接口GET /api/v1/annotations/list返回标注数据记录，每条记录包含标注ID、提交数据内容、创建时间和提交用户。支持limit参数控制返回条数，默认返回最近50条记录。',
     '图8： 标注数据列表', 'p_api_annotations.png'),
    ('18.9 标注结果详情',
     '标注结果详情页面以可视化卡片形式展示诗歌标注的完整结果：（1）诗歌原文区——展示标题、作者和诗句全文；（2）意象分析区——列出意象词及其强度评分；（3）情感分析区——展示情感类别和强度；（4）结构分析区——展示体裁和用韵信息；（5）质量评分区——展示完整度和置信度指标。',
     '图9： 标注结果详情', 'p_annotations.png'),
    ('18.10 管线运行过程',
     '执行run命令启动标注管线，系统按"输入读取→文本预处理→AI标注引擎→结果解析与校验→数据输出"五步流程自动处理。输入支持TXT/JSON/CSV/MD格式，AI引擎完成结构化标注后经过多层校验确保数据质量，最终输出标准格式的标注结果。',
     '图10： 管线运行过程', 'pipeline_run.png'),
    ('18.11 标注数据输出',
     '管线输出结构化JSON格式的标注数据，包含运行元信息（ID、状态、耗时）、诗歌原文（标题、作者、诗句）、多维标注结果（意象分析、情感分析、结构分析）及质量评分（完整度、置信度）。支持JSON/CSV/XML三种导出格式。',
     '图11： 标注数据输出', 'pipeline_output.png'),
]

# ─── Insert new content ───
# Build all paragraphs in REVERSE order (last inserted = first displayed)
pos = next(i for i, p in enumerate(doc.paragraphs)
           if p.text.strip() == '十九、Web前端模块')

# Page break before section 19
ins_page_break(pos)

# Screenshots in reverse order
for h2, desc, fig, img_file in reversed(screenshots):
    ins_img(os.path.join(IMG, img_file), pos)
    ins_fig(fig, pos)
    ins_text(desc, pos)
    ins_heading2(h2, pos)

# Section heading
ins_heading1('十八、系统运行截图', pos)

print(f'Inserted {len(screenshots)} screenshots')

# ─── Save ───
doc.save(DST)
print(f'Saved to {DST}')

# ─── Verify ───
from docx import Document as D2
d2 = D2(DST)
figs = []
for p in d2.paragraphs:
    t = p.text.strip()
    if '图' in t and '：' in t:
        figs.append(t)
print(f'Figure labels: {len(figs)}')
for f in figs:
    print(f'  {f}')
