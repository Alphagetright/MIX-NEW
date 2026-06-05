# -*- coding: utf-8 -*-
"""生成数据生产管线软著所需的docx文件"""

import os
import math
import hashlib
import time
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT_DIR = r"C:\Users\Administrator\Desktop\All Mix"
BASE_DIR = r"C:\Users\Administrator\Desktop\All Mix"
_DIAGRAM_DIR = os.path.join(OUTPUT_DIR, "_diagrams")

# PIL
try:
    from PIL import Image, ImageDraw, ImageFont
    _HAS_PIL = True
except ImportError:
    _HAS_PIL = False


def _get_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\msyhbd.ttc",
        r"C:\Windows\Fonts\simsun.ttc", r"C:\Windows\Fonts\simhei.ttf",
    ]
    for fp in candidates:
        if os.path.exists(fp):
            try:
                return ImageFont.truetype(fp, size, encoding="unic")
            except Exception:
                continue
    return ImageFont.load_default()


def _draw_rounded_box(draw, xy, text, fill, font, text_color=(255, 255, 255), radius=8):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle([x1, y1, x2, y2], radius=radius, fill=fill)
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    tx, ty = (x1 + x2 - tw) / 2, (y1 + y2 - th) / 2
    draw.text((tx, ty), text, fill=text_color, font=font)


def _draw_arrow(draw, x1, y1, x2, y2, color=(100, 100, 100), width=2):
    draw.line([(x1, y1), (x2, y2)], fill=color, width=width)
    arrow_size = 8
    dx, dy = x2 - x1, y2 - y1
    angle = math.atan2(dy, dx)
    ax1 = x2 - arrow_size * math.cos(angle - 0.4)
    ay1 = y2 - arrow_size * math.sin(angle - 0.4)
    ax2 = x2 - arrow_size * math.cos(angle + 0.4)
    ay2 = y2 - arrow_size * math.sin(angle + 0.4)
    draw.polygon([(x2, y2), (ax1, ay1), (ax2, ay2)], fill=color)


# ═══ Diagram Generators ═══

def _ensure_diagram_dir():
    os.makedirs(_DIAGRAM_DIR, exist_ok=True)


def generate_architecture_diagram():
    """系统分层架构图"""
    _ensure_diagram_dir()
    fp = os.path.join(_DIAGRAM_DIR, "pipeline_arch.png")
    w, h = 850, 520
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    ftitle = _get_font(20, bold=True)
    flayer = _get_font(15, bold=True)
    fmod = _get_font(12)
    fn = _get_font(11)

    draw.text((w // 2 - 120, 6), "数据生产管线系统架构图", fill=(50, 50, 50), font=ftitle)

    layers = [
        ("核心处理层\n(core/)", 50, (173, 216, 230), [
            "input", "preprocessing", "engine", "parsing",
            "validation", "quality", "output"
        ]),
        ("工具支撑层\n(lib/)", 200, (144, 238, 144), [
            "cache", "config", "errors", "log",
            "pipeline", "reporting", "cost", "utils"
        ]),
        ("接口层\n(cli+web)", 370, (255, 255, 150), [
            "cli_parser", "cli_commands", "cli_formatter",
            "web_app", "API", "demo"
        ]),
    ]

    box_w, box_h, gap = 100, 28, 8
    for layer_name, y_start, color, modules in layers:
        draw.text((12, y_start + 20), layer_name, fill=(60, 60, 60), font=flayer)
        n_per_row = 7
        for i, mod in enumerate(modules):
            row, col = i // n_per_row, i % n_per_row
            total_w = min(len(modules), n_per_row) * (box_w + gap) - gap
            x_start = (w - total_w) // 2 + 30
            x = x_start + col * (box_w + gap)
            y = y_start + row * (box_h + 4)
            _draw_rounded_box(draw, (x, y, x + box_w, y + box_h),
                              mod, color, fmod, text_color=(30, 30, 30))

    _draw_arrow(draw, 60, 168, 60, 190)
    _draw_arrow(draw, 60, 335, 60, 360)
    draw.text((12, 440), "数据流向: 输入 → 预处理 → 引擎生成 → 解析 → 校验 → 质控 → 输出",
              fill=(120, 120, 120), font=fn)
    draw.text((12, 458), "设计原则: 下层无上层依赖，上层依赖下层，单向依赖",
              fill=(120, 120, 120), font=fn)
    img.save(fp, "PNG")
    return fp


def generate_pipeline_flow_diagram():
    """管线数据流图"""
    _ensure_diagram_dir()
    fp = os.path.join(_DIAGRAM_DIR, "pipeline_flow.png")
    w, h = 800, 380
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    ft = _get_font(18, bold=True)
    fs = _get_font(13)
    fn = _get_font(11)

    draw.text((w // 2 - 100, 6), "管线数据流与处理阶段", fill=(50, 50, 50), font=ft)

    stages = [
        ("原始文本\n输入", (30, 60, 150, 110), (173, 216, 230)),
        ("格式检测\n编码转换", (180, 60, 300, 110), (200, 220, 240)),
        ("文本预处理\n清洗归一", (330, 60, 450, 110), (144, 238, 144)),
        ("AI引擎\n模板生成", (480, 60, 600, 110), (255, 200, 100)),
        ("结果解析\nJSON提取", (630, 60, 750, 110), (255, 180, 150)),
        ("数据校验\n质量检查", (330, 160, 450, 210), (255, 150, 150)),
        ("质量控制\n评分过滤", (480, 160, 600, 210), (200, 150, 255)),
        ("多格式输出\n归档打包", (630, 160, 750, 210), (150, 200, 150)),
    ]
    for text, box, color in stages:
        _draw_rounded_box(draw, box, text, color, fs, text_color=(30, 30, 30))

    arrows = [(150, 85, 175), (300, 85, 325), (450, 85, 475),
              (600, 85, 625), (160, 185, 325), (450, 185, 475),
              (600, 185, 625)]
    for x1, y, x2 in arrows:
        draw.line([(x1, y), (x2 - 5, y)], fill=(80, 80, 80), width=2)
        draw.polygon([(x2, y), (x2 - 8, y - 4), (x2 - 8, y + 4)], fill=(80, 80, 80))

    draw.line([(390, 110), (390, 155)], fill=(80, 80, 80), width=2)
    draw.polygon([(390, 155), (386, 148), (394, 148)], fill=(80, 80, 80))

    draw.text((12, 240), "支持断点续跑：任一阶段失败后修复可从中断处继续，无需重头执行",
              fill=(120, 120, 120), font=fn)
    draw.text((12, 258), "全程日志追踪：每个阶段记录耗时、输入输出量、异常信息",
              fill=(120, 120, 120), font=fn)
    img.save(fp, "PNG")
    return fp


def generate_cli_screenshot():
    """CLI运行模拟截图"""
    _ensure_diagram_dir()
    fp = os.path.join(_DIAGRAM_DIR, "cli_screenshot.png")
    w, h = 700, 380
    img = Image.new("RGB", (w, h), (30, 30, 40))
    draw = ImageDraw.Draw(img)
    f = _get_font(13)
    fb = _get_font(14, bold=True)

    lines = [
        ("╔══════════════════════════════════════════════╗", (100, 150, 150)),
        ("║  古典诗歌文本结构化标注生产系统  v1.0        ║", (100, 200, 150)),
        ("╚══════════════════════════════════════════════╝", (100, 200, 150)),
        ("", None),
        ("[用户] python demo.py run --input poem.txt --format json", (80, 180, 220)),
        ("", None),
        ("[系统] 正在加载配置文件...", (80, 180, 220)),
        ("[系统] 输入文件: poem.txt (UTF-8, 312 chars)", (80, 180, 220)),
        ("[系统] 启动管线: 8个阶段, 最大并发3", (80, 180, 220)),
        ("", None),
        (" > 阶段1/8 [输入适配]    ████████████ 100%", (80, 200, 120)),
        (" > 阶段2/8 [文本预处理]  ████████████ 100%", (80, 200, 120)),
        (" > 阶段3/8 [引擎生成]    ████████░░░  80%", (80, 200, 120)),
        (" > 阶段4/8 [结果解析]    ░░░░░░░░░░░   0%  (等待中)", (80, 200, 120)),
        ("", None),
        ("[系统] 当前进度: 3/8 阶段完成 | 已用: 12.3s | 预计剩余: 5.1s", (80, 180, 220)),
    ]
    y = 20
    for text, color in lines:
        if text:
            r, g, b = color if color else (180, 180, 180)
            draw.text((20, y), text, fill=(r, g, b), font=fb if "═" in text else f)
        y += 22

    draw.text((20, y + 10), "按 Ctrl+C 停止管线  |  按 R 重新运行  |  按 Q 退出",
              fill=(120, 120, 120), font=_get_font(12))
    img.save(fp, "PNG")
    return fp


def generate_web_screenshot():
    """Web界面模拟截图"""
    _ensure_diagram_dir()
    fp = os.path.join(_DIAGRAM_DIR, "web_screenshot.png")
    w, h = 750, 400
    img = Image.new("RGB", (w, h), (240, 242, 245))
    draw = ImageDraw.Draw(img)
    f = _get_font(12)
    fb = _get_font(14, bold=True)
    ft = _get_font(16, bold=True)

    # Navbar
    draw.rectangle([0, 0, w, 45], fill=(26, 26, 46))
    draw.text((15, 12), "📜 诗歌标注平台", fill=(255, 255, 255), font=fb)
    draw.text((600, 14), "欢迎，admin", fill=(180, 180, 180), font=f)

    # Hero
    draw.rectangle([15, 55, w - 15, 140], fill=(15, 52, 96))
    draw.text((30, 65), "古典诗歌文本结构化标注生产系统", fill=(255, 255, 255), font=ft)
    draw.text((30, 90), "输入原始诗歌文本 → 自动生成高维度结构化标注数据", fill=(180, 200, 220), font=f)
    draw.text((30, 110), "支持 TXT / JSON / CSV / MD 输入 | JSON / CSV / XML 输出", fill=(150, 180, 210), font=f)

    # Cards
    cards = [("📥 数据导入", "批量导入诗歌文本\n自动格式检测"), ("⚙️ AI标注引擎", "模板引擎驱动\n结构化标注生成"), ("✅ 校验质控", "多层级校验\n置信度评分")]
    for i, (title, desc) in enumerate(cards):
        x = 20 + i * 245
        draw.rectangle([x, 150, x + 230, 220], fill=(255, 255, 255), outline=(220, 220, 220))
        draw.text((x + 10, 158), title, fill=(26, 26, 46), font=fb)
        draw.text((x + 10, 183), desc, fill=(120, 120, 120), font=f)

    # API section
    draw.rectangle([15, 230, w - 15, 320], fill=(255, 255, 255), outline=(220, 220, 220))
    draw.text((25, 238), "API 接口列表", fill=(26, 26, 46), font=fb)
    apis = [
        ("POST  /api/v1/pipeline/run", "启动标注管线"),
        ("GET   /api/v1/pipeline/status/:id", "查询任务状态"),
        ("POST  /api/v1/annotations/submit", "提交标注数据"),
        ("GET   /api/v1/annotations/list", "标注数据列表"),
    ]
    for i, (endpoint, desc) in enumerate(apis):
        y = 260 + i * 18
        draw.text((25, y), endpoint, fill=(21, 101, 192), font=_get_font(11))
        draw.text((280, y), desc, fill=(100, 100, 100), font=_get_font(11))

    draw.text((15, 335), "登录注册功能：支持用户注册、登录、会话管理、权限控制",
              fill=(120, 120, 120), font=_get_font(11))
    draw.text((15, 353), "二级接口：管线任务管理（运行/状态/列表） + 标注数据管理（提交/查询/导出）",
              fill=(120, 120, 120), font=_get_font(11))
    img.save(fp, "PNG")
    return fp


def generate_api_diagram():
    """API接口结构图"""
    _ensure_diagram_dir()
    fp = os.path.join(_DIAGRAM_DIR, "api_structure.png")
    w, h = 700, 350
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    ft = _get_font(18, bold=True)
    fs = _get_font(13)
    fn = _get_font(11)

    draw.text((w // 2 - 80, 6), "API 接口层级结构", fill=(50, 50, 50), font=ft)

    # Level 1: 一级接口（管线管理）
    draw.text((20, 50), "一级接口 — 管线管理", fill=(15, 52, 96), font=_get_font(14, bold=True))
    l1 = [
        ("POST /api/v1/pipeline/run", (20, 80, 260, 115)),
        ("GET /api/v1/pipeline/status/:id", (290, 80, 530, 115)),
        ("GET /api/v1/pipeline/list", (560, 80, 680, 115)),
    ]
    for text, box in l1:
        _draw_rounded_box(draw, box, text, (173, 216, 230), _get_font(11), text_color=(30, 30, 30))

    # Level 2: 二级接口（标注数据管理）
    draw.text((20, 145), "二级接口 — 标注数据管理", fill=(46, 125, 50), font=_get_font(14, bold=True))
    l2 = [
        ("POST /api/v1/annotations/submit", (20, 175, 260, 210)),
        ("GET /api/v1/annotations/list", (290, 175, 460, 210)),
        ("GET /api/v1/annotations/:id", (490, 175, 660, 210)),
    ]
    for text, box in l2:
        _draw_rounded_box(draw, box, text, (200, 230, 200), _get_font(11), text_color=(30, 30, 30))

    # Level 3: 系统接口
    draw.text((20, 240), "系统接口 — 运维管理", fill=(230, 100, 50), font=_get_font(14, bold=True))
    l3 = [
        ("GET /api/v1/system/info", (20, 270, 220, 305)),
        ("GET /api/v1/system/health", (250, 270, 450, 305)),
        ("GET /api/v1/annotations/export/:fmt", (480, 270, 680, 305)),
    ]
    for text, box in l3:
        _draw_rounded_box(draw, box, text, (255, 220, 180), _get_font(11), text_color=(30, 30, 30))

    img.save(fp, "PNG")
    return fp


def generate_all_diagrams():
    return {
        "architecture": generate_architecture_diagram(),
        "pipeline_flow": generate_pipeline_flow_diagram(),
        "cli": generate_cli_screenshot(),
        "web": generate_web_screenshot(),
        "api": generate_api_diagram(),
    }

# Directories containing source files and their labels
SOURCE_GROUPS = {
    "core": os.path.join(BASE_DIR, "core"),
    "cli": os.path.join(BASE_DIR, "cli"),
    "lib": os.path.join(BASE_DIR, "lib"),
}


def _set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_font(run, "黑体", size=16 if level == 0 else 14 if level == 1 else 12, bold=True)
    return h


def _add_para(doc, text, size=10.5, bold=False, align=None, indent=True):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    _set_font(run, "宋体", size, bold)
    return p


def _add_bullet(doc, text, size=10.5, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    _set_font(run, "宋体", size)
    return p


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                _set_font(run, "黑体", 9, bold=True)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_font(run, "宋体", 9)
    return table


def _add_code_block(doc, code_text, size=7):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    _set_font(run, "Courier New", size)
    p.paragraph_format.line_spacing = Pt(9)
    return p


def _add_image(doc, img_path, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    return p


def _collect_all_source_files():
    """Collect all .py files from core/, cli/, lib/ directories, return sorted list of (rel_path, lines)"""
    all_files = []
    for group_name, group_dir in SOURCE_GROUPS.items():
        if not os.path.isdir(group_dir):
            continue
        for root, dirs, files in os.walk(group_dir):
            for fname in sorted(files):
                if not fname.endswith(".py"):
                    continue
                full_path = os.path.join(root, fname)
                rel_path = os.path.relpath(full_path, BASE_DIR)
                with open(full_path, "r", encoding="utf-8") as f:
                    lines = f.readlines()
                all_files.append((rel_path, lines))
    return sorted(all_files, key=lambda x: x[0])


def _build_concatenated_source():
    """Concatenate all source files into one big list of lines with file headers"""
    files = _collect_all_source_files()
    all_lines = []
    for rel_path, lines in files:
        header = f"# ===== {rel_path} =====\n"
        all_lines.append(header)
        all_lines.extend(lines)
        all_lines.append("\n")
    return all_lines, files


# ═══════════════════════════════════════════════
#  软件说明书
# ═══════════════════════════════════════════════

def generate_manual(diagrams=None):
    """生成软件说明书（>20页）"""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    # ═══ 封面 ═══
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("古典诗歌文本结构化标注生产系统")
    _set_font(run, "黑体", 26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("软 件 说 明 书")
    _set_font(run, "黑体", 22, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    for line in ["版本：1.0", "开发语言：Python 3.10+", "开发环境：Windows 10"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        _set_font(run, "宋体", 14)

    doc.add_page_break()

    # ═══ 目录 ═══
    _add_heading(doc, "目  录", level=0)
    toc_items = [
        "一、引言",
        "二、系统架构概述",
        "三、输入适配层",
        "四、文本预处理层",
        "五、核心生成引擎",
        "六、结果解析层",
        "七、数据校验层",
        "八、质量控制层",
        "九、成本与性能追踪层",
        "十、数据输出层",
        "十一、报告生成层",
        "十二、管线编排层",
        "十三、缓存与持久化层",
        "十四、异常与重试层",
        "十五、日志与诊断层",
        "十六、配置管理层",
        "十七、通用工具层",
        "十八、CLI接口层",
        "十九、Web前端模块",
        "二十、文件清单",
    ]
    for item in toc_items:
        _add_para(doc, item, size=12, indent=False)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第一章 引言
    # ════════════════════════════════════════════
    _add_heading(doc, "一、引言", level=1)

    _add_heading(doc, "1.1 项目背景", level=2)
    _add_para(doc, "古典诗歌是中华优秀传统文化的重要组成部分，对诗歌文本进行结构化标注是数字人文研究的基础性工作。传统的人工标注方式效率低、成本高、标准难以统一。本系统旨在构建一套自动化的古典诗歌文本结构化标注生产管线，通过整合大语言模型的生成能力与多级后处理校验机制，实现从原始诗歌文本到高维度结构化标注数据的高效生产。")

    _add_heading(doc, "1.2 设计目标", level=2)
    _add_para(doc, "本系统的设计目标包括：（1）支持多种输入格式的原始诗歌文本加载（TXT/JSON/CSV/MD/XML）；（2）提供完备的文本预处理能力（编码检测、文本归一、行级清洗、结构检测）；（3）基于模板引擎与大语言模型推理，生成结构化标注数据；（4）实现多层级结果解析与错误恢复机制；（5）构建全面的数据校验体系（类型检查、范围校验、交叉验证）；（6）实现质量控制与置信度评分；（7）支持多种输出格式（JSON/CSV/XML）；（8）提供完整的管线编排与断点续跑能力。")

    _add_heading(doc, "1.3 运行环境", level=2)
    _add_para(doc, "操作系统：Windows 10 / Linux / macOS。运行环境：Python 3.10 及以上版本。核心依赖：无（纯Python标准库实现）。文档生成依赖：python-docx（用于生成Word格式报告）、Pillow（用于生成示意图）。")

    _add_heading(doc, "1.4 系统定位", level=2)
    _add_para(doc, "本系统定位为数据生产环节的上游工具——输入原始诗歌文本，输出高维度结构化标注数据。系统核心提示词由外部配置文件管理，不纳入源代码范畴。本系统与已有系统形成上下游关系，功能零重叠。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第二章 系统架构
    # ════════════════════════════════════════════
    _add_heading(doc, "二、系统架构概述", level=1)

    _add_heading(doc, "2.1 总体架构", level=2)
    _add_para(doc, "本系统采用分层模块化架构设计，共包含130个Python模块（约12,665行代码），从底至上分为十六个子系统。系统采用管线式数据流设计，数据从输入层进入，依次经过预处理、引擎生成、解析、校验、质量控制等阶段，最终通过输出层导出为结构化文件。")

    _add_para(doc, "十六个子系统分别为：输入适配层、文本预处理层、核心生成引擎、结果解析层、数据校验层、质量控制层、成本与性能追踪层、数据输出层、报告生成层、管线编排层、缓存与持久化层、异常与重试层、日志与诊断层、配置管理层、通用工具层、CLI接口层。各层职责清晰，下层无上层依赖，上层依赖下层。")

    _add_heading(doc, "2.2 数据流设计", level=2)
    _add_para(doc, "数据在管线中的流转方向为：原始文本→输入适配（格式检测、编码转换）→文本预处理（清洗、归一、结构检测）→引擎生成（模板渲染、上下文组装、推理调用）→结果解析（JSON提取、清洗、字段映射）→数据校验（类型检查、范围校验、交叉验证）→质量控制（一致性检查、置信度评分）→数据输出（序列化、格式转换、归档）。管线编排层负责各阶段的调度、执行上下文传递和错误传播。")

    _add_heading(doc, "2.3 模块组织", level=2)
    _add_para(doc, "全部130个模块按功能划分为三个文件夹：（1）core/——核心处理部分，共59个模块，包含引擎、解析、校验、质量、输入、预处理、输出等核心流程；（2）cli/——命令行接口部分，共6个模块，包含CLI解析、命令实现、格式化输出和演示运行；（3）lib/——工具支撑部分，共65个模块，包含缓存、配置、成本、异常、日志、管线编排、报告、通用工具等基础设施。")

    _add_heading(doc, "2.4 设计原则", level=2)
    _add_para(doc, "系统设计遵循以下原则：（1）单一职责——每个模块聚焦于一个功能领域；（2）依赖反转——核心业务逻辑不依赖具体实现，通过接口抽象解耦；（3）可配置化——通过外部配置管理模板、策略、参数等可变部分；（4）防御性编程——完善的异常处理、重试机制和降级策略；（5）可观测性——全面的日志、性能追踪和质量度量。")

    if diagrams and "architecture" in diagrams:
        _add_image(doc, diagrams["architecture"], width_cm=14)

    if diagrams and "pipeline_flow" in diagrams:
        _add_image(doc, diagrams["pipeline_flow"], width_cm=14)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第三章 输入适配层
    # ════════════════════════════════════════════
    _add_heading(doc, "三、输入适配层", level=1)

    _add_heading(doc, "3.1 模块概述", level=2)
    _add_para(doc, "输入适配层负责将各种格式的原始诗歌文本文件加载到系统中，包含6个模块：input_reader（多格式读取）、encoding_detector（编码检测）、text_normalizer（文本归一）、text_splitter（文本分割）、batch_reader（目录扫描）、file_validator（文件校验）。")

    _add_heading(doc, "3.2 多格式读取", level=2)
    _add_para(doc, "InputReader类支持自动检测并读取五种格式的输入文件：TXT格式（纯文本，按行读取）、JSON格式（支持自动识别诗歌字段，如poems/诗歌集等关键词）、CSV格式（支持表头检测和列选择）、MD格式（Markdown文本）、XML格式（ElementTree解析）。FormatDetector类通过文件扩展名和内容特征进行格式识别，采用职责链模式依次尝试各格式解析器。")

    _add_heading(doc, "3.3 编码检测与转换", level=2)
    _add_para(doc, "EncodingDetector类实现了完整的编码检测流程：首先检测BOM标记（支持UTF-8/UTF-16LE/UTF-16BE/UTF-32LE/UTF-32BE五种BOM格式），然后使用编码启发式规则推断编码类型。EncodingConverter支持从任意检测到的编码向UTF-8的自动转换。EncodingValidator对转换后的文本进行编码正确性验证，确保无乱码残留。")

    _add_heading(doc, "3.4 文本归一化", level=2)
    _add_para(doc, "TextNormalizer通过可配置的步骤管线对文本进行归一化处理，支持以下步骤：WhitespaceNormalizer（空白字符统一化，合并连续空白）、PunctuationNormalizer（标点符号标准化，中文标点→英文标点转换）、UnicodeNormalizer（Unicode规范化，支持NFC/NFD/NFKC/NFKD四种形式）、ControlCharFilter（控制字符过滤，保留常见格式控制字符）、FullWidthConverter（全角字符转半角）。")

    _add_heading(doc, "3.5 分块与批量读取", level=2)
    _add_para(doc, "TextSplitter提供四种分割策略：LineSplitter（按行分割）、SentenceSplitter（按句号分割）、StanzaDetector（按诗联/阙的标记分割）、ChunkSplitter（按最大块大小控制分割）。BatchReader支持目录级别的批量扫描，提供文件扩展名过滤（.txt/.json/.csv/.md）、文件大小范围过滤、时间范围过滤和忽略模式匹配功能，支持遍历时的进度回调。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第四章 文本预处理层
    # ════════════════════════════════════════════
    _add_heading(doc, "四、文本预处理层", level=1)

    _add_heading(doc, "4.1 模块概述", level=2)
    _add_para(doc, "文本预处理层对原始文本进行精细化清洗和结构分析，包含6个模块：line_cleaner（行级清洗）、punctuation_normalizer（标点归一）、metadata_guesser（元数据猜测）、structure_detector（结构检测）、quality_prefilter（质量预过滤）、preprocessing_pipeline（预处理编排）。")

    _add_heading(doc, "4.2 行级清洗", level=2)
    _add_para(doc, "LineCleaner提供五种清洗功能：StripCleaner（去除行首行尾空白）、EmptyLineFilter（过滤空行）、CommentRemover（去除注释行，支持//、#、--、<!-- -->四种注释格式）、NumberingStripper（去除行号前缀，支持数字+点、数字+括号、汉字数字等多种序号格式）、LineCleaner（按顺序组合清洗步骤）。")

    _add_heading(doc, "4.3 元数据猜测", level=2)
    _add_para(doc, "MetadataGuesser综合使用多种策略从文本中猜测元数据信息：TitleGuesser通过行位置和格式模式识别标题（四类模式匹配）、AuthorGuesser通过常见作者标记（如唐·白居易、杜甫等）和内置22位唐代诗人姓名表识别作者、GenreGuesser通过体裁关键词（如五言绝句、七言律诗等）和文本统计特征（行数、每行字数）推断诗歌体裁。")

    _add_heading(doc, "4.4 结构检测", level=2)
    _add_para(doc, "StructureDetector对诗歌文本进行多维结构分析：LineCounter（行数统计）、CharCounter（字数统计，含字频分析）、PatternDetector（异常模式检测，包括行长度异常、重复行、过短行等）、CoupletDetector（联/阙模式识别，基于押韵和语义相似度分析）、StructureDetector（综合输出结构分析报告）。")

    _add_heading(doc, "4.5 质量预过滤", level=2)
    _add_para(doc, "QualityPrefilter在文本进入核心处理前进行质量筛查：LengthFilter（长度阈值过滤，过短的文本被标记）、EmptyFilter（空内容检测）、GibberishDetector（乱码检测，通过字符分布的熵值评估文本质量）、EncodingQualityChecker（编码质量检查）、QualityPrefilter（综合质量判定，支持自定义阈值）。PreprocessingPipeline提供完整的预处理步骤编排能力，支持步骤的注册、插入、移除、启用/禁用和条件跳过。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第五章 核心生成引擎
    # ════════════════════════════════════════════
    _add_heading(doc, "五、核心生成引擎", level=1)

    _add_heading(doc, "5.1 模块概述", level=2)
    _add_para(doc, "核心生成引擎是本系统的关键组成部分，负责将预处理后的诗歌文本转换为结构化标注数据。包含8个模块：template_loader、template_renderer、template_manager、context_assembler、request_builder、inference_client、retry_policy、response_collector。")

    _add_heading(doc, "5.2 模板系统", level=2)
    _add_para(doc, "模板系统由TemplateLoader、TemplateRenderer和TemplateManager三个模块组成。TemplateLoader支持从外部配置文件加载YAML/JSON/TXT格式的模板，自动检测文件格式。TemplateRenderer实现了完整的模板渲染能力，包括VariableReplacer（双花括号变量替换{{var}}）、ConditionalBlock（条件块{% if %} ... {% endif %}，支持if-else分支）和LoopBlock（循环展开{% for item in list %} ... {% endfor %}）。TemplateManager提供模板注册、版本管理、选择策略（latest/stable）和多模板管理功能。")

    _add_heading(doc, "5.3 上下文组装", level=2)
    _add_para(doc, "ContextAssembler负责构建发送给大语言模型的完整上下文，支持三种组装模式：（1）assemble方法——将模板内容、示例和诗歌文本拼接为纯文本格式；（2）assemble_messages方法——构建OpenAI格式的消息列表（role: system/user）；（3）通过SystemInstructionBuilder灵活构建系统指令，支持添加约束（add_constraint）、示例（add_example）和格式说明（add_format_spec）。")

    _add_heading(doc, "5.4 请求构建与推理调用", level=2)
    _add_para(doc, "RequestBuilder封装API请求的构建过程，支持MessageBuilder链式调用构建消息列表，APIParameters管理温度、最大Token数、TopP等推理参数。InferenceClient提供HTTP推理调用功能，包含TimeoutConfig超时管理、ConnectionStats连接统计、StreamHandler流式响应处理以及ResponseParser响应解析功能。支持单次推理和批量推理两种模式。")

    _add_heading(doc, "5.5 重试策略", level=2)
    _add_para(doc, "引擎的重试策略专门针对推理调用场景设计，定义了专用的异常体系（RateLimitError、ServerOverloadError、TokenLimitError、TimeoutError）。RetryPolicy支持可配置的最大重试次数、基础延迟、最大延迟和抖动开关。RetryExecutor提供同步和异步两种执行方式，自动捕获异常、计算延迟（指数退避+随机抖动）、判断可重试性和超时判定。")

    _add_heading(doc, "5.6 响应收集", level=2)
    _add_para(doc, "ResponseCollector负责收集和处理推理响应。ResponseBuffer管理流式响应块的缓冲，ResponseAssembler支持文本组装、JSON解析和自定义解析函数。ResponseMetadata记录模型信息、完成原因、Token用量、耗时等元数据。RawResponseCache提供基于LRU的原始响应缓存，支持按内容键值缓存和访问顺序淘汰。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第六章 结果解析层
    # ════════════════════════════════════════════
    _add_heading(doc, "六、结果解析层", level=1)

    _add_heading(doc, "6.1 模块概述", level=2)
    _add_para(doc, "结果解析层负责将大语言模型返回的原始文本解析为结构化数据，包含8个模块：json_extractor、json_cleaner、json_validator、field_mapper、nested_builder、error_recovery、parsing_metrics、batch_parser。")

    _add_heading(doc, "6.2 JSON提取与清洗", level=2)
    _add_para(doc, "JsonExtractor从模型返回的文本中检测和提取JSON内容，支持代码围栏检测（```json ... ```）、边界自动识别和多JSON块提取。JsonCleaner提供四级清洗能力：CommentRemover（删除JSON中的//和/**/注释，非标准JSON但常见于LLM输出）、TrailingCommaFixer（修复数组和对象末尾的多余逗号）、QuoteFixer（修复未转义的引号）、TruncationRecovery（检测并修复被截断的JSON，自动补全缺失的括号和引号）。CleanerPipeline支持按需组合清洗步骤。")

    _add_heading(doc, "6.3 字段映射与嵌套构建", level=2)
    _add_para(doc, "FieldMapper实现从源字段到目标字段的灵活映射，支持MappingRule配置源路径、目标路径、默认值和类型转换器。PathResolver通过点分路径在嵌套字典中定位值；TypeConverter提供安全的str/int/float/bool类型转换。NestedBuilder从扁平映射构建层级结构，ListAssembler支持列表装配、DictMerger支持字典合并、TreeBuilder通过父-子关系构建树结构。")

    _add_heading(doc, "6.4 错误恢复与批量解析", level=2)
    _add_para(doc, "ErrorRecovery提供三级恢复策略链：RetryStrategy（JSON修复重试，尝试多种修复组合）、PartialResultStrategy（字段级部分结果，成功字段保留，失败字段用默认值）、DegradeStrategy（降级处理，整体失败时返回最小有效结果）。BatchParser使用WorkerPool线程池进行并发解析，支持进度追踪和结果归集。ParsingMetrics记录各类解析指标，包括成功率、错误类型分布和字段级成功率。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第七章 数据校验层
    # ════════════════════════════════════════════
    _add_heading(doc, "七、数据校验层", level=1)

    _add_heading(doc, "7.1 模块概述", level=2)
    _add_para(doc, "数据校验层对解析后的结构化数据进行多维度校验，确保数据质量符合预期标准。包含8个模块：schema_definition、type_checker、required_checker、range_checker、cross_validator、integrity_reporter、threshold_config、validation_summary。")

    _add_heading(doc, "7.2 Schema定义与类型检查", level=2)
    _add_para(doc, "SchemaDefinition提供完整的字段规格定义能力，包括FieldSpec（字段名、类型、必填性、默认值、约束条件）、SchemaRegistry（Schema注册与查找）、以及多种内置约束（RangeConstraint数值范围、LengthConstraint字符串长度、EnumConstraint枚举取值、PatternConstraint正则模式）。TypeChecker对字段值进行递归类型检查，支持str/int/float/bool/list/dict六种基本类型，TypeConverter提供带错误处理和默认值的类型转换。")

    _add_heading(doc, "7.3 必填检查与范围校验", level=2)
    _add_para(doc, "RequiredChecker支持三种必填规则：简单必填（字段必须存在）、嵌套必填（嵌套结构中字段必须存在）、条件必填（ConditionBuilder构建的条件触发型必填规则，如A字段为某值时B字段必填）。RangeChecker针对数值、字符串、列表和枚举四类数据类型进行范围校验：RangeRule（数值最小/最大值，支持开闭区间）、LengthRule（字符串/列表长度范围）、EnumRule（枚举值白名单校验）。")

    _add_heading(doc, "7.4 交叉验证与汇总", level=2)
    _add_para(doc, "CrossValidator实现字段间依赖和引用一致性验证，包含DependencyRule（字段间依赖规则，如A=X时B=Y）、ReferenceRule（引用一致性规则，字段值必须在引用列表中存在）、BusinessRule（自定义业务规则函数）。ValidationGraph构建字段间的依赖关系图，支持循环依赖检测。IntegrityReporter计算覆盖率统计、缺失分布和完整度评分。ValidationSummary汇总校验结果，包含通过率、错误分类、Top错误排序和趋势比较。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第八章 质量控制层
    # ════════════════════════════════════════════
    _add_heading(doc, "八、质量控制层", level=1)

    _add_heading(doc, "8.1 模块概述", level=2)
    _add_para(doc, "质量控制层对整批数据的质量进行综合评估和控制，包含8个模块：consistency_checker、pattern_detector、confidence_scorer、error_taxonomy、severity_evaluator、correction_builder、quality_summary、dashboard_data。")

    _add_heading(doc, "8.2 一致性检查与模式检测", level=2)
    _add_para(doc, "ConsistencyChecker执行跨记录对比分析：DuplicateDetector基于关键字段检测重复记录；DistributionAnalyzer分析各字段的值分布统计信息（最大值、最小值、均值、中位数、标准差）；OutlierDetector使用IQR（四分位距）方法识别统计异常值。PatternDetector检测数据中的系统性问题：ErrorPattern通过注册的错误模式进行匹配；BiasDetector使用信息熵检测分类偏差；AnomalyDetector通过Z-Score检测时序异常。")

    _add_heading(doc, "8.3 置信度评分", level=2)
    _add_para(doc, "ConfidenceScorer计算每条数据的综合置信度，评分公式：综合分 = 字段完善率 × 一致性 × 历史准确率。CompletenessFactor计算必填字段的填充比例；ConsistencyFactor评估字段值在跨记录范围内的分布一致性；HistoricalFactor比较当前数据与历史批次的一致性。最终得分归一化至[0,1]区间。")

    _add_heading(doc, "8.4 错误分类与严重评估", level=2)
    _add_para(doc, "ErrorTaxonomy建立层级化错误分类体系，支持ErrorCategory的父子层次结构、SeverityLevel严重级别枚举（critical/major/minor/warning/info）、RootCauseMapper根因映射和FixStrategyMapper修复策略映射。SeverityEvaluator从影响范围、修复成本、数据质量影响和风险等级四个维度综合评估每个错误的严重程度，生成SeverityMatrix综合矩阵。")

    _add_heading(doc, "8.5 修正建议与质量汇总", level=2)
    _add_para(doc, "CorrectionBuilder基于检测到的质量问题生成自动修复方案。FixSuggestion描述单个修复建议，BatchFixPlan按策略分组的批量修复方案，AutoFixer自动应用常见修复。QualitySummary生成总体质量报告，包含QualityScore（综合评分和各维度评分）、TrendIndicator（改善/稳定/下滑趋势指示）和ImprovementSuggestion（可操作的改进建议）。DashboardData为看板展示提供格式化数据。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第九章 成本与性能追踪层
    # ════════════════════════════════════════════
    _add_heading(doc, "九、成本与性能追踪层", level=1)

    _add_heading(doc, "9.1 模块概述", level=2)
    _add_para(doc, "成本与性能追踪层对管线运行过程中的Token消耗、计算成本、性能指标进行追踪分析，包含7个模块：token_counter、cost_estimator、performance_timer、throughput_meter、usage_logger、budget_controller、cost_report。")

    _add_heading(doc, "9.2 Token计数与成本估算", level=2)
    _add_para(doc, "TokenCounter提供多种估算策略：字符基础估算（面向CJK文字约每2字符1Token）、模型特定估算（按不同模型配置的比率）、词基础估算。TokenStats维护按模型和输入类型分类的累计统计信息。CostEstimator基于PriceConfig（按模型配置的每千Token单价）计算输入/输出分别计费的成本，支持批量的汇总计算和货币单位配置。")

    _add_heading(doc, "9.3 性能计时与吞吐计量", level=2)
    _add_para(doc, "PerformanceTimer支持阶段级别的性能计时，通过TimerContext上下文管理器自动计时代码块，PercentileCalculator计算P50/P95/P99等百分位耗时，TimingReport输出各命名阶段的详细统计。FlameNode支持火焰图数据采集。ThroughputMeter实时计量管线吞吐量，支持条/秒、字/秒、Token/秒三种速率单位，SlidingWindow滑动窗口算法计算实时速率。")

    _add_heading(doc, "9.4 预算控制与成本报告", level=2)
    _add_para(doc, "UsageLogger提供结构化的用量记录，支持按时间/操作/元数据维度的查询分析和日/周/月趋势聚合。BudgetController实现预算控制功能，支持配置多个预算限额（总预算、周期、告警阈值），跟踪实际支出并与预算对比，超限时触发告警。CostReport生成多维度成本报告，包含CostDetail（逐项成本明细）、CostTrend（时序趋势数据）、CostComparison（多次运行对比）和ROICalculator（成本vs质量改进的ROI估算）。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十章 数据输出层
    # ════════════════════════════════════════════
    _add_heading(doc, "十、数据输出层", level=1)

    _add_heading(doc, "10.1 模块概述", level=2)
    _add_para(doc, "数据输出层负责将结构化标注数据以多种格式写入文件系统，包含8个模块：json_serializer、csv_serializer、xml_serializer、formatter_registry、batch_writer、output_merger、output_validator、file_archiver。")

    _add_heading(doc, "10.2 JSON序列化", level=2)
    _add_para(doc, "JsonSerializer支持中文编码控制（选择使用全ASCII转义或保留中文字符）、缩进格式控制、排序键控制和大文件流式写入。SerializeConfig提供完整的配置选项。StreamWriter逐块写入大型JSON结构，避免内存溢出。ShardManager支持将超大JSON文件自动分割为多个分片。")

    _add_heading(doc, "10.3 CSV与XML序列化", level=2)
    _add_para(doc, "CsvSerializer提供完整的CSV写入能力，包括UTF-8 BOM写入（确保Excel正确识别中文）、引号策略选择（始终/必要/从不）、分隔符配置（逗号/制表符/自定义）和多Sheet支持（通过分文件模拟）。XmlSerializer支持XML格式输出，包括编码声明、命名空间处理、CDATA包裹和格式化输出（pretty-print）。XmlElement提供简洁的XML元素构建API。")

    _add_heading(doc, "10.4 批量写入与合并", level=2)
    _add_para(doc, "BatchWriter实现高效的批量写入：WriteBuffer在内存中累积到阈值后批量落盘；RetryWriter在写入失败时自动重试；WriteProgress实时追踪写入进度。OutputMerger支持多文件合并，提供concat连接、interleave交错和merge-by-key三种合并策略，以及keep_first/keep_last/keep_all_unique三种去重策略。FileArchiver实现输出文件的归档管理，支持ZIP打包、时间戳命名、目录结构组织和旧文件清理。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十一章 报告生成层
    # ════════════════════════════════════════════
    _add_heading(doc, "十一、报告生成层", level=1)

    _add_heading(doc, "11.1 模块概述", level=2)
    _add_para(doc, "报告生成层将管线运行的结果和质量信息汇总为多种格式的报告文档，包含8个模块：run_report、quality_report、cost_report、comparison_report、text_report、json_report、html_report、report_archiver。")

    _add_heading(doc, "11.2 运行报告与质量报告", level=2)
    _add_para(doc, "RunReport记录每次管线运行的基本信息，包括运行元数据（时间戳、持续时间、命令、配置快照）和结果统计（输入数量、成功/失败数、平均耗时）。RunReportBuilder通过ReportSection构建分节式报告。QualityReport汇总质量校验结果，包括分维度质量评分（QualityMetrics）、错误类型分布（ErrorDistribution）和字段完善度统计表（CompletenessTable）。")

    _add_heading(doc, "11.3 多格式输出", level=2)
    _add_para(doc, "系统支持四种报告输出格式：TextReport生成纯文本格式，包含TextTable（列对齐的ASCII表格）、TextAlignment（左/中/右对齐）和TextDivider（水平分隔线）等格式化工具；JsonReport生成JSON结构化的报告数据，适合程序化读取和二次处理；HtmlReport生成带CSS样式的独立HTML页面，支持表格和图表占位，可直接在浏览器中查看。")

    _add_heading(doc, "11.4 对比分析与归档", level=2)
    _add_para(doc, "ComparisonReport支持多维度对比：BatchCompare（批次间比较）、ConfigCompare（配置差异比较）、StrategyCompare（不同策略效果对比），ComparisonTable以表格形式呈现对比结果。ReportArchiver实现报告的自动归档管理，支持NamingStrategy（多种命名策略）、RetentionPolicy（保留策略，按数量/时间/重要性）和ArchiveIndex（归档索引，基于hashlib的搜索）。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十二章 管线编排层
    # ════════════════════════════════════════════
    _add_heading(doc, "十二、管线编排层", level=1)

    _add_heading(doc, "12.1 模块概述", level=2)
    _add_para(doc, "管线编排层是整个系统的调度中枢，负责将各处理阶段编排为可执行的流水线。包含9个模块：stage_definition、pipeline_builder、pipeline_runner、task_definition、task_queue、checkpoint_store、resume_handler、progress_tracker、pipeline_metrics。")

    _add_heading(doc, "12.2 阶段定义与管线构建", level=2)
    _add_para(doc, "StageDefinition定义管线的单个处理阶段，包含StageIO（输入/输出端口声明）、StageDependency（阶段间依赖关系）、StageConfig（配置接口）和StageStatus枚举（pending/running/completed/failed/skipped）。PipelineBuilder通过DependencyGraph构建DAG执行图，使用TopologicalSorter（Kahn算法）对阶段进行拓扑排序，并自动检测循环依赖。BuildResult包含完整的图结构、执行顺序和警告信息。")

    _add_heading(doc, "12.3 管线执行与任务管理", level=2)
    _add_para(doc, "PipelineRunner负责阶段的实际调度执行，包含RunContext（阶段间共享的执行上下文，支持结果传递）、StageExecutor（单阶段执行器，处理输入输出绑定）、ErrorPropagator（错误传播器，终止依赖该阶段的所有下游阶段）、RunnerStats（执行统计）。TaskDefinition和TaskQueue提供任务级管理能力，支持FIFO/优先级/LIFO三种队列模式，PersistentQueue支持队列的磁盘持久化，ConcurrencyController控制最大并发任务数。")

    _add_heading(doc, "12.4 断点续跑", level=2)
    _add_para(doc, "CheckpointStore实现执行进度的序列化存储与恢复。Checkpoint记录管线执行状态的快照，包含每个阶段的完成状态和中间结果；ConsistencyCheck通过SHA-256完整性校验确保检查点的可靠性。ResumeHandler支持完整的恢复流程：CheckpointDetector自动检测最近的检查点、ResultCollector归集已完成阶段的输出结果、CompletedChecker验证各阶段的完成状态。支持全量恢复、部分恢复和演练三种恢复策略。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十三章 缓存与持久化
    # ════════════════════════════════════════════
    _add_heading(doc, "十三、缓存与持久化层", level=1)

    _add_heading(doc, "13.1 模块概述", level=2)
    _add_para(doc, "缓存与持久化层为系统提供多级缓存能力和数据持久化支持，包含7个模块：cache_entry、memory_cache、disk_cache、cache_policy、cache_manager、result_cache、template_cache。")

    _add_heading(doc, "13.2 内存缓存与磁盘缓存", level=2)
    _add_para(doc, "MemoryCache基于OrderedDict实现LRU（最近最少使用）淘汰策略，支持TTL过期、容量上限（最大条目数）、线程安全（RLock可重入锁）和过期条目自动清理（expire_stale）。DiskCache采用文件存储方式，以MD5哈希作为文件键名，JSON序列化存储值，支持目录散列管理（避免单目录文件过多）和TTL过期。CacheEntry封装单个缓存条目，记录访问计数和时间戳。")

    _add_heading(doc, "13.3 缓存策略与管理", level=2)
    _add_para(doc, "CachePolicy提供四种策略体系：TTLPolicy（生存时间策略，支持固定TTL和自适应TTL）、EvictionPolicy（淘汰策略，含LRU/LFU/FIFO三种算法）、CapacityPolicy（容量策略，按条目数或内存大小限制）、TieredCachePolicy（分级策略，配置热/温/冷三级缓存）。CacheManager整合内存和磁盘两级缓存，提供统一的get/set/delete接口，自动在两级间查找和写入。")

    _add_heading(doc, "13.4 专用缓存", level=2)
    _add_para(doc, "ResultCache基于内容哈希实现去重缓存，对相同输入的请求自动返回缓存结果，支持幂等性保障（IdempotencyGuard防止同一条数据被重复处理）。TemplateCache专门缓存预编译的模板对象，支持版本跟踪（warm_up预加载、get_or_compile懒加载编译、invalidate失效处理）和is_fresh版本新鲜度检查。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十四章 异常与重试
    # ════════════════════════════════════════════
    _add_heading(doc, "十四、异常与重试层", level=1)

    _add_heading(doc, "14.1 模块概述", level=2)
    _add_para(doc, "异常与重试层提供系统的异常处理框架和容错机制，包含5个模块：exceptions、error_handler、retry_policy、circuit_breaker、fallback_handler。")

    _add_heading(doc, "14.2 分层异常体系", level=2)
    _add_para(doc, "系统采用分层异常体系设计，以PipelineError（基类，错误码1000）为根，向下分为六大子类和22种具体异常：InputError（输入异常，含EncodingError、FormatError）、ParseError（解析异常，含JSONExtractError、JSONRepairError、FieldMappingError）、ValidateError（校验异常，含SchemaError、TypeCheckError、CrossValidateError）、GenerateError（生成异常，含TemplateError、InferenceError、RetryExhaustedError）、OutputError（输出异常，含SerializeError、WriteError）、ConfigError（配置异常，含ConfigLoadError、ConfigValidateError），以及其他辅助异常（CacheError、RetryError、CircuitBreakerError、LogError、PipelineBuildError、PipelineRunError）。ERROR_CODE_MAP提供错误码到异常类的双向映射。")

    _add_heading(doc, "14.3 重试策略与熔断器", level=2)
    _add_para(doc, "RetryPolicy支持三种重试延迟算法：FixedIntervalRetry（固定间隔）、ExponentialBackoffRetry（指数退避，含随机抖动）、LinearBackoffRetry（线性递增）。RetryExecutor提供同步和异步两种执行方式。CircuitBreaker实现熔断器模式，维护CLOSED/OPEN/HALF_OPEN三种状态，支持failure_threshold失败阈值、recovery_timeout恢复超时、half_open_max_calls半开状态最大探测请求数，以及CircuitBreakerRegistry熔断器注册管理。")

    _add_heading(doc, "14.4 降级处理", level=2)
    _add_para(doc, "FallbackHandler提供四种降级策略：DefaultValueStrategy（在调用失败时返回配置的默认值）、CacheFallbackStrategy（优先从缓存获取，缓存未命中时才执行实际调用）、EmptyResultStrategy（返回空结果结构）、GracefulDegrader（组合多级降级策略，按优先级依次尝试，高级策略失败后自动降级到低级策略）。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十五章 日志与诊断
    # ════════════════════════════════════════════
    _add_heading(doc, "十五、日志与诊断层", level=1)

    _add_heading(doc, "15.1 模块概述", level=2)
    _add_para(doc, "日志与诊断层提供系统的日志记录、审计追踪和诊断信息收集能力，包含6个模块：log_config、log_writer、structured_logger、audit_logger、diagnostic_collector、log_cleaner。")

    _add_heading(doc, "15.2 日志配置与写入", level=2)
    _add_para(doc, "LogConfig提供统一的日志配置管理，支持formatter（格式器）、handler（处理器）和logger（记录器）三级配置体系，DEFAULT_LOG_CONFIG内置默认配置。LogWriter支持三种输出目标：ConsoleWriter（控制台输出，适合交互式运行）、FileWriter（文件写入，支持按大小滚动）、MultiWriter（同时输出到多个目标）。LogRotator提供时间滚动和按大小滚动的日志文件管理。")

    _add_heading(doc, "15.3 结构化日志与审计", level=2)
    _add_para(doc, "StructuredLogger提供结构化日志记录能力，通过thread-local变量实现trace_id追踪ID的自动注入和上下文绑定，支持性能埋点。LoggerManager以单例模式管理全局日志实例。AuditLogger提供操作审计功能，AuditEntry记录每次操作的时间戳、操作类型、操作人和数据变更详情，支持MemoryAuditBackend和FileAuditBackend两种存储后端。")

    _add_heading(doc, "15.4 诊断收集与日志清理", level=2)
    _add_para(doc, "DiagnosticCollector收集运行时诊断信息，包括平台信息（操作系统、版本、架构）、环境信息（Python版本、PATH、环境变量）、配置快照（当前加载的配置）和错误快照（最近错误的上下文）。LogCleaner提供日志文件的清理策略管理，支持按文件数量、按总大小、按保留天数的三种清理策略，以及旧日志的压缩归档。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十六章 配置管理
    # ════════════════════════════════════════════
    _add_heading(doc, "十六、配置管理层", level=1)

    _add_heading(doc, "16.1 模块概述", level=2)
    _add_para(doc, "配置管理层负责系统配置的加载、校验、监视和序列化，包含7个模块：config_source、config_loader、config_validator、config_watcher、profile_manager、config_serializer、config_reporter。")

    _add_heading(doc, "16.2 多源配置加载", level=2)
    _add_para(doc, "ConfigSource定义了配置源接口，系统支持五种配置源：DefaultConfigSource（默认配置，硬编码的基础参数）、FileConfigSource（从JSON/YAML文件加载配置）、EnvConfigSource（从环境变量读取配置，支持前缀过滤）、CLIConfigSource（从命令行参数解析配置）、DictConfigSource（直接从Python字典创建配置源）。各配置源具有优先级顺序，高优先级覆盖低优先级配置。")

    _add_heading(doc, "16.3 配置加载与校验", level=2)
    _add_para(doc, "ConfigLoader通过add_source/add_defaults/add_file/add_env/add_cli_args方法链式添加配置源，load方法执行多源合并（deep_merge深度合并策略）并返回统一配置。ConfigAccessor提供安全的嵌套键访问（如get('engine.temperature')）。ConfigValidator通过add_rule方法逐条添加校验规则，validate方法执行全部校验并返回ConfigValidationReport（含校验结果、警告和错误列表）。")

    _add_heading(doc, "16.4 配置监视与管理", level=2)
    _add_para(doc, "ConfigWatcher监视配置文件的变化，支持文件修改检测和自动热加载。ConfigReloader封装了重新加载配置并通知监听器的完整流程。ProfileManager支持多profile（配置集）管理，允许用户保存、切换和删除不同的配置方案，profile之间支持继承（profile可基于另一个profile扩展）。ConfigSerializer提供配置的导出和导入功能，ConfigReporter生成配置报告，支持文本/JSON/表格三种格式和配置差异对比。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十七章 通用工具层
    # ════════════════════════════════════════════
    _add_heading(doc, "十七、通用工具层", level=1)

    _add_heading(doc, "17.1 模块概述", level=2)
    _add_para(doc, "通用工具层为全系统提供基础工具函数和实用类，包含8个模块：string_utils、file_utils、dict_utils、list_utils、time_utils、validation_utils、serialization_utils、system_utils。")

    _add_heading(doc, "17.2 字符串与文件工具", level=2)
    _add_para(doc, "StringUtils提供20+字符串处理函数，包括truncate截断、pad填充、strip_all全面去除空白、normalize_spaces空白归一化、escape_regex正则转义、unicode_normalize（NFC/NFD/NFKC/NFKD）、extract_chinese提取中文字符、is_chinese_char中文判断、split_by_punctuation按标点分割等。StringBuilder提供高效的字符串拼接。FileUtils提供文件系统操作工具，包括ensure_dir递归创建目录、get_file_size/extension/name/path等路径操作、read/write_text_file和read/write_json_file读写封装、file_hash文件哈希计算、TempDir临时目录上下文管理器、FileLock跨进程文件锁。")

    _add_heading(doc, "17.3 字典与列表工具", level=2)
    _add_para(doc, "DictUtils提供15+字典操作函数，包括deep_get/deep_set嵌套键路径存取、deep_merge深度合并（递归合并嵌套字典，支持覆盖/递归/保留三种策略）、filter_keys/exclude_keys/pick键过滤、flatten/unflatten展开与还原、sort_by_key/sort_by_value排序、rename_key重命名键、DefaultDict带默认值的字典工厂。ListUtils提供15+列表操作函数，包括chunk分块、flatten扁平化（支持多级嵌套）、dedup去重（稳定保留首次出现）、group_by键分组、top_n前N个、sample随机采样、partition分区（满足条件/不满足条件）、sliding_window滑动窗口、Paginator分页器。")

    _add_heading(doc, "17.4 时间与序列化工具", level=2)
    _add_para(doc, "TimeUtils提供时间相关工具，包括now/now_str当前时间获取、format_timestamp格式化（支持多种格式模板）、parse_time解析（自动适应多种格式）、time_diff时间差计算、format_duration友好化的持续时间显示、Timer计时器上下文管理器、RateLimiter速率限制器（令牌桶算法）、Throttle节流器。SerializationUtils提供序列化相关工具，包括deep_clone深拷贝、to_json/from_json的JSON编解码、to/from_base64的Base64编解码、to/from_hex的十六进制编解码、compute_hash/compute_file_hash的哈希计算。SystemUtils提供系统信息查询功能，包括平台检测、CPU核心数、进程信息、环境变量管理、内存使用查询。")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十八章 CLI接口层
    # ════════════════════════════════════════════
    _add_heading(doc, "十八、CLI接口层", level=1)

    _add_heading(doc, "18.1 模块概述", level=2)
    _add_para(doc, "CLI接口层为系统提供命令行操作界面，包含5个模块：cli_parser、cli_commands、cli_formatter、demo_runner和入口文件demo.py。")

    _add_heading(doc, "18.2 命令解析与注册", level=2)
    _add_para(doc, "CliParser基于简单的命令注册机制实现子命令解析。CommandRegistry维护命令注册表，支持register/discover/lookup操作。SubCommand定义命令名称、处理器函数、描述信息和参数列表。HelpFormatter自动生成全局帮助和逐命令帮助文本。支持--version版本显示和help <command>详细帮助。")

    _add_heading(doc, "18.3 命令实现", level=2)
    _add_para(doc, "CliCommands提供七个子命令的完整实现：RunCommand（执行管线处理，带进度显示）、ValidateCommand（校验数据文件和配置）、ReportCommand（生成指定格式的报告）、CleanCommand（清理缓存和临时文件）、ConfigCommand（管理配置参数，支持get/set/list）、StatsCommand（显示管线运行统计）、HelpCommand（显示帮助信息）。各命令遵循统一的execute(**kwargs)接口。")

    _add_heading(doc, "18.4 格式化与演示", level=2)
    _add_para(doc, "CliFormatter提供命令行输出格式化能力：ColorFormatter（ANSI颜色支持，自动检测终端能力，支持NO_COLOR环境变量）、TableFormatter（带对齐的ASCII表格）、ProgressBar（文本进度条[\#\#\#\#----]）、SummaryFormatter（键值对摘要块）、OutputWriter（统一输出到文件或控制台）。DemoRunner提供快速演示功能：自动生成测试样本数据、执行快速管线演示、输出DemoReport演示报告。")

    _add_para(doc, "入口文件demo.py使用标准库argparse实现命令行参数解析，支持七个子命令的分发执行，包含错误处理包装器和友好的退出信息。")

    if diagrams and "cli" in diagrams:
        _add_image(doc, diagrams["cli"], width_cm=14)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第十九章 Web前端模块
    # ════════════════════════════════════════════
    _add_heading(doc, "十九、Web前端模块", level=1)

    _add_heading(doc, "19.1 模块概述", level=2)
    _add_para(doc, "Web前端模块为系统提供基于Flask框架的轻量级Web界面，包含登录注册、主页控制台和RESTful API接口。用户通过浏览器即可完成标注任务的提交、查看和管理，无需使用命令行。该模块包含5个文件：web/app.py（Flask应用主逻辑）、web/templates/login.html（登录页）、web/templates/register.html（注册页）、web/templates/home.html（主页）、web/templates/dashboard.html（控制台）。")

    _add_heading(doc, "19.2 登录注册功能", level=2)
    _add_para(doc, "系统提供完整的用户认证功能。登录页面（/login）支持用户名密码验证，使用SHA-256哈希存储密码；注册页面（/register）支持新用户注册，包含密码一致性检查和用户名唯一性检查；退出登录（/logout）清除会话状态；未登录用户自动重定向至登录页面。所有密码均经过哈希处理后存储，不保存明文密码。")

    _add_heading(doc, "19.3 主页与控制台", level=2)
    _add_para(doc, "登录后进入系统主页（/home），页面包含以下区域：（1）导航栏——显示系统Logo、导航链接（主页、控制台、退出）和当前登录用户；（2）Hero区域——系统名称、功能描述和输入输出格式说明；（3）功能卡片——数据导入、AI标注引擎、校验质控三大核心功能简介；（4）快捷操作区——新建标注任务、查看任务列表、导出数据、系统诊断四个快捷按钮；（5）API接口列表——展示所有可用API端点。")

    _add_para(doc, "控制台页面（/dashboard）提供系统运行状态总览：（1）统计卡片——标注数据总量、管线运行次数、今日新增标注；（2）最近管线运行——展示最近运行的任务ID、状态和时间；（3）系统功能模块——列出所有子系统模块及其运行状态。")

    if diagrams and "web" in diagrams:
        _add_image(doc, diagrams["web"], width_cm=14)

    _add_heading(doc, "19.4 API接口设计", level=2)
    _add_para(doc, "系统提供三级RESTful API接口体系：（1）一级接口（管线管理）——POST /api/v1/pipeline/run启动标注管线、GET /api/v1/pipeline/status/:id查询任务状态、GET /api/v1/pipeline/list查看任务列表、GET /api/v1/pipeline/stats查看管线统计；（2）二级接口（标注数据管理）——POST /api/v1/annotations/submit提交标注数据、GET /api/v1/annotations/list获取标注列表、GET /api/v1/annotations/:id查询单条标注、GET /api/v1/annotations/export/:fmt导出标注数据；（3）系统接口（运维管理）——GET /api/v1/system/info系统信息、GET /api/v1/system/health健康检查。")

    _add_para(doc, "所有业务接口均需要用户登录认证（基于Flask session），系统接口中的health检查免认证。API返回格式统一为JSON，包含必要的状态码和错误信息。列表接口支持limit参数控制返回条数。")

    if diagrams and "api" in diagrams:
        _add_image(doc, diagrams["api"], width_cm=14)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 第二十章 文件清单
    # ════════════════════════════════════════════
    _add_heading(doc, "二十、文件清单", level=1)

    _add_para(doc, "本软件共包含130个Python源文件，总计约12,665行代码（含注释和空行）。按文件夹分类如下：")

    _add_heading(doc, "20.1 core/ — 核心管线（59个文件，6,198行）", level=2)

    core_files = [
        ("core/engine/", "9个", "模板加载、渲染、管理、上下文组装、请求构建、推理调用、重试策略、响应收集", "862"),
        ("core/input/", "7个", "多格式读取、编码检测、文本归一、分块、批量读取、文件校验", "615"),
        ("core/preprocessing/", "7个", "行级清洗、标点归一、元数据猜测、结构检测、质量预过滤、管线编排", "533"),
        ("core/parsing/", "9个", "JSON提取、清洗、校验、字段映射、嵌套构建、错误恢复、批量解析", "971"),
        ("core/validation/", "9个", "Schema定义、类型检查、必填检查、范围校验、交叉验证、完整性报告", "1,472"),
        ("core/quality/", "9个", "一致性检查、模式检测、置信度评分、错误分类、严重评估、修正建议", "901"),
        ("core/output/", "9个", "JSON/CSV/XML序列化、格式注册、批量写入、合并、校验、归档", "844"),
    ]
    _add_table(doc, ["目录", "文件数", "功能描述", "行数"], core_files)

    _add_heading(doc, "20.2 cli/ — CLI接口（6个文件，1,046行）", level=2)
    _add_table(doc, ["文件", "功能描述", "行数"], [
        ("cli/cli_parser.py", "命令解析与注册", "~130"),
        ("cli/cli_commands.py", "七个子命令实现", "~175"),
        ("cli/cli_formatter.py", "格式化输出与着色", "~190"),
        ("cli/demo_runner.py", "演示运行器", "~155"),
        ("cli/__init__.py", "包初始化", "~24"),
        ("demo.py", "命令行入口", "~110"),
    ])

    _add_heading(doc, "20.3 web/ — Web前端（5个文件，~350行）", level=2)
    web_files = [
        ("web/app.py", "Flask应用主逻辑（登录/注册/主页/API）", "~180"),
        ("web/templates/login.html", "登录页面模板", "~60"),
        ("web/templates/register.html", "注册页面模板", "~55"),
        ("web/templates/home.html", "主页模板（含API列表展示）", "~120"),
        ("web/templates/dashboard.html", "控制台模板（统计/状态）", "~100"),
    ]
    _add_table(doc, ["文件", "功能描述", "行数"], web_files)

    _add_heading(doc, "20.4 lib/ — 支撑库（65个文件，5,421行）", level=2)
    lib_files = [
        ("lib/cache/", "8个", "缓存条目、内存缓存、磁盘缓存、策略、管理器、结果缓存、模板缓存", "445"),
        ("lib/config/", "8个", "配置源、加载、校验、监视、profile管理、序列化、报告", "574"),
        ("lib/cost/", "8个", "Token计数、成本估算、性能计时、吞吐计量、用量日志、预算控制", "780"),
        ("lib/errors/", "6个", "分层异常体系(22类)、错误处理、重试策略、熔断器、降级处理", "455"),
        ("lib/log/", "7个", "日志配置、写入器、结构化日志、审计日志、诊断收集、日志清理", "484"),
        ("lib/pipeline/", "10个", "阶段定义、管线构建、执行、任务定义、队列、断点、恢复", "1,345"),
        ("lib/reporting/", "9个", "运行/质量/成本/对比报告、文本/JSON/HTML格式、归档", "703"),
        ("lib/utils/", "9个", "字符串、文件、字典、列表、时间、校验、序列化、系统工具", "635"),
    ]
    _add_table(doc, ["目录", "文件数", "功能描述", "行数"], lib_files)

    _add_para(doc, f"合计：135个文件，约13,015行代码。", bold=True)

    # 保存
    fp = os.path.join(OUTPUT_DIR, "软著_数据生产管线_软件说明书.docx")
    doc.save(fp)
    print(f"说明书已保存: {fp}")
    return fp


# ═══════════════════════════════════════════════
#  源代码文档
# ═══════════════════════════════════════════════

def generate_code_docx():
    """生成源代码docx文件（前2000行+后2000行）"""
    all_lines, files = _build_concatenated_source()
    total = len(all_lines)

    first_2000 = all_lines[:2000]
    last_2000 = all_lines[-2000:] if total > 2000 else all_lines

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(8)

    # 封面
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("古典诗歌文本结构化标注生产系统")
    _set_font(run, "黑体", 18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("源  代  码")
    _set_font(run, "黑体", 16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(f"（共计{total}行，{len(files)}个文件，此处展示前2000行与后2000行）"), "宋体", 11)

    doc.add_page_break()

    # --- 文件清单 ---
    _add_heading(doc, "文件清单", level=1)
    for rel_path, lines in files:
        _add_para(doc, f"• {rel_path}（{len(lines)}行）", size=9, indent=False)

    doc.add_page_break()

    # --- 前2000行 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第一部分：前2000行")
    _set_font(run, "黑体", 14, bold=True)
    doc.add_paragraph()

    chunk_size = 50
    for start in range(0, len(first_2000), chunk_size):
        chunk = first_2000[start:start + chunk_size]
        code_text = "".join(chunk)
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        _set_font(run, "Courier New", 6.5)
        p.paragraph_format.line_spacing = Pt(9)

    doc.add_page_break()

    # --- 后2000行 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第二部分：后2000行")
    _set_font(run, "黑体", 14, bold=True)
    doc.add_paragraph()

    for start in range(0, len(last_2000), chunk_size):
        chunk = last_2000[start:start + chunk_size]
        code_text = "".join(chunk)
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        _set_font(run, "Courier New", 6.5)
        p.paragraph_format.line_spacing = Pt(9)

    fp = os.path.join(OUTPUT_DIR, "软著_数据生产管线_源代码.docx")
    doc.save(fp)
    print(f"源代码文档已保存: {fp} (全文{total}行，前2000行+后2000行)")
    return fp


if __name__ == "__main__":
    if _HAS_PIL:
        print("正在生成架构图、流程图、界面截图...")
        diagrams = generate_all_diagrams()
    else:
        print("Pillow 不可用，跳过图示生成")
        diagrams = {}

    print("正在生成软件说明书...")
    generate_manual(diagrams=diagrams)

    print("正在生成源代码文档...")
    generate_code_docx()

    print("\n全部文档生成完成！")
