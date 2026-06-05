#!/usr/bin/env python3
"""Rebuild pipeline doc using python-docx API - fixed image replacement."""
import os, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn
from PIL import Image

SRC = r'C:\Users\Administrator\Desktop\新建文件夹 (14)\软著_数据生产管线_软件说明书.docx'
DST = r'C:\Users\Administrator\Desktop\pipeline_doc.docx'
IMG = r'C:\Users\Administrator\Desktop\All Mix'

shutil.copy2(SRC, DST)
doc = Document(DST)

# ─── Find sections ───
paras = list(enumerate(p.text.strip() for p in doc.paragraphs))
body_s18 = next(i for i, t in paras if t == '十八、CLI接口层' and i > 100)
body_s19 = next(i for i, t in paras if t == '十九、Web前端模块' and i > 100)
print(f'Section 18 at {body_s18}, Section 19 at {body_s19}')

# ─── Remove section 18 paragraphs (REVERSE order) ───
for i in range(body_s19 - 1, body_s18 - 1, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
removed = body_s19 - body_s18
print(f'Removed section 18 ({removed} paragraphs)')

# ─── Build XML helper functions ───
NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"'

def text_para(text, bold=False, sz=21, font='宋体', style=None, color=None, align=None, indent=None):
    sp = ' xml:space="preserve"' if text and (text[0] == ' ' or text[-1] == ' ') else ''
    xml = f'<w:p {NS}>\n  <w:pPr>\n'
    if style:
        xml += f'    <w:pStyle w:val="{style}"/>\n'
    if align:
        xml += f'    <w:jc w:val="{align}"/>\n'
    if indent:
        xml += f'    <w:ind w:firstLine="{int(indent * 567)}"/>\n'
    xml += f'    <w:rPr>\n      <w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}"/>\n'
    if bold:
        xml += '      <w:b/>\n'
    if color:
        xml += f'      <w:color w:val="{color}"/>\n'
    xml += f'      <w:sz w:val="{sz}"/>\n      <w:szCs w:val="{sz}"/>\n    </w:rPr>\n  </w:pPr>\n'
    xml += f'  <w:r>\n    <w:rPr>\n      <w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}"/>\n'
    if bold:
        xml += '      <w:b/>\n'
    if color:
        xml += f'      <w:color w:val="{color}"/>\n'
    xml += f'      <w:sz w:val="{sz}"/>\n      <w:szCs w:val="{sz}"/>\n    </w:rPr>\n'
    xml += f'    <w:t{sp}>{text}</w:t>\n  </w:r>\n</w:p>'
    return xml

def img_para(img_path, rid, docpr_id, cx_inches=5.2):
    with Image.open(img_path) as img:
        w, h = img.size
    cx = int(cx_inches * 914400)
    cy = int(cx * h / w)
    fname = os.path.basename(img_path)
    return f'''<w:p {NS}>
  <w:pPr><w:jc w:val="center"/></w:pPr>
  <w:r>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="{docpr_id}" name="{fname}"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
        <a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:pic>
            <pic:nvPicPr><pic:cNvPr id="0" name="{fname}"/><pic:cNvPicPr/></pic:nvPicPr>
            <pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>
            <pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>
          </pic:pic>
        </a:graphicData></a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>'''

def page_break():
    return f'<w:p {NS}><w:r><w:br w:type="page"/></w:r></w:p>'

# ─── Find section 19 heading element ───
s19_elem = None
for p in doc.paragraphs:
    if p.text.strip() == '十九、Web前端模块':
        s19_elem = p._element
        break

if s19_elem is None:
    print('ERROR: Could not find section 19 heading!')
    exit(1)

# ─── Build new section 18 content (in display order) ───
content = []

content.append(text_para('十八、系统运行截图', bold=True, sz=28, font='黑体', style='Heading1'))

content.append(text_para('18.1 管线CLI命令帮助', bold=True, sz=24, font='黑体', style='Heading2'))
content.append(text_para('系统提供完整的CLI命令接口，支持run（执行标注管线）、validate（校验数据）、report（生成报告）、clean（清理缓存）、config（管理配置）、stats（运行统计）七个子命令。通过demo.py --help查看所有命令及参数说明。', indent=0.75))
content.append(text_para('图1： 管线CLI命令帮助', bold=True, sz=24, font='宋体', color='2C3E50', align='center'))
content.append(img_para(f'{IMG}/pipeline_cli_help.png', 'rId14', 14))

content.append(text_para('18.2 管线运行过程', bold=True, sz=24, font='黑体', style='Heading2'))
content.append(text_para('执行run命令启动标注管线，系统按"输入读取→文本预处理→AI标注引擎→结果解析与校验→数据输出"五步流程自动处理。输入支持TXT/JSON/CSV/MD格式，AI引擎完成标注后经过多层校验确保质量。', indent=0.75))
content.append(text_para('图2： 管线运行过程', bold=True, sz=24, font='宋体', color='2C3E50', align='center'))
content.append(img_para(f'{IMG}/pipeline_run.png', 'rId15', 15))

content.append(text_para('18.3 标注数据输出', bold=True, sz=24, font='黑体', style='Heading2'))
content.append(text_para('管线输出结构化JSON格式的标注数据，包含运行元信息、诗歌原文、多维标注结果（意象分析/情感分析/结构分析）及质量评分。支持JSON/CSV/XML三种导出格式。', indent=0.75))
content.append(text_para('图3： 标注数据输出', bold=True, sz=24, font='宋体', color='2C3E50', align='center'))
content.append(img_para(f'{IMG}/pipeline_output.png', 'rId16', 16))

content.append(page_break())

# Insert before section 19 heading (reverse order so first = first in doc)
parent = s19_elem.getparent()
idx = list(parent).index(s19_elem)
for xml_str in reversed(content):
    elem = etree.fromstring(xml_str)
    parent.insert(idx, elem)

print(f'Inserted new section 18 ({len(content)} elements)')

# ─── Replace old images in section 19 ───
# Find old image elements by rId in the body
body_elem = doc.element.body

for old_rid, replacements in [
    ('rId12', [  # in 19.3 - replace with login + register
        ('label', '图4： 登录页面'),
        ('img', f'{IMG}/p_login.png', 'rId17', 17),
        ('label', '图5： 注册页面'),
        ('img', f'{IMG}/p_register.png', 'rId18', 18),
    ]),
    ('rId13', [  # in 19.4 - replace with dashboard + home + api health
        ('label', '图6： 控制台主页'),
        ('img', f'{IMG}/p_dashboard.png', 'rId19', 19),
        ('img', f'{IMG}/p_home.png', 'rId20', 20),
        ('label', '图7： API接口响应'),
        ('img', f'{IMG}/p_api_health.png', 'rId21', 21),
    ]),
]:
    # Find the blip element that references old_rid
    xpath = f'.//{{http://schemas.openxmlformats.org/drawingml/2006/main}}blip[@{{http://schemas.openxmlformats.org/officeDocument/2006/relationships}}embed="{old_rid}"]'
    blips = body_elem.findall(xpath)
    if not blips:
        print(f'WARNING: No blip with {old_rid} found')
        continue

    for blip in blips:
        # Find parent w:p
        p_elem = blip
        while p_elem is not None and p_elem.tag != f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}p':
            p_elem = p_elem.getparent()
        if p_elem is None:
            continue

        parent_of_p = p_elem.getparent()
        p_idx = list(parent_of_p).index(p_elem)

        # Build replacement XML
        new_xmls = []
        for item in replacements:
            if item[0] == 'label':
                new_xmls.append(text_para(item[1], bold=True, sz=24, font='宋体', color='2C3E50', align='center'))
            elif item[0] == 'img':
                new_xmls.append(img_para(item[1], item[2], item[3]))

        # Insert in reverse order so first in list = first in doc
        for xml_str in reversed(new_xmls):
            new_elem = etree.fromstring(xml_str)
            parent_of_p.insert(p_idx, new_elem)

        # Remove old paragraph
        parent_of_p.remove(p_elem)
        print(f'Replaced {old_rid} with {len(replacements)} elements')

# ─── Save ───
doc.save(DST)
print(f'Saved to {DST}')

# ─── Verify ───
doc2 = Document(DST)
figs = []
for p in doc2.paragraphs:
    t = p.text.strip()
    if '图' in t and '：' in t:
        figs.append(t)
    if '十八、系统运行截图' in t:
        print(f'  Found: {t}')
print(f'Figure labels ({len(figs)}): {figs}')
