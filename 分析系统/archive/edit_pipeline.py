#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Rebuild pipeline doc with python-docx - clean approach."""
import os, zipfile
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn

SRC = 'C:/Users/Administrator/Desktop/新建文件夹 (14)/软著_数据生产管线_软件说明书.docx'
DST = 'C:/Users/Administrator/Desktop/新建文件夹 (14)/软著_古典诗歌文本结构化标注生产系统_软件说明书.docx'
IMG = 'C:/Users/Administrator/Desktop/All Mix'

doc = Document(SRC)

# ─── Find sections ───
paras = list(enumerate(p.text.strip() for p in doc.paragraphs))
body_s18 = next(i for i, t in paras if t == '十八、CLI接口层' and i > 100)
body_s19 = next(i for i, t in paras if t == '十九、Web前端模块' and i > 100)
body_s20 = next(i for i, t in paras if t == '二十、文件清单' and i > 100)
print(f'BODY: 18={body_s18} 19={body_s19} 20={body_s20}')

# Also find figure indices within section 19
# Look for "图6" etc in body area (indices > body_s19)
fig6 = None; fig7 = None; fig8 = None
for i, t in paras:
    if i < body_s19: continue
    if '图6' in t and fig6 is None: fig6 = i
    if '图7' in t and fig7 is None: fig7 = i
    if '图8' in t and fig8 is None: fig8 = i
    if i > body_s20: break
print(f'Figures: 6={fig6} 7={fig7} 8={fig8}')

# ─── Remove section 18 body paragraphs ───
removed = 0
for i in range(body_s19 - 1, body_s18 - 1, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    removed += 1
print(f'Removed {removed} paragraphs (section 18)')

# ─── Find new position of section 19 ───
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '十九、Web前端模块' and i > 100:
        ins_pos = i
        break

def ins_t(text, pos, bold=False, sz=12, style=None, indent=None, align=None):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    if style: p.style = style
    if align: p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(sz)
    if indent: p.paragraph_format.first_line_indent = Cm(indent)
    return p

def ins_img(path, pos, w=5.5):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(w))
    return p

# ─── Insert new section 18: system screenshots ───
ins_t('十八、系统运行截图', ins_pos, bold=True, sz=16, style=doc.styles['Heading 1'])

# 18.1
ins_t('18.1 管线CLI命令帮助', ins_pos, bold=True, sz=14, style=doc.styles['Heading 2'])
ins_t('系统提供完整的CLI命令接口，支持run（执行标注管线）、validate（校验数据）、report（生成报告）、clean（清理缓存）、config（管理配置）、stats（运行统计）七个子命令。通过demo.py --help查看所有命令及参数说明。', ins_pos, indent=0.75)
ins_t('图1： 管线CLI命令帮助', ins_pos, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
ins_img(f'{IMG}/pipeline_cli_help.png', ins_pos)

# 18.2
ins_t('18.2 管线运行过程', ins_pos, bold=True, sz=14, style=doc.styles['Heading 2'])
ins_t('执行run命令启动标注管线，系统按"输入读取→文本预处理→AI标注引擎→结果解析与校验→数据输出"五步流程自动处理。输入支持TXT/JSON/CSV/MD格式，AI引擎完成标注后经过多层校验确保质量。', ins_pos, indent=0.75)
ins_t('图2： 管线运行过程', ins_pos, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
ins_img(f'{IMG}/pipeline_run.png', ins_pos)

# 18.3
ins_t('18.3 标注数据输出', ins_pos, bold=True, sz=14, style=doc.styles['Heading 2'])
ins_t('管线输出结构化JSON格式的标注数据，包含运行元信息、诗歌原文、多维标注结果（意象分析/情感分析/结构分析）及质量评分。支持JSON/CSV/XML三种导出格式。', ins_pos, indent=0.75)
ins_t('图3： 标注数据输出', ins_pos, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
ins_img(f'{IMG}/pipeline_output.png', ins_pos)

# Page break before section 19
pb = doc.paragraphs[ins_pos].insert_paragraph_before('')
br_elem = etree.SubElement(pb._element.find(qn('w:r')) or etree.SubElement(pb._element, qn('w:r')), qn('w:br'))
br_elem.set(qn('w:type'), 'page')

print(f'Inserted new section 18 content before paragraph {ins_pos}')

# ─── Save ───
doc.save(DST)
print(f'Saved to {DST}')

# ─── Update: Replace old images in the ZIP ───
# Old images from original doc: image3 (CLI section), image4 (old web), image5 (api diagram)
# New images added by python-docx: image6, image7, image8 (our new pipeline screenshots)
# We need to ALSO add the web screenshots (login, register, etc.)
# Since python-docx already handles images for us, the web section still has old figure images

# Approach: Unpack and swap image files directly in the zip
tmp = DST + '.tmp'
os.rename(DST, tmp)

# Map: find which internal names python-docx gave the old images
with zipfile.ZipFile(tmp, 'r') as z:
    media = [n for n in z.namelist() if 'media' in n]
    print('Media files:', media)

# We need to check which of the new files (pipeline_cli_help etc.) used which image numbers
# and also replace the old section 19 images
# Let me find them by looking in document.xml

with zipfile.ZipFile(tmp, 'r') as z:
    doc_xml = z.read('word/document.xml').decode('utf-8')
    # Find image references
    import re
    for m in re.finditer(r'r:embed="([^"]+)"[^>]*/>', doc_xml):
        rid = m.group(1)
        # Find the Target for this rid
        rels_xml = z.read('word/_rels/document.xml.rels').decode('utf-8')
        tmatch = re.search(f'Id="{rid}"[^>]*Target="([^"]+)"', rels_xml)
        if tmatch:
            print(f'  {rid} -> {tmatch.group(1)}')

z.close()
