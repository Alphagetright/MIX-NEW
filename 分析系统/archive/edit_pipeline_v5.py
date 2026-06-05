#!/usr/bin/env python3
"""Rebuild pipeline doc: all screenshots in section 18, login/register first."""
import os, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn
from PIL import Image

SRC = r'C:\Users\Administrator\Desktop\新建文件夹 (14)\软著_数据生产管线_软件说明书.docx'
DST = r'C:\Users\Administrator\Desktop\pipeline_doc.docx'
IMG = r'C:\Users\Administrator\Desktop\All Mix'

shutil.copy2(SRC, DST)

# Use python-docx API for clean relationship management
doc = Document(DST)

# ─── Find sections ───
paras = list(enumerate(p.text.strip() for p in doc.paragraphs))
body_s18 = next(i for i, t in paras if t == '十八、CLI接口层' and i > 100)
body_s19 = next(i for i, t in paras if t == '十九、Web前端模块' and i > 100)
print(f'Section 18 at {body_s18}, Section 19 at {body_s19}')

# ─── Remove OLD section 18 paragraphs ───
for i in range(body_s19 - 1, body_s18 - 1, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
print('Removed old section 18')

# ─── Remove old image paragraphs from section 19 ───
# Find image paragraphs in section 19 that reference old images
s19_para_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '十九、Web前端模块':
        s19_para_idx = i
        break

old_img_paras = []
for i, p in enumerate(doc.paragraphs):
    if i < s19_para_idx: continue
    if i > s19_para_idx + 100: break
    for b in p._element.findall('.//' + qn('a:blip')):
        rid = b.get(qn('r:embed'))
        if rid in ('rId12', 'rId13'):
            old_img_paras.append(i)

# Remove in reverse order
for i in reversed(old_img_paras):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
print(f'Removed {len(old_img_paras)} old image paragraphs from section 19')

# Also remove figure labels that were next to those images
# Look for "图6"/"图7"/"图8" text paragraphs near those areas
# (these were original figure labels in the source doc)

# ─── Find new section 19 insertion point ───
new_s19_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '十九、Web前端模块':
        new_s19_idx = i
        break
print(f'Section 19 now at paragraph {new_s19_idx}')

# ─── Helper functions using python-docx API (auto-manages rels & media) ───
def insert_heading1(text, pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.style = doc.styles['Heading 1']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    return p

def insert_heading2(text, pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return p

def insert_text(text, pos, indent_cm=0.75):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    return p

def insert_figure_label(text, pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p

def insert_image(img_path, pos, width_inches=5.2):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    return p

def insert_page_break(pos):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    run = p.add_run()
    br = etree.SubElement(run._element, qn('w:br'))
    br.set(qn('w:type'), 'page')

# ─── Screenshot data: ordered list ───
# (heading2_title, description, figure_label, img_filename)
screenshots = [
    ('18.1 用户登录',
     '系统登录页面（/login）提供用户认证入口，支持用户名密码验证。输入正确的用户名和密码后点击"登录"按钮，系统验证通过后跳转至系统主页。登录状态基于Flask session管理，未登录用户自动重定向至登录页面。密码使用SHA-256哈希存储，不保存明文密码。',
     '图1： 用户登录页面', 'p_login.png'),
    ('18.2 用户注册',
     '注册页面（/register）支持新用户自助注册，需填写用户名、密码和确认密码。系统进行两项校验：（1）密码一致性检查——两次输入的密码必须一致；（2）用户名唯一性检查——用户名不可与已有用户重复。注册成功后自动跳转至登录页面。',
     '图2： 用户注册页面', 'p_register.png'),
    ('18.3 控制台主页',
     '控制台页面（/dashboard）提供系统运行状态总览，包含三个区域：（1）统计卡片——标注数据总量、管线运行次数、今日新增标注；（2）最近管线运行——展示最近运行的任务ID、状态和时间；（3）系统功能模块——列出所有子系统模块及运行状态。',
     '图3： 控制台主页', 'p_dashboard.png'),
    ('18.4 系统主页',
     '登录后进入系统主页（/home），页面包含导航栏（系统Logo、导航链接和当前用户）、Hero区域（系统名称和功能描述）、功能卡片（数据导入、AI标注引擎、校验质控三大核心功能）、快捷操作区（新建标注任务、查看任务列表、导出数据、系统诊断）及API接口列表。',
     '图4： 系统主页', 'p_home.png'),
    ('18.5 API接口响应',
     '系统提供RESTful API接口，通过GET /api/v1/system/health可验证服务运行状态，返回JSON格式的系统健康信息，包括状态、运行时间、版本号等。所有业务接口均需用户登录认证（基于Flask session），API返回格式统一为JSON，包含状态码和错误信息。',
     '图5： API接口响应', 'p_api_health.png'),
    ('18.6 管线CLI命令帮助',
     '系统提供完整的CLI命令接口，支持run（执行标注管线）、validate（校验数据）、report（生成报告）、clean（清理缓存）、config（管理配置）、stats（运行统计）七个子命令。通过python demo.py --help可查看所有命令及参数说明。',
     '图6： 管线CLI命令帮助', 'pipeline_cli_help.png'),
    ('18.7 管线运行过程',
     '执行run命令启动标注管线，系统按"输入读取→文本预处理→AI标注引擎→结果解析与校验→数据输出"五步流程自动处理。输入支持TXT/JSON/CSV/MD格式，AI引擎完成结构化标注后经过多层校验确保数据质量，最终输出标准格式的标注结果。',
     '图7： 管线运行过程', 'pipeline_run.png'),
    ('18.8 标注数据输出',
     '管线输出结构化JSON格式的标注数据，包含运行元信息（ID、状态、耗时）、诗歌原文（标题、作者、诗句）、多维标注结果（意象分析、情感分析、结构分析）及质量评分（完整度、置信度）。支持JSON/CSV/XML三种导出格式。',
     '图8： 标注数据输出', 'pipeline_output.png'),
]

# ─── Insert new content in REVERSE order ───
# (last inserted = first displayed)
pos = new_s19_idx

# Page break before section 19
insert_page_break(pos)

# Insert screenshots in reverse order
for h2_title, desc, fig_label, img_file in reversed(screenshots):
    img_path = os.path.join(IMG, img_file)
    insert_image(img_path, pos)
    insert_figure_label(fig_label, pos)
    insert_text(desc, pos)
    insert_heading2(h2_title, pos)

# Section heading (inserted first but appears last due to reverse)
insert_heading1('十八、系统运行截图', pos)

print(f'Inserted {len(screenshots)} screenshots in section 18')

# ─── Save ───
doc.save(DST)
print(f'Saved to {DST}')

# ─── Verify ───
doc2 = Document(DST)
figs = []
for p in doc2.paragraphs:
    t = p.text.strip()
    if '图' in t and '：' in t:
        figs.append(t)
    if '十八、系统运行截图' in t:
        print(f'  Found section heading')
print(f'Figure labels ({len(figs)}): {figs}')
