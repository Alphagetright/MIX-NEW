#!/usr/bin/env python3
"""Rebuild pipeline doc using python-docx API cleanly."""
import os, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn

SRC = r'C:\Users\Administrator\Desktop\新建文件夹 (14)\软著_数据生产管线_软件说明书.docx'
DST = r'C:\Users\Administrator\Desktop\pipeline_doc.docx'
IMG = r'C:\Users\Administrator\Desktop\All Mix'

# Copy original first
shutil.copy2(SRC, DST)
doc = Document(DST)

# ─── Find sections ───
paras = list(enumerate(p.text.strip() for p in doc.paragraphs))
body_s18 = next(i for i, t in paras if t == '十八、CLI接口层' and i > 100)
body_s19 = next(i for i, t in paras if t == '十九、Web前端模块' and i > 100)
print(f'Section 18 at index {body_s18}, Section 19 at index {body_s19}')

# Also find image paragraphs in section 19
fig_imgs = []
for i, p in enumerate(doc.paragraphs):
    if i < body_s19: continue
    # Check if paragraph contains a drawing (image)
    drawings = p._element.findall('.//'+qn('w:drawing'))
    if drawings:
        blips = p._element.findall('.//'+qn('a:blip'))
        for b in blips:
            rid = b.get(qn('r:embed'))
            fig_imgs.append((i, rid))
    if i > body_s19 + 100: break  # don't go past section 19

print(f'Found {len(fig_imgs)} images in section 19: {fig_imgs}')

# ─── Remove section 18 paragraphs ───
removed = 0
for i in range(body_s19 - 1, body_s18 - 1, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    removed += 1
print(f'Removed {removed} paragraphs (section 18)')

# ─── Find insertion point (new position of section 19 heading) ───
ins_pos = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '十九、Web前端模块' and i > 100:
        ins_pos = i
        break

print(f'Insertion point at paragraph {ins_pos}')

def ins_t(text, pos, bold=False, sz=12, style=None, indent=None, align=None):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    if style: p.style = style
    if align: p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(sz)
    if indent: p.paragraph_format.first_line_indent = Cm(indent)
    return p

def ins_img(path, pos, w=5.5):
    p = doc.paragraphs[pos].insert_paragraph_before('')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(w))
    return p

def add_page_break(pos):
    """Add a page break before a paragraph."""
    p = doc.paragraphs[pos].insert_paragraph_before('')
    r = p.add_run()
    br = etree.SubElement(r._element, qn('w:br'))
    br.set(qn('w:type'), 'page')

# ─── Insert new section 18: system screenshots ───
ins_t('十八、系统运行截图', ins_pos, bold=True, sz=16, style=doc.styles['Heading 1'])

# 18.1
ins_t('18.1 管线CLI命令帮助', ins_pos, bold=True, sz=14, style=doc.styles['Heading 2'])
ins_t('系统提供完整的CLI命令接口，支持run（执行标注管线）、validate（校验数据）、report（生成报告）、clean（清理缓存）、config（管理配置）、stats（运行统计）七个子命令。通过demo.py --help查看所有命令及参数说明。', ins_pos, indent=0.75)
ins_t('图1： 管线CLI命令帮助', ins_pos, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
ins_img(f'{IMG}/pipeline_cli_help.png', ins_pos)

# 18.2
ins_t('18.2 管线运行过程', ins_pos, bold=True, sz=14, style=doc.styles['Heading 2'])
ins_t('执行run命令启动标注管线，系统按"输入读取→文本预处理→AI标注引擎→结果解析与校验→数据输出"五步流程自动处理。输入支持TXT/JSON/CSV/MD格式，AI引擎完成标注后经过多层校验确保质量。', ins_pos, indent=0.75)
ins_t('图2： 管线运行过程', ins_pos, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
ins_img(f'{IMG}/pipeline_run.png', ins_pos)

# 18.3
ins_t('18.3 标注数据输出', ins_pos, bold=True, sz=14, style=doc.styles['Heading 2'])
ins_t('管线输出结构化JSON格式的标注数据，包含运行元信息、诗歌原文、多维标注结果（意象分析/情感分析/结构分析）及质量评分。支持JSON/CSV/XML三种导出格式。', ins_pos, indent=0.75)
ins_t('图3： 标注数据输出', ins_pos, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
ins_img(f'{IMG}/pipeline_output.png', ins_pos)

# Page break before section 19
add_page_break(ins_pos)

print('Inserted new section 18')

# ─── Now handle section 19 images ───
# Find the old images in section 19 by re-scanning
section19_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '十九、Web前端模块' and i > 100:
        section19_start = i
        break

print(f'Section 19 starts at paragraph {section19_start}')

# Find the image paragraphs in section 19 body (between 19.2/19.3/19.4)
# Strategy: find image paragraphs by checking for blip elements
img_paras_19 = []
for i, p in enumerate(doc.paragraphs):
    if i < section19_start: continue
    if i > section19_start + 100: break
    blips = p._element.findall('.//'+qn('a:blip'))
    if blips:
        for b in blips:
            rid = b.get(qn('r:embed'))
            # Only target old image rIds (rId12=rId13 from original)
            if rid in ('rId12', 'rId13'):
                img_paras_19.append((i, rid))

print(f'Old images to replace in section 19: {img_paras_19}')

# For each old image paragraph, insert new images before it, then remove the old one
if len(img_paras_19) >= 1:
    idx, rid = img_paras_19[0]  # First old image (in 19.3 主页与控制台)
    # Insert before the old image paragraph
    ins_t('图4： 登录页面', idx, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    ins_img(f'{IMG}/p_login.png', idx)
    ins_t('图5： 注册页面', idx, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    ins_img(f'{IMG}/p_register.png', idx)
    # Remove the old paragraph (its index shifted by 4 insertions)
    old_p = doc.paragraphs[idx + 4]._element
    old_p.getparent().remove(old_p)
    print('Replaced old image 1 with login + register screenshots')

if len(img_paras_19) >= 2:
    # After removing first old image, the second image paragraph may have shifted
    # Re-scan to find it
    img_paras_19_v2 = []
    for i, p in enumerate(doc.paragraphs):
        if i < section19_start: continue
        if i > section19_start + 100: break
        blips = p._element.findall('.//'+qn('a:blip'))
        if blips:
            for b in blips:
                rid = b.get(qn('r:embed'))
                if rid in ('rId12', 'rId13'):
                    img_paras_19_v2.append((i, rid))

    if img_paras_19_v2:
        idx, rid = img_paras_19_v2[0]
        ins_t('图6： 控制台主页', idx, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        ins_img(f'{IMG}/p_dashboard.png', idx)
        ins_img(f'{IMG}/p_home.png', idx)
        ins_t('图7： API接口响应', idx, bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        ins_img(f'{IMG}/p_api_health.png', idx)
        # Remove old paragraph (shifted by 4 insertions)
        old_p = doc.paragraphs[idx + 4]._element
        old_p.getparent().remove(old_p)
        print('Replaced old image 2 with dashboard + home + api health screenshots')
    else:
        print('WARNING: Second old image not found after first replacement')

# Also clean up TOC if possible - remove stale entry from TOC field
# This would require manipulating field codes which python-docx can't easily do
# Word will auto-update TOC on open

# ─── Save ───
doc.save(DST)
print(f'Saved to {DST}')
