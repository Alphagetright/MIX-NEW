# -*- coding: utf-8 -*-
"""生成软著源代码申报 .docx 文件 — 前2000行 + 后2000行"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

BASE = r"C:\Users\Administrator\Desktop\All Mix"
OUT = os.path.join(BASE, "软著_源代码.docx")

# ─── 核心源文件清单（仅含应用代码，排除构建脚本和旧页面） ───

PYTHON_FILES = [
    "config.py",
    "models.py",
    "users.py",
    "cache.py",
    "logger.py",
    "errors.py",
    "validators.py",
    "middleware.py",
    "monitor.py",
    "utils.py",
    "preprocessor.py",
    "export_service.py",
    "build_rag.py",
    "query_rag.py",
    "analytics.py",
    "admin.py",
    "cli.py",
    "app.py",
]

TEMPLATE_FILES = [
    "templates/base.html",
    "templates/login.html",
    "templates/register.html",
    "templates/dashboard.html",
    "templates/graph_charts.html",
    "templates/graph_table.html",
    "templates/ai.html",
    "templates/recycle.html",
    "templates/admin.html",
]

ALL_FILES = PYTHON_FILES + TEMPLATE_FILES


def read_file_lines(path):
    """读取文件行，返回 list[str]"""
    full = os.path.join(BASE, path)
    if not os.path.exists(full):
        return []
    with open(full, "r", encoding="utf-8", errors="ignore") as f:
        return f.readlines()


def build_all_lines():
    """拼接所有源文件为一个带标记的行列表"""
    all_lines = []
    for fname in ALL_FILES:
        lines = read_file_lines(fname)
        if not lines:
            continue
        # 文件头标记
        all_lines.append(f"# {'='*70}\n")
        all_lines.append(f"# 文件: {fname}  行数: {len(lines)}\n")
        all_lines.append(f"# {'='*70}\n")
        all_lines.append("\n")
        all_lines.extend(lines)
        all_lines.append("\n")
    return all_lines


def main():
    all_lines = build_all_lines()
    total = len(all_lines)
    print(f"总行数（含分隔标记）: {total}")

    front = all_lines[:2000]
    back = all_lines[-2000:]

    print(f"前2000行: {len(front)}")
    print(f"后2000行: {len(back)}")

    doc = Document()

    # ── 页面设置 ──
    for section in doc.sections:
        section.page_width = Cm(21)       # A4
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ── 样式设置 ──
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(8)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.1

    # ── 封面页 ──
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("唐诗意象智能分析系统")
    r.font.size = Pt(26)
    r.bold = True
    r.font.name = "黑体"
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("计算机软件著作权登记 — 源程序")
    r.font.size = Pt(16)
    r.font.name = "黑体"
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        "提交源码行数：前2000行 + 后2000行",
        "总源码行数：约 5,200 行（不含测试和构建脚本）",
        "开发语言：Python 3.10+ / HTML / JavaScript",
        "开发完成日期：2026年5月",
    ]
    for text in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(12)

    doc.add_page_break()

    # ── 前2000行 ──
    doc.add_heading("第一部分：源代码前2000行", level=1)

    for i, line in enumerate(front, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1

        # 行号
        r_num = p.add_run(f"{i:5d} ")
        r_num.font.name = "Courier New"
        r_num.font.size = Pt(7)
        r_num.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 代码内容（替换制表符为4空格）
        display = line.rstrip("\n").replace("\t", "    ")
        r_code = p.add_run(display if display else " ")
        r_code.font.name = "Courier New"
        r_code.font.size = Pt(8)

    # ── 分隔页 ──
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（中间部分略）")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.name = "宋体"
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.add_page_break()

    # ── 后2000行 ──
    doc.add_heading("第二部分：源代码后2000行", level=1)

    start_line = total - 2000 + 1
    for i, line in enumerate(back, start_line):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1

        r_num = p.add_run(f"{i:5d} ")
        r_num.font.name = "Courier New"
        r_num.font.size = Pt(7)
        r_num.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        display = line.rstrip("\n").replace("\t", "    ")
        r_code = p.add_run(display if display else " ")
        r_code.font.name = "Courier New"
        r_code.font.size = Pt(8)

    doc.save(OUT)
    print(f"\n已生成: {OUT}")
    print(f"文件大小: {os.path.getsize(OUT) / 1024:.0f} KB")


if __name__ == "__main__":
    main()
