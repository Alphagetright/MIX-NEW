#!/usr/bin/env python3
"""Rebuild pipeline doc using python-docx API - fixed insertion order."""
import os, shutil, copy
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
from docx.oxml.ns import qn

SRC = r'C:\Users\Administrator\Desktop\新建文件夹 (14)\软著_数据生产管线_软件说明书.docx'
DST = r'C:\Users\Administrator\Desktop\pipeline_doc.docx'
IMG = r'C:\Users\Administrator\Desktop\All Mix'

shutil.copy2(SRC, DST)
doc = Document(DST)

# ─── Find sections ───
paras = list(enumerate(p.text.strip() for p in doc.paragraphs))
body_s18 = next(i for i, t in paras if t == '十八、CLI接口层' and i > 100)
body_s19 = next(i for i, t in paras if t == '十九、Web前端模块' and i > 100)
print(f'Section 18 at {body_s18}, Section 19 at {body_s19}')

# ─── Find old image paragraphs in section 19 (by rId) ───
old_img_indices = []
for i, p in enumerate(doc.paragraphs):
    if i < body_s19: continue
    if i > body_s19 + 100: break
    for b in p._element.findall('.//'+qn('a:blip')):
        rid = b.get(qn('r:embed'))
        if rid in ('rId12', 'rId13'):
            old_img_indices.append(i)
print(f'Old image paragraphs at indices: {old_img_indices}')

# ─── Remove section 18 paragraphs (REVERSE order) ───
for i in range(body_s19 - 1, body_s18 - 1, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
print('Removed section 18 paragraphs')

# ─── Helper: batch insert XML elements before section 19 heading ───
def make_heading1(text):
    p = etree.SubElement(doc.styles['Heading 1']._element, qn('w:p'))  # won't work
    return None

# Better approach: Build all new paragraphs as XML, insert in one batch
def build_para(children_xml, style=None, align=None, indent=None):
    """Build a <w:p> element with given children XML string."""
    p = etree.SubElement(etree.Element('dummy'), qn('w:p'))
    pPr = etree.SubElement(p, qn('w:pPr'))
    if style:
        style_elem = etree.SubElement(pPr, qn('w:pStyle'))
        style_elem.set(qn('w:val'), style)
    if align:
        jc = etree.SubElement(pPr, qn('w:jc'))
        jc.set(qn('w:val'), align)
    if indent:
        ind = etree.SubElement(pPr, qn('w:ind'))
        ind.set(qn('w:firstLine'), str(int(indent * 567)))  # cm to twips
    return p

def text_para(text, bold=False, sz=21, font='宋体', style=None, color=None, align=None, indent=None, space_preserve=True):
    """Create a w:p element with a single w:r containing w:t."""
    # Use etree HTML-like construction
    p_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    p_xml += '<w:pPr>'
    if style:
        p_xml += f'<w:pStyle w:val="{style}"/>'
    if align:
        p_xml += f'<w:jc w:val="{align}"/>'
    if indent:
        twips = int(indent * 567)
        p_xml += f'<w:ind w:firstLine="{twips}"/>'
    p_xml += '<w:rPr>'
    p_xml += f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}"/>'
    if bold:
        p_xml += '<w:b/>'
    if color:
        p_xml += f'<w:color w:val="{color}"/>'
    p_xml += f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/>'
    p_xml += '</w:rPr></w:pPr>'
    sp = ' xml:space="preserve"' if space_preserve else ''
    p_xml += f'<w:r><w:rPr><w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}"/>'
    if bold:
        p_xml += '<w:b/>'
    if color:
        p_xml += f'<w:color w:val="{color}"/>'
    p_xml += f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/>'
    p_xml += f'</w:rPr><w:t{sp}>{text}</w:t></w:r></w:p>'
    return p_xml

def img_para(img_path, rid, docpr_id, cx_inches=5.0):
    """Create a w:p element with a centered image."""
    from docx.shared import Emu
    cx = int(cx_inches * 914400)  # inches to EMU
    # Actually get image dimensions
    from PIL import Image
    with Image.open(img_path) as img:
        w, h = img.size
    # Maintain aspect ratio
    cy = int(cx * h / w)
    fname = os.path.basename(img_path)

    return f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
  <w:pPr><w:jc w:val="center"/></w:pPr>
  <w:r>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="{docpr_id}" name="{fname}"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
        <a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:pic>
            <pic:nvPicPr><pic:cNvPr id="0" name="{fname}"/><pic:cNvPicPr/></pic:nvPicPr>
            <pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>
            <pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>
          </pic:pic>
        </a:graphicData></a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>'''

# ─── Build new section 18 content ───
N = 'w'  # namespace prefix in the doc

# Find section 19 heading element in the XML tree
s19_elem = None
for p in doc.paragraphs:
    if p.text.strip() == '十九、Web前端模块':
        s19_elem = p._element
        break

if s19_elem is None:
    print('ERROR: Could not find section 19 heading!')
    exit(1)

# Build all new paragraphs as XML strings (in correct display order)
new_content = []

# Section heading
new_content.append(text_para('十八、系统运行截图', bold=True, sz=28, font='黑体', style='Heading1'))

# 18.1
new_content.append(text_para('18.1 管线CLI命令帮助', bold=True, sz=24, font='黑体', style='Heading2'))
new_content.append(text_para('系统提供完整的CLI命令接口，支持run（执行标注管线）、validate（校验数据）、report（生成报告）、clean（清理缓存）、config（管理配置）、stats（运行统计）七个子命令。通过demo.py --help查看所有命令及参数说明。', indent=0.75))
new_content.append(text_para('图1： 管线CLI命令帮助', bold=True, sz=24, font='宋体', color='2C3E50', align='center'))
new_content.append(img_para(f'{IMG}/pipeline_cli_help.png', 'rId14', 14, 5.0))

# 18.2
new_content.append(text_para('18.2 管线运行过程', bold=True, sz=24, font='黑体', style='Heading2'))
new_content.append(text_para('执行run命令启动标注管线，系统按"输入读取→文本预处理→AI标注引擎→结果解析与校验→数据输出"五步流程自动处理。输入支持TXT/JSON/CSV/MD格式，AI引擎完成标注后经过多层校验确保质量。', indent=0.75))
new_content.append(text_para('图2： 管线运行过程', bold=True, sz=24, font='宋体', color='2C3E50', align='center'))
new_content.append(img_para(f'{IMG}/pipeline_run.png', 'rId15', 15, 5.0))

# 18.3
new_content.append(text_para('18.3 标注数据输出', bold=True, sz=24, font='黑体', style='Heading2'))
new_content.append(text_para('管线输出结构化JSON格式的标注数据，包含运行元信息、诗歌原文、多维标注结果（意象分析/情感分析/结构分析）及质量评分。支持JSON/CSV/XML三种导出格式。', indent=0.75))
new_content.append(text_para('图3： 标注数据输出', bold=True, sz=24, font='宋体', color='2C3E50', align='center'))
new_content.append(img_para(f'{IMG}/pipeline_output.png', 'rId16', 16, 5.0))

# Page break before section 19
new_content.append('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:r><w:br w:type="page"/></w:r></w:p>')

# Insert before section 19 heading (in reverse order so first appears first)
parent = s19_elem.getparent()
idx = list(parent).index(s19_elem)
for xml_str in reversed(new_content):
    elem = etree.fromstring(xml_str)
    parent.insert(idx, elem)

print(f'Inserted new section 18 ({len(new_content)} paragraphs)')

# ─── Replace old images in section 19 ───
# Now find the old images again (indices may have shifted)
old_imgs = []
for i, p in enumerate(doc.paragraphs):
    if i > 300: break
    for b in p._element.findall('.//'+qn('a:blip')):
        rid = b.get(qn('r:embed'))
        if rid in ('rId12', 'rId13'):
            old_imgs.append((i, rid, p._element))

print(f'Found {len(old_imgs)} old images to replace: {[(i, rid) for i,rid,_ in old_imgs]}')

rid_map = {
    'rId12': [  # was the single old web screenshot in 19.3
        ('图4： 登录页面', 'rId17', f'{IMG}/p_login.png', 17),
        ('login_img', 'rId17', f'{IMG}/p_login.png', 17),
        ('图5： 注册页面', 'rId18', f'{IMG}/p_register.png', 18),
        ('register_img', 'rId18', f'{IMG}/p_register.png', 18),
    ],
    'rId13': [  # was the old API diagram in 19.4
        ('图6： 控制台主页', 'rId19', f'{IMG}/p_dashboard.png', 19),
        ('dash_img', 'rId19', f'{IMG}/p_dashboard.png', 19),
        ('home_img', 'rId20', f'{IMG}/p_home.png', 20),
        ('图7： API接口响应', 'rId21', f'{IMG}/p_api_health.png', 21),
        ('api_img', 'rId21', f'{IMG}/p_api_health.png', 21),
    ]
}

for img_idx, rid, elem in old_imgs:
    replacements = rid_map.get(rid, [])
    parent = elem.getparent()
    elem_idx = list(parent).index(elem)

    # Build replacement XML in REVERSE order (last inserted = first displayed)
    replacement_xmls = []
    for label, new_rid, img_path, docpr_id in replacements:
        if label.endswith('_img'):
            replacement_xmls.append(img_para(img_path, new_rid, docpr_id, 5.0))
        else:
            replacement_xmls.append(text_para(label, bold=True, sz=24, font='宋体', color='2C3E50', align='center'))

    # Insert in reverse order so first xml appears first
    for xml_str in reversed(replacement_xmls):
        new_elem = etree.fromstring(xml_str)
        parent.insert(elem_idx, new_elem)

    # Remove old element
    parent.remove(elem)
    print(f'Replaced old {rid} with {len(replacements)} new paragraphs')

# ─── Save ───
doc.save(DST)
print(f'Saved to {DST}')

# Quick verify
doc2 = Document(DST)
figs = [p.text.strip() for p in doc2.paragraphs if '图' in p.text and ('：' in p.text)]
print(f'Figure labels found: {figs}')
