# -*- coding: utf-8 -*-
"""Extract summaries from three soft-copyright docx files."""
import os, sys, glob
from docx import Document

# Force UTF-8 for output
sys.stdout.reconfigure(encoding='utf-8') if hasattr(sys.stdout, 'reconfigure') else None

folder = r"C:\Users\Administrator\Desktop\新建文件夹 (14)"
files = sorted(glob.glob(os.path.join(folder, "*.docx")))

labels = {
    "古典诗歌文本结构化标注生产系统": "一、古典诗歌文本结构化标注生产系统",
    "唐诗意象数据CLI运维管理系统": "二、唐诗意象数据CLI运维管理系统",
    "软件说明书": "三、唐诗意象智能分析系统（Web端）",
}

for fp in files:
    basename = os.path.basename(fp)
    label = "未知"
    for key, val in labels.items():
        if key in basename:
            label = val
            break

    # Only process the three main documents (skip source-code files)
    if label == "未知":
        continue

    print()
    print("=" * 60)
    print(f"  {label}")
    print(f"  文件: {basename}")
    print("=" * 60)

    doc = Document(fp)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Find project background / introduction
    for i, t in enumerate(paras):
        if "项目背景" in t or "项目目标" in t or "功能概述" in t or "系统定位" in t or "项目简介" in t:
            print(f"\n  【{t}】")
            if i + 1 < len(paras):
                print(f"  {paras[i+1][:500]}")
            break
    else:
        # Fallback: print first meaningful paragraph
        for t in paras:
            if len(t) > 50 and "项目" in t:
                print(f"\n  {t[:500]}")
                break

    # Print core features
    print(f"\n  【核心功能】")
    found_feature = False
    feature_count = 0
    for i, t in enumerate(paras):
        if "功能" in t and ("概述" in t or "介绍" in t or "设计" in t):
            found_feature = True
            continue
        if found_feature and feature_count < 5:
            if t.startswith("•") or t.startswith("-") or t.startswith("1.") or t.startswith("2."):
                print(f"    {t[:200]}")
                feature_count += 1
            elif len(t) > 30 and feature_count == 0:
                print(f"    {t[:300]}")
                feature_count += 1
                break

    # Print document structure (headings)
    print(f"\n  【文档章节】")
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and 'Heading' in p.style.name and count < 30:
            print(f"    {t[:70]}")
            count += 1
